"""Output, revision, and deliverable module runners."""

from __future__ import annotations

import traceback
from pathlib import Path

from agent.document_store import get_document, resolve_agent_context
from agent.executor_common import (
    _fail_result,
    _get_solve_data,
    _module_fingerprint,
    _ok_result,
)
from agent.executor_dirty import mark_dirty_from_revise
from agent.types import ModuleResult
from config import TEMP_DIR, UML_RENDER_OK
from log_util import loge
from modules.fill_report import do_fill
from modules.uml import format_render_summary, render_uml_diagrams


def _run_render_uml(ctx: dict, params: dict) -> ModuleResult:
    if not UML_RENDER_OK:
        return _fail_result("render_uml", "UML 模块不可用", params)
    solve = _get_solve_data(ctx)
    parsed = solve.get("parsed") or {}
    diagrams = parsed.get("diagrams") or []
    if not diagrams:
        skipped = {"images_b64": [], "skipped": True, "kind_stats": {}}
        skipped["summary"] = format_render_summary(skipped)
        return _ok_result("render_uml", skipped, params)
    try:
        out = render_uml_diagrams(
            diagrams,
            allow_online=params.get("allow_online", True),
            code=solve.get("code") or parsed.get("code") or "",
            language=solve.get("language") or parsed.get("language") or "java",
        )
        ok = bool(out.get("success"))
        if ok:
            return _ok_result("render_uml", out, params)
        return ModuleResult(
            ok=False,
            data=out,
            logs=out.get("errors") or [out.get("summary") or "图表渲染未完全成功"],
            fingerprint=_module_fingerprint("render_uml", params, out),
            cacheable=False,
        )
    except Exception as e:
        return _fail_result("render_uml", str(e), params)


def _run_revise_answer(ctx: dict, params: dict) -> ModuleResult:
    from agent.executor_dirty import (
        apply_revise_to_module_results,
        code_status_from_ctx,
        mark_dirty_from_revise,
    )
    from modules.revise_answer import revise_answer

    solve_mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    if not solve_mr.get("ok"):
        return _fail_result("revise_answer", "无 solve_lab 结果可修订", params)

    solve_data = dict(solve_mr.get("data") or {})
    parsed = dict(solve_data.get("parsed") or solve_data)
    verification = ctx.get("verification_report") or {}
    dirty_fields = (ctx.get("dirty_fields") or {}).get("solve_lab") or []

    scope = params.get("scope")
    if not scope:
        if code_status_from_ctx(ctx) == "verified":
            scope = ["steps", "result", "summary"]
        elif dirty_fields and dirty_fields != ["full"]:
            scope = dirty_fields
        else:
            scope = ["full"]

    failed_checks = [c for c in verification.get("checks", []) if not c.get("ok")]
    feedback = (params.get("feedback") or "").strip() or "; ".join(
        f"{c.get('id')}: {c.get('message')}" for c in failed_checks[:5]
    ) or "请根据校验结果改进答案文字"

    try:
        rev = revise_answer(
            ctx["settings"],
            parsed=parsed,
            report_excerpt=ctx.get("planner_input_text") or ctx.get("report_text") or "",
            scope=scope,
            feedback=feedback,
            verification_report=verification,
            format_spec=ctx.get("format_spec"),
        )
        merged = dict(solve_data)
        merged["parsed"] = rev.get("parsed") or parsed
        rev_parsed = merged["parsed"]
        for key in (
            "steps_analysis",
            "result_description",
            "summary",
            "expected_output",
            "code",
            "code_files",
            "diagrams",
            "language",
            "main_file",
        ):
            if key in rev_parsed:
                merged[key] = rev_parsed[key]
        changed = rev.get("changed_fields") or (scope if isinstance(scope, list) else [scope])
        apply_revise_to_module_results(ctx, merged, changed_fields=changed)
        mark_dirty_from_revise(
            ctx,
            changed_fields=changed,
            scope=scope if isinstance(scope, list) else [scope],
        )
        return _ok_result("revise_answer", rev, params)
    except Exception as e:
        return _fail_result("revise_answer", str(e), params)


def _run_fix_diagrams(ctx: dict, params: dict) -> ModuleResult:
    from modules.diagram_verify import verify_diagrams
    from modules.fix_diagrams import fix_diagrams

    solve_mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    if not solve_mr.get("ok"):
        return _fail_result("fix_diagrams", "无 solve_lab 结果可修复", params)

    solve_data = dict(solve_mr.get("data") or {})
    parsed = dict(solve_data.get("parsed") or {})
    if not parsed.get("diagrams"):
        return _fail_result("fix_diagrams", "答案中无 diagrams 可修复", params)

    render_data = ((ctx.get("module_results") or {}).get("render_uml") or {}).get("data") or {}
    validation = render_data.get("validation") or verify_diagrams(
        solve_data,
        render_result=render_data if render_data.get("errors") is not None else None,
    )
    issues = validation.get("issues") or []
    feedback = (params.get("feedback") or "").strip()

    try:
        fixed = fix_diagrams(
            ctx["settings"],
            parsed=parsed,
            report_excerpt=ctx.get("planner_input_text") or ctx.get("report_text") or "",
            feedback=feedback,
            issues=issues,
            verification_report=ctx.get("verification_report"),
            format_spec=ctx.get("format_spec"),
        )
        merged = dict(solve_data)
        merged["parsed"] = fixed.get("parsed") or parsed
        from agent.executor_dirty import apply_revise_to_module_results

        apply_revise_to_module_results(
            ctx,
            merged,
            changed_fields=fixed.get("changed_fields") or ["diagrams"],
        )
        mark_dirty_from_revise(ctx, changed_fields=["diagrams"], scope=["diagrams"])
        return _ok_result(
            "fix_diagrams",
            {
                "changed_fields": fixed.get("changed_fields") or ["diagrams"],
                "diagrams": (merged.get("parsed") or {}).get("diagrams"),
                "issue_count": fixed.get("issue_count", 0),
            },
            params,
        )
    except Exception as e:
        return _fail_result("fix_diagrams", str(e), params)


def _verify_fill_output(output_path: str, answers: list, mode: str) -> tuple[bool, str]:
    """Read back the filled docx and verify it contains actual answer content.

    Returns (ok, note). A thin docx with only section headings means fill_lab
    silently matched nothing — the caller should treat this as a warning.
    """
    try:
        from docx import Document

        doc = Document(output_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
        full_text = "\n".join(parts).strip()
        char_count = len(full_text)

        if char_count < 100:
            return False, (
                f"输出文档仅 {char_count} 字符，fill_report 可能未生效"
                f"（文档节标题与答案不匹配？）"
            )

        # Check at least one content word from the answers appears in the output
        answer_words = set()
        for ans in answers:
            parsed = ans.get("parsed") or {}
            for key in ("steps_analysis", "result_description", "summary"):
                text = parsed.get(key, "")
                for w in text.split():
                    if len(w) >= 6 and w.isalpha() and not w.isascii():
                        answer_words.add(w)
        if answer_words:
            sample = list(answer_words)[:5]
            if not any(w in full_text for w in sample):
                return False, "文档中未找到答案关键字，fill_report 可能未生效（内容未匹配到任何节）"

        return True, f"文档 {char_count} 字符，填充验证通过"
    except Exception as e:
        return False, f"验证填充结果时出错: {e}"


def _run_fill_report(ctx: dict, params: dict) -> ModuleResult:
    from document.pdf_export import prepare_fill_docx_for_fill

    doc_ids = ctx.get("document_ids") or []
    file_path = ""
    file_name = "report.docx"
    fill_body_text = ctx.get("report_text") or ""
    metadata = dict(ctx.get("metadata") or {})
    if ctx.get("assignment_text") and not metadata.get("assignment_text"):
        metadata["assignment_text"] = ctx["assignment_text"]
    if ctx.get("planner_input_text") and not metadata.get("planner_input_text"):
        metadata["planner_input_text"] = ctx["planner_input_text"]
    fill_info = ctx.get("fill_target_info")

    if doc_ids:
        if not fill_info:
            try:
                doc_ctx = resolve_agent_context(doc_ids)
                fill_info = doc_ctx.get("fill_target_info") or {}
                fill_body_text = (
                    doc_ctx.get("fill_target", {}).get("full_text")
                    or doc_ctx.get("report_text")
                    or fill_body_text
                )
                metadata = doc_ctx.get("metadata") or metadata
            except ValueError:
                fill_info = {}
        rec = get_document(doc_ids[0])
        if rec and not fill_info:
            file_path = rec.get("file_path") or ""
            file_name = rec.get("file_name") or file_name
            fill_body_text = rec.get("fill_body_text") or rec.get("report_text") or fill_body_text
            metadata = rec.get("metadata") or metadata
        elif fill_info:
            file_path = fill_info.get("path") or ""
            file_name = fill_info.get("file_name") or file_name

    solve = _get_solve_data(ctx)
    if not solve and (ctx.get("module_results") or {}).get("solve_code_cloze", {}).get("ok"):
        solve = (ctx["module_results"]["solve_code_cloze"]).get("data") or {}
    if not solve and (ctx.get("module_results") or {}).get("solve_theory", {}).get("ok"):
        solve = (ctx["module_results"]["solve_theory"]).get("data") or {}

    answers = []
    if solve:
        ans = dict(solve)
        ans.setdefault("type", "lab_report" if solve.get("parsed") else "theory")
        uml = (ctx.get("module_results") or {}).get("render_uml")
        if uml and uml.get("ok"):
            uml_imgs = (uml.get("data") or {}).get("images_b64") or []
            if uml_imgs:
                ans.setdefault("uml_images_b64", uml_imgs)
        answers.append(ans)

    from modules.deliverable import is_content_only_output_mode

    output_mode = ctx.get("output_mode", "deliverable")
    content_only = is_content_only_output_mode(output_mode)

    if not answers:
        if content_only:
            return _ok_result(
                "fill_report",
                {"success": True, "mode": "answer_only", "note": "答案已显示在界面中，未写入文件"},
                params,
            )
        return _fail_result("fill_report", "无可填充的解题结果", params)

    if content_only:
        return _ok_result(
            "fill_report",
            {
                "success": True,
                "mode": output_mode,
                "note": "答案工作区模式：未写入文件，请自行复制粘贴",
            },
            params,
        )

    if output_mode == "new_document":
        report_text = ctx.get("planner_input_text") or ctx.get("report_text") or ""
        if not report_text:
            return _fail_result("fill_report", "无报告文本，无法生成新文档", params)
        from document.pdf_export import generate_docx_shell
        shell_path = generate_docx_shell(
            report_text,
            metadata=metadata,
            output_path=TEMP_DIR / f"run_{ctx.get('run_id', 'x')}_new_shell.docx",
        )
        out_name = f"run_{ctx.get('run_id', 'x')}_实验报告_已完成.docx"
        out_path = str(TEMP_DIR / out_name)
        try:
            path = do_fill(shell_path, answers, out_path,
                          fill_sections=ctx.get("fill_sections"), metadata=metadata,
                          settings=ctx.get("settings"))
            verified, verify_note = _verify_fill_output(str(path), answers, "new_document")
            if not verified:
                return _fail_result(
                    "fill_report",
                    f"填充后验证失败: {verify_note}。请检查文档节标题格式，或调整 solve_lab 输出使其匹配文档结构后重新 fill_report。",
                    params,
                )
            return _ok_result(
                "fill_report",
                {"output_path": str(path), "success": True, "mode": "new_document"},
                params,
            )
        except Exception as e:
            from log_util import loge
            loge("fill", f"fill_report (new_document) 异常: {e}\n{traceback.format_exc()}")
            return _fail_result("fill_report", f"{type(e).__name__}: {e}", params)

    try:
        inp, resolved = prepare_fill_docx_for_fill(
            Path(file_path) if file_path else None,
            file_name,
            source_format=(fill_info or {}).get("source_format")
            or metadata.get("source_format"),
            fill_body_text=fill_body_text,
            metadata=metadata,
            shell_output_path=TEMP_DIR / f"run_{ctx.get('run_id', 'x')}_shell.docx",
        )
        out_name = f"run_{ctx.get('run_id', 'x')}_{Path(file_name).stem}_已完成.docx"
        out_path = str(TEMP_DIR / out_name)
        path = do_fill(inp, answers, out_path, fill_sections=ctx.get("fill_sections"), metadata=metadata,
                      settings=ctx.get("settings"))
        verified, verify_note = _verify_fill_output(str(path), answers, "fill_original")
        if not verified:
            return _fail_result(
                "fill_report",
                f"填充后验证失败: {verify_note}。请检查文档节标题格式，或调整 solve_lab 输出使其匹配文档结构后重新 fill_report。",
                params,
            )
        data = {"output_path": str(path), "success": True, "fill_target": resolved}
        return _ok_result("fill_report", data, params)
    except Exception as e:
        from log_util import loge
        loge("fill", f"fill_report (fill_original) 异常: {e}\n{traceback.format_exc()}")
        return _fail_result("fill_report", f"{type(e).__name__}: {e}", params)


def _run_present_deliverable(ctx: dict, params: dict) -> ModuleResult:
    """Assemble LabDeliverable (zero LLM). V5 default pipeline tail."""
    from modules.deliverable import build_deliverable

    try:
        if not ((ctx.get("module_results") or {}).get("solve_lab") or {}).get("ok"):
            cloze_mr = (ctx.get("module_results") or {}).get("solve_code_cloze") or {}
            if cloze_mr.get("ok"):
                ctx.setdefault("module_results", {})["solve_lab"] = {
                    "ok": True,
                    "data": cloze_mr.get("data") or {},
                }
        dlv = build_deliverable(ctx)
    except ValueError as e:
        return _fail_result("present_deliverable", str(e), params)
    except Exception as e:
        loge("deliverable", f"present_deliverable 异常: {e}\n{traceback.format_exc()}")
        return _fail_result("present_deliverable", f"{type(e).__name__}: {e}", params)

    ctx["deliverable"] = dlv
    return _ok_result(
        "present_deliverable",
        {"deliverable": dlv, "deliverable_id": dlv.get("id")},
        params,
    )
