"""
Standard-mode agent executor (Phase 2a.1).
"""

from __future__ import annotations

import base64
import hashlib
import json
import threading
import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable, Optional

from agent.decision_log import append_decision
from agent.document_store import clear_run_temp, get_document, resolve_agent_context
from agent.fallback import fallback_to_solve
from agent.planner import (
    MAX_CONSECUTIVE_FAILURES,
    replan_incremental,
    verify_plan_fingerprint,
)
from agent.run_control import (
    emit_event,
    is_cancelled,
    map_api_error,
    pop_retry_module,
    release_run,
    set_last_error,
)
from agent.executor_dirty import (
    compute_sub_fingerprints,
    mark_dirty_from_revise,
    note_module_completed,
    should_rerun_module,
)
from agent.types import ModuleResult, PlanStep
from config import IDE_RENDER_OK, TEMP_DIR, UML_RENDER_OK
from log_util import loge, logi
from modules.fill_report import do_fill
from modules.run_code import execute_code, get_java_exe
from modules.screenshot import paths_to_b64, render_ide_screenshot_file, render_terminal_image
from modules.solve_lab import solve_lab
from modules.uml import format_render_summary, render_uml_diagrams


def _module_fingerprint(module: str, params: dict, data: dict) -> str:
    payload = {"module": module, "params": params, "keys": sorted(data.keys())[:20]}
    digest = hashlib.sha256(
        json.dumps(payload, sort_keys=True, ensure_ascii=False).encode("utf-8")
    ).hexdigest()
    return f"sha256:{digest[:24]}"


def _ok_result(module: str, data: dict, params: dict | None = None) -> ModuleResult:
    subs: dict[str, str] = {}
    if module == "solve_lab":
        subs = compute_sub_fingerprints(data)
    return ModuleResult(
        ok=True,
        data=data,
        logs=[],
        fingerprint=_module_fingerprint(module, params or {}, data),
        sub_fingerprints=subs,
        cacheable=True,
    )


def _fail_result(module: str, message: str, params: dict | None = None) -> ModuleResult:
    return ModuleResult(
        ok=False,
        data={"error": message},
        logs=[message],
        fingerprint=_module_fingerprint(module, params or {}, {"error": message}),
        cacheable=False,
    )


def _get_solve_data(ctx: dict) -> dict:
    mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    if mr.get("ok"):
        return mr.get("data") or {}
    return {}


def _build_error_meta(result_data: dict, module: str) -> dict | None:
    """Build error_meta for SSE progress events from module result data."""
    if module != "run_code":
        return None
    degraded = result_data.get("degraded") or result_data.get("degraded_reason")
    meta = {
        "degraded": bool(degraded),
        "degraded_reason": result_data.get("degraded_reason", ""),
    }
    if result_data.get("error_category"):
        meta["category"] = result_data["error_category"]
    return meta if (meta["degraded"] or meta.get("category")) else None


def _run_solve_lab(ctx: dict, params: dict) -> ModuleResult:
    settings = ctx["settings"]
    question = dict(ctx.get("question") or {})
    question["type"] = "lab_report"
    # Use planner_input_text when available — it includes assignment + fill_target
    # for combined documents, while report_text is only the template half.
    full_text = ctx.get("planner_input_text") or ctx.get("report_text") or question.get("full_text") or ""
    question["full_text"] = full_text
    lang = params.get("language") or (ctx.get("user_profile") or {}).get("default_language", "java")
    question["preferred_lang"] = lang
    include_uml = bool(params.get("include_uml"))
    from agent.skill_store import match_skills

    matched = match_skills(
        {"language": lang, "full_text": full_text, "report_text": full_text}
    )
    fired = ctx.setdefault("skills_fired", [])
    for skill in matched:
        sid = skill.get("id")
        if sid and sid not in fired:
            fired.append(sid)
    fmt_spec = ctx.get("format_spec")
    if fmt_spec:
        from agent.template_analyzer import to_format_constraints

        question["format_constraints"] = to_format_constraints(fmt_spec)
    try:
        from modules.user_constraints import constraints_from_ctx

        user_constraints = constraints_from_ctx(ctx)
        if user_constraints:
            ctx["user_constraints"] = user_constraints

        def _on_pipeline_phase(event: dict) -> None:
            ctx.setdefault("pipeline_phases", []).append(event)

        result = solve_lab(
            settings["api_key"],
            settings.get("provider", "deepseek"),
            settings.get("model", "deepseek-chat"),
            question,
            custom_url=settings.get("custom_url") or settings.get("customUrl") or "",
            include_uml=include_uml,
            format_spec=fmt_spec,
            settings=settings,
            user_constraints=user_constraints,
            on_phase=_on_pipeline_phase,
        )
        if result.get("solve_session"):
            ctx["solve_session"] = result["solve_session"]
        if result.get("pipeline_meta"):
            ctx["pipeline_meta"] = result["pipeline_meta"]
        return _ok_result("solve_lab", result, params)
    except Exception as e:
        return _fail_result("solve_lab", str(e), params)


def _run_solve_theory(ctx: dict, params: dict) -> ModuleResult:
    from llm_client import call_ai

    settings = ctx["settings"]
    txt = ctx.get("planner_input_text") or ctx.get("report_text") or ""
    question = {
        "type": "theory",
        "content": txt,
        "full_text": txt,
        "preferred_lang": params.get("language", "python"),
    }
    try:
        result = call_ai(
            settings["api_key"],
            settings.get("provider", "deepseek"),
            settings.get("model", "deepseek-chat"),
            question,
            custom_url=settings.get("custom_url") or "",
        )
        return _ok_result("solve_theory", result, params)
    except Exception as e:
        return _fail_result("solve_theory", str(e), params)


def _guess_filename_from_lang(language: str) -> str:
    ext = {"python": ".py", "java": ".java", "c": ".c", "cpp": ".cpp",
           "javascript": ".js"}.get((language or "python").lower(), ".py")
    return f"main{ext}"


def _run_run_code(ctx: dict, params: dict) -> ModuleResult:
    """Execute code with preflight check → classify → fix → retry → degrade loop."""
    from modules.fix_code import apply_fix_to_solve_data, fix_code_from_error
    from modules.preflight import _check_execution_pattern
    from modules.run_code import classify_run_error, execute_code, execute_multi_file

    solve = _get_solve_data(ctx)
    code_files = solve.get("code_files") or []
    main_file = solve.get("main_file") or ""
    code = solve.get("code") or ""
    language = solve.get("language") or params.get("language") or "java"

    # Normalize: single code string → code_files
    if not code_files and code:
        main_file = main_file or _guess_filename_from_lang(language)
        code_files = [{"name": main_file, "code": code}]
    if not code_files:
        return _fail_result("run_code", "无代码可运行", params)
    if not main_file:
        main_file = code_files[0].get("name", "main.py")

    # Preflight: detect execution pattern in all files
    combined_code = "\n".join(f.get("code", "") for f in code_files)
    pattern_check = _check_execution_pattern(combined_code, language)
    pattern = pattern_check.get("pattern", "script")
    pattern_ok = pattern_check.get("ok", True)

    is_multi_file = len(code_files) > 1

    # Patterns that can't run as-is — skip execution entirely, go direct to fix
    if not pattern_ok:
        logi("executor", f"run_code skipped: pattern={pattern} → fix_code")
        # In ReAct mode, surface the preflight failure to the LLM directly
        # so it can observe the specific reason and call fix_code itself.
        # Standard/deep mode keeps the internal auto-fix loop.
        if ctx.get("run_mode") == "react":
            return _fail_result("run_code", pattern_check["message"], params)
        category = {
            "web_server": "timeout_blocking",
            "interactive": "timeout_blocking",
            "jsp_template": "compile_error",
            "jsp_tags": "compile_error",
            "html_in_java": "compile_error",
            "servlet_api": "compile_error",
            "emoji_in_code": "compile_error",
        }.get(pattern, "timeout_blocking")
        result = _fix_and_retry(ctx, params, code or combined_code, language,
                                pattern_check["message"],
                                category=category, pattern=pattern)
        if result.get("data") and isinstance(result["data"], dict):
            result["data"].setdefault("error_category", category)
        return result

    if language == "java" and not get_java_exe():
        return _fail_result("run_code", "需要 JRE", params)

    try:
        if is_multi_file:
            output, is_error = execute_multi_file(code_files, language, main_file)
        else:
            output, is_error = execute_code(code_files[0]["code"], language)
    except Exception as e:
        return _fail_result("run_code", str(e), params)

    if not is_error:
        return _ok_result(
            "run_code",
            {"output": output, "error": False, "is_error": False, "degraded": False,
             "code_files": code_files, "main_file": main_file},
            params,
        )

    # Classify and fix
    error_class = classify_run_error(output, pattern)
    category = error_class.get("category", "runtime_exception")
    error_msg = f"{error_class.get('message', '')}: {output[:200]}"

    return _fix_and_retry(ctx, params, language, error_msg,
                          code=code, code_files=code_files, main_file=main_file,
                          category=category, pattern=pattern)


def _regenerate_code(ctx: dict, params: dict, language: str,
                     error_msg: str, category: str, retry_count: int) -> ModuleResult:
    """Abandon incremental fixes — regenerate answer from scratch with error context.

    When the same error category persists through 2+ fix_code attempts,
    patching is making things worse. Re-solve with the accumulated failure
    knowledge injected as a stern constraint.
    """
    from modules.fix_code import apply_fix_to_solve_data
    from modules.run_code import execute_code, execute_multi_file

    logi("executor", f"regenerating code (same {category} x2+, retry={retry_count})")

    # Inject error as a hard constraint into the question
    question = dict(ctx.get("question") or {})
    question["type"] = "lab_report"
    full_text = ctx.get("planner_input_text") or ctx.get("report_text") or question.get("full_text") or ""
    question["full_text"] = full_text + (
        f"\n\n【上次代码被拒绝 — 必须避免】"
        f"上轮生成的 {language} 代码编译/运行失败（{category}），经过多次修复仍无法通过。"
        f"请完全放弃上一轮的代码思路，从零重新生成。"
        f"错误信息: {error_msg[:500]}"
        f"\n特别注意：确保代码是纯 {language} 语法，不要混入其他语言，不要引入不可用的库。"
    )
    question["preferred_lang"] = language

    settings = ctx["settings"]
    try:
        result = solve_lab(
            settings["api_key"],
            settings.get("provider", "deepseek"),
            settings.get("model", "deepseek-chat"),
            question,
            custom_url=settings.get("custom_url") or settings.get("customUrl") or "",
            include_uml=False,
            format_spec=ctx.get("format_spec"),
        )
    except Exception as e:
        return _degrade_run_code(ctx, params, category, f"regenerate failed: {e}")

    ctx.setdefault("module_results", {})["solve_lab"] = _ok_result(
        "solve_lab", result, params
    )
    mark_dirty_from_revise(ctx, changed_fields=["code", "language"], scope=["code"])

    # Re-run the regenerated code
    code_files = result.get("code_files") or []
    main_file = result.get("main_file") or ""
    code = result.get("code") or ""
    new_lang = result.get("language") or language
    try:
        if code_files:
            output, is_error = execute_multi_file(code_files, new_lang, main_file)
        else:
            output, is_error = execute_code(code, new_lang)
        if not is_error:
            return _ok_result(
                "fix_code",
                {"fixed": True, "language": new_lang, "retries": retry_count + 1,
                 "regenerated": True},
                params,
            )
        # Still failing after regeneration → degrade
        from modules.run_code import classify_run_error
        new_class = classify_run_error(output, "")
        return _degrade_run_code(
            ctx, params,
            new_class.get("category", category),
            f"regenerated still fails: {output[:200]}",
        )
    except Exception as e:
        return _degrade_run_code(ctx, params, category, f"regenerated run error: {e}")


def _fix_and_retry(ctx: dict, params: dict, language: str,
                   error_msg: str, *, code: str = "",
                   code_files: list[dict] | None = None,
                   main_file: str = "",
                   category: str = "", pattern: str = "",
                   retry_count: int = 0,
                   same_error_count: int = 0) -> ModuleResult:
    """Fix code → retry → regenerate → degrade loop.

    When the same error category persists through 2 fix_code rounds,
    incremental patching is abandoned in favour of a fresh solve_lab
    with the error context injected as a hard constraint.
    """
    from modules.fix_code import apply_fix_to_solve_data, fix_code_from_error
    from modules.run_code import execute_code, execute_multi_file

    MAX_FIX_RETRIES = 3
    REGEN_THRESHOLD = 2  # Same error category × N consecutive rounds → regenerate

    if retry_count >= MAX_FIX_RETRIES:
        return _degrade_run_code(ctx, params, category, error_msg)

    # Same error keeps coming back — patching is making things worse
    if same_error_count >= REGEN_THRESHOLD:
        return _regenerate_code(ctx, params, language, error_msg, category, retry_count)

    solve_mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    solve_data = dict(solve_mr.get("data") or {})
    files = code_files or []
    old_code = code or (files[0].get("code", "") if files else "")

    try:
        fix = fix_code_from_error(
            ctx["settings"],
            code=code,
            code_files=files if files else None,
            main_file=main_file,
            language=language,
            error_output=str(error_msg),
            report_excerpt=ctx.get("planner_input_text") or ctx.get("report_text") or "",
            category=category,
            pattern=pattern,
        )
        new_code = fix.get("code") or old_code
        new_files = fix.get("code_files") or []
        new_main = fix.get("main_file") or main_file
        new_lang = fix.get("language") or language

        is_multi = bool(new_files)

        # If fix didn't actually change code → count as same error
        if new_code.strip() == old_code.strip() and not new_files and retry_count > 0:
            return _fix_and_retry(ctx, params, language, error_msg,
                                  code=code, code_files=files, main_file=main_file,
                                  category=category, pattern=pattern,
                                  retry_count=retry_count + 1,
                                  same_error_count=same_error_count + 1)

        updated = apply_fix_to_solve_data(solve_data, fix)
        ctx.setdefault("module_results", {})["solve_lab"] = _ok_result(
            "solve_lab", updated, params
        )
        mark_dirty_from_revise(ctx, changed_fields=["code", "language"], scope=["code"])

        # Re-run with fixed code
        try:
            if is_multi:
                output, is_error = execute_multi_file(new_files, new_lang, new_main)
            else:
                output, is_error = execute_code(new_code, new_lang)
            if not is_error:
                return _ok_result(
                    "fix_code",
                    {"fixed": True, "language": new_lang, "retries": retry_count + 1},
                    params,
                )
            # Still failing → check if same error, then retry or regenerate
            from modules.run_code import classify_run_error
            new_class = classify_run_error(output, pattern)
            new_category = new_class.get("category", category)
            repeat = same_error_count + 1 if new_category == category else 0
            return _fix_and_retry(
                ctx, params, new_lang,
                f"{new_class.get('message', '')}: {output[:200]}",
                code=new_code, code_files=new_files, main_file=new_main,
                category=new_category,
                pattern=pattern,
                retry_count=retry_count + 1,
                same_error_count=repeat,
            )
        except Exception as e:
            return _fix_and_retry(ctx, params, new_lang, str(e),
                                  code=new_code, code_files=new_files,
                                  main_file=new_main,
                                  category=category, pattern=pattern,
                                  retry_count=retry_count + 1,
                                  same_error_count=same_error_count + 1)

    except Exception as e:
        if retry_count + 1 >= MAX_FIX_RETRIES:
            return _degrade_run_code(ctx, params, category, str(e))
        return _fix_and_retry(ctx, params, language, str(e),
                              code=code, code_files=files, main_file=main_file,
                              category=category, pattern=pattern,
                              retry_count=retry_count + 1,
                              same_error_count=same_error_count + 1)


def _degrade_run_code(ctx: dict, params: dict, category: str, error_msg: str) -> ModuleResult:
    """Final degradation: skip code execution, use expected_output as substitute."""
    solve = _get_solve_data(ctx)
    fallback_output = solve.get("parsed", {}).get("expected_output", "") or "(无输出)"
    logi("executor", f"degraded run_code: category={category} error={error_msg[:80]}")

    return _ok_result(
        "run_code",
        {
            "output": fallback_output,
            "error": False,
            "is_error": False,
            "degraded": True,
            "degraded_reason": f"代码执行失败({category}): {error_msg[:200]}",
            "error_category": category,
        },
        params,
    )


def _run_fix_code(ctx: dict, params: dict) -> ModuleResult:
    """LLM fix_code (standalone — called when fix_code is a separate plan step)."""
    from modules.fix_code import apply_fix_to_solve_data, fix_code_from_error

    solve_mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    if not solve_mr.get("ok"):
        return _fail_result("fix_code", "无 solve_lab 结果可修复", params)

    solve_data = dict(solve_mr.get("data") or {})
    run_mr = (ctx.get("module_results") or {}).get("run_code") or {}
    err = (run_mr.get("data") or {}).get("output") or params.get("error") or "未知运行错误"
    code = solve_data.get("code") or ""
    code_files = solve_data.get("code_files") or []
    main_file = solve_data.get("main_file") or ""
    language = solve_data.get("language") or "java"
    if not code and not code_files:
        return _fail_result("fix_code", "无代码可修复", params)

    category = params.get("error_category") or ""

    try:
        fix = fix_code_from_error(
            ctx["settings"],
            code=code,
            code_files=code_files if code_files else None,
            main_file=main_file,
            language=language,
            error_output=str(err),
            report_excerpt=ctx.get("planner_input_text") or ctx.get("report_text") or "",
            category=category,
        )
        updated = apply_fix_to_solve_data(solve_data, fix)
        ctx.setdefault("module_results", {})["solve_lab"] = _ok_result(
            "solve_lab", updated, params
        )
        mark_dirty_from_revise(ctx, changed_fields=["code", "language"], scope=["code"])
        return _ok_result(
            "fix_code",
            {"fixed": True, "language": updated.get("language")},
            params,
        )
    except Exception as e:
        return _fail_result("fix_code", str(e), params)


def _run_screenshot(ctx: dict, params: dict, module: str) -> ModuleResult:
    solve = _get_solve_data(ctx)
    code = solve.get("code") or ""
    language = solve.get("language") or "java"
    run_mr = (ctx.get("module_results") or {}).get("run_code") or {}
    terminal_text = (run_mr.get("data") or {}).get("output") or solve.get("parsed", {}).get(
        "expected_output", ""
    )
    style = params.get("style", "ide")
    try:
        if module == "screenshot_terminal" or style == "terminal":
            img_path = render_terminal_image(terminal_text, "Terminal")
            images_b64 = [base64.b64encode(Path(img_path).read_bytes()).decode()]
        elif IDE_RENDER_OK and code:
            paths = render_ide_screenshot_file(code, terminal_text, language)
            images_b64 = paths_to_b64(paths)
            img_path = paths[0] if paths else ""
        else:
            img_path = render_terminal_image(terminal_text, "Output")
            images_b64 = [base64.b64encode(Path(img_path).read_bytes()).decode()]
        return _ok_result(
            module,
            {"images_b64": images_b64, "image_b64": images_b64[0] if images_b64 else None},
            params,
        )
    except Exception as e:
        return _fail_result(module, str(e), params)


def _run_render_uml(ctx: dict, params: dict) -> ModuleResult:
    if not UML_RENDER_OK:
        return _fail_result("render_uml", "UML 模块不可用", params)
    solve = _get_solve_data(ctx)
    parsed = solve.get("parsed") or {}
    diagrams = parsed.get("diagrams") or []
    if not diagrams:
        skipped = {"images_b64": [], "skipped": True, "kind_stats": {}}
        skipped["summary"] = format_render_summary(skipped)
        return _ok_result("render_uml", skipped, params)
    try:
        out = render_uml_diagrams(
            diagrams,
            allow_online=params.get("allow_online", True),
            code=solve.get("code") or parsed.get("code") or "",
            language=solve.get("language") or parsed.get("language") or "java",
        )
        ok = bool(out.get("success"))
        if ok:
            return _ok_result("render_uml", out, params)
        return ModuleResult(
            ok=False,
            data=out,
            logs=out.get("errors") or [out.get("summary") or "图表渲染未完全成功"],
            fingerprint=_module_fingerprint("render_uml", params, out),
            cacheable=False,
        )
    except Exception as e:
        return _fail_result("render_uml", str(e), params)


def _run_fix_diagrams(ctx: dict, params: dict) -> ModuleResult:
    from modules.diagram_verify import verify_diagrams
    from modules.fix_diagrams import fix_diagrams

    solve_mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    if not solve_mr.get("ok"):
        return _fail_result("fix_diagrams", "无 solve_lab 结果可修复", params)

    solve_data = dict(solve_mr.get("data") or {})
    parsed = dict(solve_data.get("parsed") or {})
    if not parsed.get("diagrams"):
        return _fail_result("fix_diagrams", "答案中无 diagrams 可修复", params)

    render_data = ((ctx.get("module_results") or {}).get("render_uml") or {}).get("data") or {}
    validation = render_data.get("validation") or verify_diagrams(
        solve_data,
        render_result=render_data if render_data.get("errors") is not None else None,
    )
    issues = validation.get("issues") or []
    feedback = (params.get("feedback") or "").strip()

    try:
        fixed = fix_diagrams(
            ctx["settings"],
            parsed=parsed,
            report_excerpt=ctx.get("planner_input_text") or ctx.get("report_text") or "",
            feedback=feedback,
            issues=issues,
            verification_report=ctx.get("verification_report"),
            format_spec=ctx.get("format_spec"),
        )
        merged = dict(solve_data)
        merged["parsed"] = fixed.get("parsed") or parsed
        from agent.executor_dirty import apply_revise_to_module_results

        apply_revise_to_module_results(
            ctx,
            merged,
            changed_fields=fixed.get("changed_fields") or ["diagrams"],
        )
        mark_dirty_from_revise(ctx, changed_fields=["diagrams"], scope=["diagrams"])
        return _ok_result(
            "fix_diagrams",
            {
                "changed_fields": fixed.get("changed_fields") or ["diagrams"],
                "diagrams": (merged.get("parsed") or {}).get("diagrams"),
                "issue_count": fixed.get("issue_count", 0),
            },
            params,
        )
    except Exception as e:
        return _fail_result("fix_diagrams", str(e), params)


def _verify_fill_output(output_path: str, answers: list, mode: str) -> tuple[bool, str]:
    """Read back the filled docx and verify it contains actual answer content.

    Returns (ok, note). A thin docx with only section headings means fill_lab
    silently matched nothing — the caller should treat this as a warning.
    """
    try:
        from docx import Document

        doc = Document(output_path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        parts.append(p.text)
        full_text = "\n".join(parts).strip()
        char_count = len(full_text)

        if char_count < 100:
            return False, (
                f"输出文档仅 {char_count} 字符，fill_report 可能未生效"
                f"（文档节标题与答案不匹配？）"
            )

        # Check at least one content word from the answers appears in the output
        answer_words = set()
        for ans in answers:
            parsed = ans.get("parsed") or {}
            for key in ("steps_analysis", "result_description", "summary"):
                text = parsed.get(key, "")
                for w in text.split():
                    if len(w) >= 6 and w.isalpha() and not w.isascii():
                        answer_words.add(w)
        if answer_words:
            sample = list(answer_words)[:5]
            if not any(w in full_text for w in sample):
                return False, "文档中未找到答案关键字，fill_report 可能未生效（内容未匹配到任何节）"

        return True, f"文档 {char_count} 字符，填充验证通过"
    except Exception as e:
        return False, f"验证填充结果时出错: {e}"


def _run_fill_report(ctx: dict, params: dict) -> ModuleResult:
    from document.pdf_export import prepare_fill_docx_for_fill

    doc_ids = ctx.get("document_ids") or []
    file_path = ""
    file_name = "report.docx"
    fill_body_text = ctx.get("report_text") or ""
    metadata = dict(ctx.get("metadata") or {})
    if ctx.get("assignment_text") and not metadata.get("assignment_text"):
        metadata["assignment_text"] = ctx["assignment_text"]
    if ctx.get("planner_input_text") and not metadata.get("planner_input_text"):
        metadata["planner_input_text"] = ctx["planner_input_text"]
    fill_info = ctx.get("fill_target_info")

    if doc_ids:
        if not fill_info:
            try:
                doc_ctx = resolve_agent_context(doc_ids)
                fill_info = doc_ctx.get("fill_target_info") or {}
                fill_body_text = (
                    doc_ctx.get("fill_target", {}).get("full_text")
                    or doc_ctx.get("report_text")
                    or fill_body_text
                )
                metadata = doc_ctx.get("metadata") or metadata
            except ValueError:
                fill_info = {}
        rec = get_document(doc_ids[0])
        if rec and not fill_info:
            file_path = rec.get("file_path") or ""
            file_name = rec.get("file_name") or file_name
            fill_body_text = rec.get("fill_body_text") or rec.get("report_text") or fill_body_text
            metadata = rec.get("metadata") or metadata
        elif fill_info:
            file_path = fill_info.get("path") or ""
            file_name = fill_info.get("file_name") or file_name

    solve = _get_solve_data(ctx)
    if not solve and (ctx.get("module_results") or {}).get("solve_theory", {}).get("ok"):
        solve = (ctx["module_results"]["solve_theory"]).get("data") or {}

    answers = []
    if solve:
        ans = dict(solve)
        ans.setdefault("type", "lab_report" if solve.get("parsed") else "theory")
        shot = (ctx.get("module_results") or {}).get("screenshot_ide") or (
            ctx.get("module_results") or {}
        ).get("screenshot_terminal")
        if shot and shot.get("ok"):
            imgs = (shot.get("data") or {}).get("images_b64") or []
            if imgs:
                ans["images_b64"] = imgs
        uml = (ctx.get("module_results") or {}).get("render_uml")
        if uml and uml.get("ok"):
            uml_imgs = (uml.get("data") or {}).get("images_b64") or []
            if uml_imgs:
                ans.setdefault("uml_images_b64", uml_imgs)
        answers.append(ans)

    from modules.deliverable import is_content_only_output_mode

    output_mode = ctx.get("output_mode", "deliverable")
    content_only = is_content_only_output_mode(output_mode)

    if not answers:
        if content_only:
            return _ok_result(
                "fill_report",
                {"success": True, "mode": "answer_only", "note": "答案已显示在界面中，未写入文件"},
                params,
            )
        return _fail_result("fill_report", "无可填充的解题结果", params)

    if content_only:
        return _ok_result(
            "fill_report",
            {
                "success": True,
                "mode": output_mode,
                "note": "答案工作区模式：未写入文件，请自行复制粘贴",
            },
            params,
        )

    if output_mode == "new_document":
        report_text = ctx.get("planner_input_text") or ctx.get("report_text") or ""
        if not report_text:
            return _fail_result("fill_report", "无报告文本，无法生成新文档", params)
        from document.pdf_export import generate_docx_shell
        shell_path = generate_docx_shell(
            report_text,
            metadata=metadata,
            output_path=TEMP_DIR / f"run_{ctx.get('run_id', 'x')}_new_shell.docx",
        )
        out_name = f"run_{ctx.get('run_id', 'x')}_实验报告_已完成.docx"
        out_path = str(TEMP_DIR / out_name)
        try:
            path = do_fill(shell_path, answers, out_path,
                          fill_sections=ctx.get("fill_sections"), metadata=metadata,
                          settings=ctx.get("settings"))
            verified, verify_note = _verify_fill_output(str(path), answers, "new_document")
            if not verified:
                return _fail_result(
                    "fill_report",
                    f"填充后验证失败: {verify_note}。请检查文档节标题格式，或调整 solve_lab 输出使其匹配文档结构后重新 fill_report。",
                    params,
                )
            return _ok_result(
                "fill_report",
                {"output_path": str(path), "success": True, "mode": "new_document"},
                params,
            )
        except Exception as e:
            from log_util import loge
            loge("fill", f"fill_report (new_document) 异常: {e}\n{traceback.format_exc()}")
            return _fail_result("fill_report", f"{type(e).__name__}: {e}", params)

    try:
        inp, resolved = prepare_fill_docx_for_fill(
            Path(file_path) if file_path else None,
            file_name,
            source_format=(fill_info or {}).get("source_format")
            or metadata.get("source_format"),
            fill_body_text=fill_body_text,
            metadata=metadata,
            shell_output_path=TEMP_DIR / f"run_{ctx.get('run_id', 'x')}_shell.docx",
        )
        out_name = f"run_{ctx.get('run_id', 'x')}_{Path(file_name).stem}_已完成.docx"
        out_path = str(TEMP_DIR / out_name)
        path = do_fill(inp, answers, out_path, fill_sections=ctx.get("fill_sections"), metadata=metadata,
                      settings=ctx.get("settings"))
        verified, verify_note = _verify_fill_output(str(path), answers, "fill_original")
        if not verified:
            return _fail_result(
                "fill_report",
                f"填充后验证失败: {verify_note}。请检查文档节标题格式，或调整 solve_lab 输出使其匹配文档结构后重新 fill_report。",
                params,
            )
        data = {"output_path": str(path), "success": True, "fill_target": resolved}
        return _ok_result("fill_report", data, params)
    except Exception as e:
        from log_util import loge
        loge("fill", f"fill_report (fill_original) 异常: {e}\n{traceback.format_exc()}")
        return _fail_result("fill_report", f"{type(e).__name__}: {e}", params)


def _run_present_deliverable(ctx: dict, params: dict) -> ModuleResult:
    """Assemble LabDeliverable (zero LLM). V5 default pipeline tail."""
    from modules.deliverable import build_deliverable

    try:
        dlv = build_deliverable(ctx)
    except ValueError as e:
        return _fail_result("present_deliverable", str(e), params)
    except Exception as e:
        loge("deliverable", f"present_deliverable 异常: {e}\n{traceback.format_exc()}")
        return _fail_result("present_deliverable", f"{type(e).__name__}: {e}", params)

    ctx["deliverable"] = dlv
    return _ok_result(
        "present_deliverable",
        {"deliverable": dlv, "deliverable_id": dlv.get("id")},
        params,
    )


_MODULE_RUNNERS = {
    "solve_lab": _run_solve_lab,
    "solve_theory": _run_solve_theory,
    "run_code": _run_run_code,
    "fix_code": _run_fix_code,
    "screenshot_ide": lambda ctx, p: _run_screenshot(ctx, p, "screenshot_ide"),
    "screenshot_terminal": lambda ctx, p: _run_screenshot(ctx, p, "screenshot_terminal"),
    "render_uml": _run_render_uml,
    "fix_diagrams": _run_fix_diagrams,
    "fill_report": _run_fill_report,
    "present_deliverable": _run_present_deliverable,
}


def run_module(ctx: dict, step: PlanStep) -> ModuleResult:
    module = step.get("module") or ""
    params = step.get("params") or {}
    runner = _MODULE_RUNNERS.get(module)
    if not runner:
        return _fail_result(module, f"未知模块: {module}", params)
    return runner(ctx, params)


def execute_standard_run(
    run_id: str,
    ctx: dict,
    steps: list[PlanStep],
    *,
    use_fallback: bool = True,
) -> dict[str, Any]:
    """Execute confirmed steps; emit SSE via run_control."""
    from agent.orchestrator import orchestrator_enabled

    if orchestrator_enabled(ctx):
        return _execute_standard_via_orchestrator(
            run_id, ctx, steps, use_fallback=use_fallback
        )
    return _execute_standard_run_legacy(run_id, ctx, steps, use_fallback=use_fallback)


def _execute_standard_via_orchestrator(
    run_id: str,
    ctx: dict,
    steps: list[PlanStep],
    *,
    use_fallback: bool = True,
) -> dict[str, Any]:
    from agent.orchestrator import RunOrchestrator, RunStepsOptions, finalize_run_payload

    emit = lambda ev: emit_event(run_id, ev)
    ctx["run_id"] = run_id

    def on_decision(entry):
        emit({"type": "decision", **entry})

    append_decision(
        ctx,
        agent="executor",
        decision="run_start",
        target="run",
        reason=f"{len(steps)} steps",
        emit=on_decision,
    )

    orch = RunOrchestrator(run_id, ctx, emit=emit, on_decision=on_decision)
    completed, cancelled = orch.run_steps(steps, options=RunStepsOptions())
    if cancelled:
        release_run(run_id, "cancelled")
        return {"cancelled": True, "run_id": run_id}

    verification = orch.run_verify(auto_remediate=bool(ctx.get("auto_remediate")))
    fill_mr = (ctx.get("module_results") or {}).get("fill_report")
    present_mr = (ctx.get("module_results") or {}).get("present_deliverable")
    deliverable = ctx.get("deliverable") or (present_mr or {}).get("data", {}).get("deliverable")
    final = {
        "run_id": run_id,
        "module_results": {
            k: {"ok": v.get("ok"), "data": v.get("data")}
            for k, v in (ctx.get("module_results") or {}).items()
        },
        "decision_log": ctx.get("decision_log"),
        "verification_report": verification,
        "output_path": (fill_mr or {}).get("data", {}).get("output_path") if fill_mr else None,
        "deliverable": deliverable,
    }

    any_ok = any(
        (ctx.get("module_results") or {}).get(m, {}).get("ok")
        for m in ("solve_lab", "solve_theory")
    )
    if not any_ok and use_fallback:
        try:
            fallback_to_solve(ctx, emit=on_decision)
            final["fallback"] = True
            final["module_results"]["solve_lab"] = ctx["module_results"].get("solve_lab")
        except Exception as e:
            mapped = map_api_error(e)
            emit({"type": "error", **mapped})
            release_run(run_id, "error")
            clear_run_temp(run_id)
            emit({"type": "done", "ok": False, **final})
            return final

    release_run(run_id, "completed")
    clear_run_temp(run_id)
    final = finalize_run_payload(orch, final)
    emit({"type": "done", "ok": True, **final})
    logi("executor", f"run_id={run_id} completed")
    return final


def _execute_standard_run_legacy(
    run_id: str,
    ctx: dict,
    steps: list[PlanStep],
    *,
    use_fallback: bool = True,
) -> dict[str, Any]:
    """Pre-V3-2 inline loop (rollback when use_orchestrator=false)."""
    emit = lambda ev: emit_event(run_id, ev)
    completed_modules: list[str] = []
    results_summary: list[dict] = []
    ctx["run_id"] = run_id

    def on_decision(entry):
        emit({"type": "decision", **entry})

    append_decision(
        ctx,
        agent="executor",
        decision="run_start",
        target="run",
        reason=f"{len(steps)} steps",
        emit=on_decision,
    )

    def _first_pending_index(step_list: list[PlanStep]) -> int:
        for idx, s in enumerate(step_list):
            mod = s.get("module") or ""
            if mod not in completed_modules:
                return idx
        return len(step_list)

    i = 0
    while i < len(steps):
        if is_cancelled(run_id):
            release_run(run_id, "cancelled")
            return {"cancelled": True, "run_id": run_id}

        step = steps[i]
        retry_only = pop_retry_module(run_id)
        module = step.get("module") or ""
        if retry_only and module != retry_only:
            i += 1
            continue

        if not step.get("default_checked", True):
            append_decision(
                ctx,
                agent="executor",
                decision="skip_module",
                target=module,
                reason="用户未勾选",
                emit=on_decision,
            )
            emit(
                {
                    "type": "progress",
                    "module": module,
                    "index": i,
                    "status": "skipped",
                }
            )
            i += 1
            continue

        prior = (ctx.get("module_results") or {}).get(module)
        if prior and prior.get("ok") and module == "solve_lab" and not should_rerun_module(
            ctx, module
        ):
            completed_modules.append(module)
            emit(
                {
                    "type": "progress",
                    "module": module,
                    "index": i,
                    "status": "done",
                    "note": "draft 已完成",
                }
            )
            i += 1
            continue

        if prior and prior.get("ok") and not should_rerun_module(ctx, module):
            append_decision(
                ctx,
                agent="executor",
                decision="reuse_cache",
                target=module,
                reason="子指纹未变 / 非 dirty_modules",
                fingerprint=(prior.get("fingerprint") or "")[:32],
                emit=on_decision,
            )
            completed_modules.append(module)
            emit(
                {
                    "type": "progress",
                    "module": module,
                    "index": i,
                    "status": "done",
                    "note": "复用缓存",
                    "reused": True,
                }
            )
            i += 1
            continue

        emit(
            {
                "type": "progress",
                "module": module,
                "index": i,
                "status": "running",
            }
        )
        append_decision(
            ctx,
            agent="executor",
            decision="run_module",
            target=module,
            reason=step.get("reason") or "",
            emit=on_decision,
        )

        try:
            result = run_module(ctx, step)
        except Exception as e:
            mapped = map_api_error(e)
            result = _fail_result(module, mapped["error"], step.get("params"))
            emit({"type": "error", "module": module, **mapped})

        ctx.setdefault("module_results", {})[module] = result
        ok = bool(result.get("ok"))

        if ok:
            ctx["consecutive_failures"] = 0
            completed_modules.append(module)
            note_module_completed(ctx, module)
        else:
            ctx["consecutive_failures"] = int(ctx.get("consecutive_failures") or 0) + 1
            result_data = result.get("data") or {}
            err_msg = result_data.get("error", "失败")
            set_last_error(run_id, module, err_msg)
            error_meta = _build_error_meta(result_data, module) if module == "run_code" else None
            emit(
                {
                    "type": "progress",
                    "module": module,
                    "index": i,
                    "status": "failed",
                    "error": err_msg,
                    "error_meta": error_meta,
                }
            )

            if ctx["consecutive_failures"] >= MAX_CONSECUTIVE_FAILURES:
                rounds_before = int(ctx.get("replan_rounds") or 0)
                new_plan = replan_incremental(
                    ctx,
                    {
                        "failed_module": module,
                        "error_summary": err_msg,
                        "completed_modules": completed_modules,
                    },
                    emit=on_decision,
                )
                # 仅在实际发生增量 replan 时推送 plan_updated（避免 replan_skipped 刷屏）
                if int(ctx.get("replan_rounds") or 0) > rounds_before:
                    emit(
                        {
                            "type": "plan_updated",
                            "plan_fingerprint": new_plan.get("plan_fingerprint"),
                            "steps": new_plan.get("steps"),
                        }
                    )
                steps = list(new_plan.get("steps") or [])
                ctx["consecutive_failures"] = 0
                i = _first_pending_index(steps)
                continue

            i += 1
            continue

        emit(
            {
                "type": "progress",
                "module": module,
                "index": i,
                "status": "done",
            }
        )
        results_summary.append({"module": module, "ok": True})
        i += 1

    from agent.quality import verify_answer

    verification = verify_answer(ctx)
    ctx["verification_report"] = verification
    emit({"type": "verification", **verification})

    fill_mr = (ctx.get("module_results") or {}).get("fill_report")
    present_mr = (ctx.get("module_results") or {}).get("present_deliverable")
    deliverable = ctx.get("deliverable") or (present_mr or {}).get("data", {}).get("deliverable")
    final = {
        "run_id": run_id,
        "module_results": {
            k: {"ok": v.get("ok"), "data": v.get("data")}
            for k, v in (ctx.get("module_results") or {}).items()
        },
        "decision_log": ctx.get("decision_log"),
        "verification_report": verification,
        "output_path": (fill_mr or {}).get("data", {}).get("output_path") if fill_mr else None,
        "deliverable": deliverable,
    }

    any_ok = any(
        (ctx.get("module_results") or {}).get(m, {}).get("ok")
        for m in ("solve_lab", "solve_theory")
    )
    if not any_ok and use_fallback:
        try:
            fallback_to_solve(ctx, emit=on_decision)
            final["fallback"] = True
            final["module_results"]["solve_lab"] = ctx["module_results"].get("solve_lab")
        except Exception as e:
            mapped = map_api_error(e)
            emit({"type": "error", **mapped})
            release_run(run_id, "error")
            clear_run_temp(run_id)
            emit({"type": "done", "ok": False, **final})
            return final

    release_run(run_id, "completed")
    clear_run_temp(run_id)
    emit({"type": "done", "ok": True, **final})
    logi("executor", f"run_id={run_id} completed")
    return final


def _save_agent_insights(ctx: dict) -> None:
    """Scan module_results for LLM notes and append to AI_INSIGHTS.md."""
    solve_mr = (ctx.get("module_results") or {}).get("solve_lab") or {}
    notes = ((solve_mr.get("data") or {}).get("parsed") or {}).get("notes", "").strip()
    if not notes:
        return
    try:
        from log_util import logi

        insights_path = Path(__file__).resolve().parent.parent.parent.parent / "docs" / "AI_INSIGHTS.md"
        today = datetime.now(timezone.utc).strftime("%Y-%m-%d")
        entry = f"\n## {today}\n\n### 自动记录（来自 AI 解题 notes — Agent 模式）\n\n{notes}\n"
        with open(insights_path, "a", encoding="utf-8") as f:
            f.write(entry)
        logi("insight", f"已保存 {len(notes)} 字 LLM 自述到 AI_INSIGHTS.md (Agent)")
    except Exception:
        pass


def start_run_async(
    run_id: str,
    ctx: dict,
    steps: list[PlanStep],
    *,
    use_fallback: bool = True,
    run_mode: str = "standard",
) -> threading.Thread:
    def _target():
        try:
            from llm_client import reset_llm_call_count

            reset_llm_call_count()
            mode = (run_mode or ctx.get("run_mode") or "standard").lower()
            if mode == "deep":
                from agent.deep_pipeline import execute_deep_run

                execute_deep_run(run_id, ctx, steps, use_fallback=use_fallback)
            elif mode == "react":
                from agent.react_loop import run_react_loop

                run_react_loop(run_id, ctx, steps, use_fallback=use_fallback)
            else:
                execute_standard_run(run_id, ctx, steps, use_fallback=use_fallback)
            _save_agent_insights(ctx)
        except Exception as e:
            loge("executor", str(e))
            mapped = map_api_error(e)
            emit_event(run_id, {"type": "error", **mapped})
            release_run(run_id, "error")
            emit_event(run_id, {"type": "done", "ok": False, "error": mapped["error"]})

    t = threading.Thread(target=_target, daemon=True)
    t.start()
    return t


def retry_single_step(run_id: str, ctx: dict, module_id: str) -> None:
    """Re-run one module from confirmed steps."""
    steps = ctx.get("confirmed_steps") or (ctx.get("plan") or {}).get("steps") or []
    target = next((s for s in steps if s.get("module") == module_id), None)
    if not target:
        raise ValueError(f"计划中无模块: {module_id}")

    from agent.run_control import set_retry_module

    set_retry_module(run_id, module_id)
    emit = lambda ev: emit_event(run_id, ev)

    def on_decision(entry):
        emit({"type": "decision", **entry})

    dirty = list(ctx.get("dirty_modules") or [])
    if module_id not in dirty:
        dirty.append(module_id)
    ctx["dirty_modules"] = dirty

    append_decision(
        ctx,
        agent="executor",
        decision="retry_step",
        target=module_id,
        reason="用户请求重试",
        emit=on_decision,
    )
    result = run_module(ctx, target)
    ctx.setdefault("module_results", {})[module_id] = result
    emit(
        {
            "type": "progress",
            "module": module_id,
            "status": "done" if result.get("ok") else "failed",
            "error": (result.get("data") or {}).get("error"),
        }
    )
