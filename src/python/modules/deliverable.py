"""LabDeliverable assembly and export (V5-0 / V5-2)."""

from __future__ import annotations

import base64
import hashlib
import io
import json
import re
import uuid
import zipfile
from datetime import datetime, timezone
from typing import Any

from config import DOCX_OK


CONTENT_ONLY_OUTPUT_MODES = frozenset({"deliverable", "answer_only"})


def is_content_only_output_mode(mode: str | None) -> bool:
    """True when fill_report should be skipped (user copies content themselves)."""
    return (mode or "deliverable").strip().lower() in CONTENT_ONLY_OUTPUT_MODES


def _get_solve_data(ctx: dict) -> dict | None:
    mr = ctx.get("module_results") or {}
    for key in ("solve_lab", "solve_theory"):
        entry = mr.get(key) or {}
        if entry.get("ok") and entry.get("data"):
            return dict(entry["data"])
    return None


def _code_package(solve: dict, parsed: dict) -> dict[str, Any]:
    language = solve.get("language") or parsed.get("language") or "python"
    code_files = solve.get("code_files") or parsed.get("code_files") or []
    main_file = solve.get("main_file") or parsed.get("main_file") or ""
    single_code = solve.get("code") or parsed.get("code") or ""

    files: list[dict[str, str]] = []
    if code_files:
        for item in code_files:
            if not isinstance(item, dict):
                continue
            name = (item.get("name") or item.get("filename") or "main.txt").strip()
            files.append({"name": name, "code": item.get("code") or item.get("content") or ""})
    elif single_code.strip():
        ext = {
            "python": ".py",
            "java": ".java",
            "c": ".c",
            "cpp": ".cpp",
            "javascript": ".js",
        }
        lang = (solve.get("language") or parsed.get("language") or "python").lower()
        fname = main_file or f"main{ext.get(lang, '.txt')}"
        files = [{"name": fname, "code": single_code}]

    if not main_file and files:
        main_file = files[0]["name"]

    return {
        "language": language if language else (solve.get("language") or parsed.get("language") or "python"),
        "files": files,
        "main_file": main_file,
    }


def _diagram_attachments(ctx: dict) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    solve = _get_solve_data(ctx) or {}
    parsed = solve.get("parsed") or {}
    raw_diagrams = parsed.get("diagrams") or solve.get("diagrams") or []

    uml_mr = (ctx.get("module_results") or {}).get("render_uml") or {}
    uml_images = (uml_mr.get("data") or {}).get("images_b64") or []

    for i, d in enumerate(raw_diagrams):
        if not isinstance(d, dict):
            continue
        kind = d.get("kind") or d.get("type") or "diagram"
        title = d.get("title") or d.get("name") or f"图 {i + 1}"
        entry: dict[str, Any] = {
            "kind": kind,
            "title": title,
            "plantuml": d.get("plantuml") or d.get("source") or "",
        }
        if i < len(uml_images):
            entry["image_b64"] = uml_images[i]
        out.append(entry)

    if not out and uml_images:
        for i, b64 in enumerate(uml_images):
            out.append({"kind": "uml", "title": f"图 {i + 1}", "image_b64": b64})

    return out


def _execution_from_session(session: dict, constraints: list[str]) -> dict[str, Any]:
    """Map SolveSession fields to LabDeliverable.execution."""
    if "skip_validation" in constraints:
        return {
            "validation_status": "not_requested",
            "validation_note": "已按你的设置跳过内化验证",
        }

    code_status = (session.get("code_status") or "skipped").lower()
    run_result = session.get("run_result") or {}
    stdout = (run_result.get("stdout") or run_result.get("output") or "")[:4000]
    stderr = (run_result.get("stderr") or "")[:2000]

    if code_status == "verified":
        return {
            "validation_status": "verified",
            "validation_note": "代码已通过内化验证沙箱（请在实验环境再次确认）",
            "sample_stdout": stdout,
            "sample_stderr": stderr or None,
        }
    if code_status == "degraded":
        note = run_result.get("stderr") or run_result.get("output") or "验证未通过"
        return {
            "validation_status": "failed",
            "validation_note": f"内化验证未通过：{str(note)[:200]}",
            "sample_stdout": stdout or None,
            "sample_stderr": stderr or str(note)[:2000],
        }
    reason = run_result.get("reason") or ""
    if reason == "no_runtime":
        note = "本机无对应语言运行时，已跳过内化验证"
    elif reason == "skip_validation":
        note = "已跳过内化验证"
    elif reason == "missing_jar":
        missing = run_result.get("missing_jars") or []
        labels = ", ".join(j.get("label") or j.get("id", "") for j in missing if isinstance(j, dict))
        note = f"验证需要白名单 jar（{labels}），等待你确认下载后重试" if labels else "验证需要白名单 jar，等待你确认下载"
    elif reason == "jar_download_declined":
        note = "未同意下载验证所需 jar，已跳过内化验证"
    else:
        note = "未执行内化验证（无代码或环境不可用）"
    return {
        "validation_status": "skipped",
        "validation_note": note,
        "sample_stdout": stdout or None,
    }


def _execution_block(ctx: dict) -> dict[str, Any]:
    """Populate from validation sandbox (V5-1) with legacy run_code fallback."""
    constraints = list(ctx.get("user_constraints") or [])
    session = ctx.get("solve_session")
    if not session:
        solve = _get_solve_data(ctx) or {}
        session = solve.get("solve_session")
    if session:
        return _execution_from_session(session, constraints)

    run_mr = (ctx.get("module_results") or {}).get("run_code") or {}
    run_data = run_mr.get("data") or {}
    output_mode = (ctx.get("output_mode") or "deliverable").strip().lower()

    if output_mode == "answer_only" or "skip_validation" in constraints:
        return {
            "validation_status": "not_requested",
            "validation_note": "未请求代码内化验证",
        }

    if run_mr.get("ok") and (run_data.get("stdout") or run_data.get("output")):
        return {
            "validation_status": "verified",
            "validation_note": "代码已在本地试跑（用户环境验证请自行完成）",
            "sample_stdout": (run_data.get("stdout") or run_data.get("output") or "")[:4000],
            "sample_stderr": (run_data.get("stderr") or "")[:2000],
        }

    return {
        "validation_status": "skipped",
        "validation_note": "未执行内化验证；请自行在实验环境运行代码",
    }


def build_deliverable(ctx: dict, *, deliverable_id: str | None = None) -> dict[str, Any]:
    """Assemble LabDeliverable from agent context module_results."""
    solve = _get_solve_data(ctx)
    if not solve:
        raise ValueError("无可用的解题结果（solve_lab / solve_theory）")

    parsed = dict(solve.get("parsed") or {})
    now = datetime.now(timezone.utc).isoformat()
    did = deliverable_id or f"dlv_{uuid.uuid4().hex[:12]}"

    sections = {
        "steps_analysis": parsed.get("steps_analysis") or "",
        "result_description": parsed.get("result_description") or "",
        "summary": parsed.get("summary") or "",
        "notes": parsed.get("notes") or solve.get("notes") or "",
    }

    code_pkg = _code_package(solve, parsed)
    diagrams = _diagram_attachments(ctx)
    execution = _execution_block(ctx)

    constraints: list[str] = list(ctx.get("user_constraints") or [])
    solve = _get_solve_data(ctx) or {}
    session = ctx.get("solve_session") or solve.get("solve_session") or {}
    for c in session.get("constraints_applied") or []:
        if c not in constraints:
            constraints.append(c)
    meta = solve.get("pipeline_meta") or ctx.get("pipeline_meta") or {}
    for c in meta.get("constraints_applied") or []:
        if c not in constraints:
            constraints.append(c)
    if ctx.get("sections_config") and "sections_config_applied" not in constraints:
        constraints.append("sections_config_applied")

    integrity_hash = compute_integrity_hash(
        {"sections": sections, "code": code_pkg, "diagrams": diagrams}
    )

    verification = ctx.get("verification_report") or {}
    quality_checks = [
        {
            "id": c.get("id") or f"check_{i}",
            "label": c.get("id") or "check",
            "passed": bool(c.get("ok")),
        }
        for i, c in enumerate(verification.get("checks") or [])
    ]

    return {
        "id": did,
        "created_at": now,
        "sections": sections,
        "code": code_pkg,
        "diagrams": diagrams,
        "execution": execution,
        "constraints_applied": constraints,
        "provenance": _provenance_from_ctx(ctx, constraints, integrity_hash, now),
        "quality": {
            "verify_passed": verification.get("passed"),
            "checks": quality_checks,
        },
    }


def compute_integrity_hash(dlv: dict[str, Any]) -> str:
    """SHA256 digest (16 hex chars) over sections + code + diagrams."""
    sections = dlv.get("sections") if "sections" in dlv else {}
    code_pkg = dlv.get("code") if "code" in dlv else {}
    diagrams = dlv.get("diagrams") if "diagrams" in dlv else []
    payload = json.dumps(
        {"sections": sections, "code": code_pkg, "diagrams": diagrams},
        sort_keys=True,
        ensure_ascii=False,
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()[:16]


def _provenance_from_ctx(
    ctx: dict,
    constraints: list[str],
    integrity_hash: str,
    generated_at: str,
) -> dict[str, Any]:
    prov: dict[str, Any] = {
        "ai_assisted": True,
        "generated_at": generated_at,
        "integrity_hash": integrity_hash,
    }
    custom = (ctx.get("provenance_custom_label") or "").strip()
    if not custom:
        settings = ctx.get("settings") or {}
        custom = (settings.get("provenanceCustomLabel") or settings.get("provenance_custom_label") or "").strip()
    if "provenance_label" in constraints:
        prov["custom_label"] = custom or "内容由 AI 辅助生成，本人已核对"
    elif custom:
        prov["custom_label"] = custom
    model = ctx.get("model") or (ctx.get("settings") or {}).get("model")
    if model:
        prov["model"] = model
    return prov


def should_include_provenance_footer(dlv: dict[str, Any]) -> bool:
    constraints = dlv.get("constraints_applied") or []
    prov = dlv.get("provenance") or {}
    return "provenance_label" in constraints or bool(prov.get("custom_label"))


def provenance_footer_lines(dlv: dict[str, Any]) -> list[str]:
    prov = dlv.get("provenance") or {}
    lines: list[str] = []
    if prov.get("custom_label"):
        lines.append(str(prov["custom_label"]))
    lines.append("本报告内容由 AI 辅助生成。")
    if prov.get("integrity_hash"):
        lines.append(f"校验码: {prov['integrity_hash']}")
    if prov.get("generated_at"):
        lines.append(f"生成时间: {prov['generated_at']}")
    return lines


def ensure_deliverable_provenance(dlv: dict[str, Any]) -> dict[str, Any]:
    """Fill missing provenance fields (e.g. client-side fallback deliverable)."""
    out = dict(dlv)
    prov = dict(out.get("provenance") or {})
    if not prov.get("integrity_hash"):
        prov["integrity_hash"] = compute_integrity_hash(out)
    prov.setdefault("ai_assisted", True)
    prov.setdefault("generated_at", out.get("created_at") or datetime.now(timezone.utc).isoformat())
    out["provenance"] = prov
    return out


def deliverable_to_markdown(dlv: dict[str, Any], *, include_footer: bool | None = None) -> str:
    """Export deliverable as copy-friendly Markdown."""
    dlv = ensure_deliverable_provenance(dlv)
    if include_footer is None:
        include_footer = should_include_provenance_footer(dlv)
    lines: list[str] = []
    sections = dlv.get("sections") or {}
    section_titles = [
        ("steps_analysis", "实验步骤 / 思路分析"),
        ("result_description", "实验结果说明"),
        ("summary", "实验总结"),
        ("notes", "备注"),
    ]
    for key, title in section_titles:
        body = (sections.get(key) or "").strip()
        if body:
            lines.append(f"## {title}")
            lines.append("")
            lines.append(body)
            lines.append("")

    code_pkg = dlv.get("code") or {}
    files = code_pkg.get("files") or []
    lang = code_pkg.get("language") or "text"
    if files:
        lines.append("## 代码")
        lines.append("")
        for f in files:
            name = f.get("name") or "code"
            lines.append(f"### {name}")
            lines.append("")
            lines.append(f"```{lang}")
            lines.append(f.get("code") or "")
            lines.append("```")
            lines.append("")

    diagrams = dlv.get("diagrams") or []
    if diagrams:
        lines.append("## 图表")
        lines.append("")
        for d in diagrams:
            title = d.get("title") or "图"
            lines.append(f"### {title}")
            if d.get("plantuml"):
                lines.append("")
                lines.append("```plantuml")
                lines.append(d["plantuml"])
                lines.append("```")
            lines.append("")

    execution = dlv.get("execution") or {}
    status = execution.get("validation_status") or "skipped"
    note = execution.get("validation_note") or ""
    lines.append("---")
    lines.append(f"*验证状态: {status}*")
    if note:
        lines.append(f"*{note}*")
    prov = dlv.get("provenance") or {}
    if prov.get("integrity_hash"):
        lines.append(f"*校验码: {prov['integrity_hash']}*")
    if include_footer and should_include_provenance_footer(dlv):
        lines.append("")
        for line in provenance_footer_lines(dlv):
            lines.append(f"*{line}*")

    return "\n".join(lines).strip() + "\n"


def deliverable_to_docx(dlv: dict[str, Any], *, include_footer: bool | None = None) -> bytes:
    """Export deliverable as a standalone .docx (fixed simple template)."""
    if not DOCX_OK:
        raise RuntimeError("python-docx 不可用，无法导出 Word")

    from docx import Document
    from docx.shared import Inches

    dlv = ensure_deliverable_provenance(dlv)
    if include_footer is None:
        include_footer = should_include_provenance_footer(dlv)

    doc = Document()
    doc.add_heading("实验报告", level=0)

    section_titles = [
        ("steps_analysis", "实验步骤 / 思路分析"),
        ("result_description", "实验结果说明"),
        ("summary", "实验总结"),
        ("notes", "备注"),
    ]
    sections = dlv.get("sections") or {}
    for key, title in section_titles:
        body = (sections.get(key) or "").strip()
        if body:
            doc.add_heading(title, level=1)
            doc.add_paragraph(body)

    code_pkg = dlv.get("code") or {}
    files = code_pkg.get("files") or []
    if files:
        doc.add_heading("代码", level=1)
        for f in files:
            name = f.get("name") or "code"
            doc.add_heading(name, level=2)
            doc.add_paragraph(f.get("code") or "")

    diagrams = dlv.get("diagrams") or []
    if diagrams:
        doc.add_heading("图表", level=1)
        for i, d in enumerate(diagrams):
            title = d.get("title") or f"图 {i + 1}"
            doc.add_heading(title, level=2)
            b64 = d.get("image_b64")
            if b64:
                try:
                    doc.add_picture(io.BytesIO(base64.b64decode(b64)), width=Inches(5.5))
                except Exception:
                    pass
            if d.get("plantuml"):
                doc.add_paragraph(d["plantuml"])

    execution = dlv.get("execution") or {}
    doc.add_paragraph("")
    doc.add_paragraph(f"验证状态: {execution.get('validation_status') or 'skipped'}")
    note = execution.get("validation_note") or ""
    if note:
        doc.add_paragraph(note)

    prov = dlv.get("provenance") or {}
    if prov.get("integrity_hash"):
        doc.add_paragraph(f"校验码: {prov['integrity_hash']}")

    if include_footer and should_include_provenance_footer(dlv):
        doc.add_paragraph("")
        for line in provenance_footer_lines(dlv):
            p = doc.add_paragraph(line)
            p.style = "Intense Quote"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def deliverable_code_zip_bytes(dlv: dict[str, Any]) -> bytes:
    files = (dlv.get("code") or {}).get("files") or []
    if not files:
        raise ValueError("交付物中无代码文件")
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        used: set[str] = set()
        for f in files:
            name = (f.get("name") or "main.txt").strip() or "main.txt"
            base = name
            n = 1
            while name in used:
                stem, dot, ext = base.rpartition(".")
                if dot:
                    name = f"{stem}_{n}.{ext}"
                else:
                    name = f"{base}_{n}"
                n += 1
            used.add(name)
            zf.writestr(name, f.get("code") or "")
    return buf.getvalue()


def deliverable_diagrams_zip_bytes(dlv: dict[str, Any]) -> tuple[bytes, int]:
    diagrams = dlv.get("diagrams") or []
    buf = io.BytesIO()
    count = 0
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, d in enumerate(diagrams):
            b64 = d.get("image_b64")
            if not b64:
                continue
            title = d.get("title") or f"diagram_{i + 1}"
            safe = re.sub(r"[^\w\u4e00-\u9fff\-]+", "_", title).strip("_")[:48] or f"diagram_{i + 1}"
            zf.writestr(f"{safe}.png", base64.b64decode(b64))
            count += 1
    if count == 0:
        raise ValueError("交付物中无可导出的图表 PNG")
    return buf.getvalue(), count


EXPORT_FORMATS = frozenset({"json", "markdown", "docx", "code_zip", "diagrams_zip"})


def export_deliverable(
    dlv: dict[str, Any],
    fmt: str,
    *,
    include_footer: bool | None = None,
) -> dict[str, Any]:
    """Unified export: returns payload dict for API response."""
    fmt = (fmt or "markdown").strip().lower()
    if fmt not in EXPORT_FORMATS:
        raise ValueError(f"不支持的导出格式: {fmt}")

    dlv = ensure_deliverable_provenance(dlv)
    if fmt == "json":
        return {"deliverable": dlv}
    if fmt == "markdown":
        return {
            "markdown": deliverable_to_markdown(dlv, include_footer=include_footer),
            "deliverable": dlv,
        }
    if fmt == "docx":
        data = deliverable_to_docx(dlv, include_footer=include_footer)
        return {
            "file_b64": base64.b64encode(data).decode("ascii"),
            "filename": f"lab_answer_{dlv.get('id', 'export')}.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "deliverable": dlv,
        }
    if fmt == "code_zip":
        data = deliverable_code_zip_bytes(dlv)
        return {
            "file_b64": base64.b64encode(data).decode("ascii"),
            "filename": f"lab_code_{dlv.get('id', 'export')}.zip",
            "mime_type": "application/zip",
            "deliverable": dlv,
        }
    data, _n = deliverable_diagrams_zip_bytes(dlv)
    return {
        "file_b64": base64.b64encode(data).decode("ascii"),
        "filename": f"lab_diagrams_{dlv.get('id', 'export')}.zip",
        "mime_type": "application/zip",
        "deliverable": dlv,
    }
