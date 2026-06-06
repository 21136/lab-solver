"""Fill Word report with AI answers and images."""

import base64
import re
from pathlib import Path
from typing import Optional

from config import DOCX_OK, TEMP_DIR
from log_util import loge, logi
from modules.lab_parse import complete_lab_parsed

if DOCX_OK:
    from docx import Document
    from docx.shared import Inches

_CORE_SEMANTICS = frozenset({"steps", "result", "summary"})

_IMAGE_FORBIDDEN_SEMANTICS = frozenset({"objective"})
_CONTENT_HEADING_KEYWORDS = ("内容", "任务")
_STEPS_HEADING_KEYWORDS = ("步骤", "过程", "操作")

SECTION_HEADER_PATTERNS = {
    "steps": re.compile(r"^三[、．.]\s*"),
    "result": re.compile(r"^四[、．.]\s*"),
    "summary": re.compile(r"^五[、．.]\s*"),
}


def _is_lab_section_header(text, key):
    if not text or len(text) > 30:
        return False
    pat = SECTION_HEADER_PATTERNS.get(key)
    if pat and pat.match(text):
        return True
    fallback = {"steps": "实验步骤", "result": "实验结果", "summary": "实验总结"}
    title = fallback.get(key, "")
    return text == title or text.startswith(title + " ")


# ── DA2: semantic section detection (keyword-based, not number-based) ──

_SECTION_NUM_RE = re.compile(r"^([一二三四五六七八九十\d]+)[、．.]")

_SEMANTIC_KEYWORDS = {
    "steps": ["步骤", "内容", "操作", "过程", "任务"],
    "result": ["结果", "分析", "数据", "输出"],
    "summary": ["总结", "思考", "心得", "讨论", "体会", "小结"],
    "objective": ["目的", "目标", "要求"],
    "principles": ["原理", "背景", "知识", "基础"],
    "discussion": ["讨论", "思考", "问答", "问题"],
    "appendix": ["附录", "附件", "参考", "源码"],
}

_METADATA_LABELS = {
    "课程名称", "实验序号", "实验名称", "实训项目", "实训名称",
    "专业", "学号", "姓名", "班级", "日期", "指导教师",
    "实验时间", "实验地点", "同组人员",
}

_FALLBACK_HEADERS = {
    "steps": "实验步骤",
    "result": "实验结果",
    "summary": "实验总结",
}


def _guess_semantic(heading_text):
    for role, keywords in _SEMANTIC_KEYWORDS.items():
        for kw in keywords:
            if kw in heading_text:
                return role
    return None


def _build_section_map(sections_detected):
    section_map = {"steps": None, "result": None, "summary": None}
    for sec in sections_detected:
        role = sec.get("semantic")
        if role and role in section_map and section_map[role] is None:
            section_map[role] = {
                "type": "paragraph",
                "heading": sec["heading"],
                "para_index": sec["index"],
            }
    return section_map, sections_detected


def _apply_semantic_overrides(sections_detected, overrides):
    """Apply DA4 UI role→heading overrides onto detected sections."""
    if not sections_detected:
        return []
    overrides = overrides or {}
    out = []
    for sec in sections_detected:
        item = dict(sec)
        heading = item.get("heading", "")
        for role, mapped_heading in overrides.items():
            if mapped_heading == heading and role in _CORE_SEMANTICS:
                item["semantic"] = role
                break
        out.append(item)
    return out


def _resolve_fill_sections(meta, paras, sections_detected=None):
    """Prefer parse-time section data + user overrides; fall back to re-detection."""
    meta = meta or {}
    overrides = meta.get("semantic_overrides") or {}
    _detected, _ = detect_sections(paras)
    working = sections_detected if sections_detected is not None else _detected
    working = _apply_semantic_overrides(working, overrides)

    cached_map = meta.get("section_map") or {}
    for key in ("steps", "result", "summary"):
        heading = (cached_map.get(key) or {}).get("heading")
        if not heading:
            continue
        for sec in working:
            if sec.get("heading") == heading:
                sec["semantic"] = key
                break

    section_map, _ = _build_section_map(working)
    return working, section_map


def _build_fill_hints(section_map, sections_detected=None):
    hints = {}
    # Determine whether all three core sections are present
    core_present = all(section_map.get(k) is not None for k in ("steps", "result", "summary"))
    non_core_count = 0
    if sections_detected:
        non_core_count = sum(
            1 for s in sections_detected
            if s.get("semantic") not in _CORE_SEMANTICS
        )

    # Screenshots target (depends only on section_map, always computed)
    if section_map.get("result") is None:
        if section_map.get("summary"):
            hints["screenshots_target"] = "summary"
        elif section_map.get("steps"):
            hints["screenshots_target"] = "steps"

    # Merge hints: skip only when all core sections present AND >= 2 non-core sections
    # (custom layout with complete core structure — each section is independent)
    if core_present and non_core_count >= 2:
        return hints

    if section_map.get("steps") is None and section_map.get("result"):
        hints["merge_steps_into"] = "result"
    if section_map.get("result") is None and section_map.get("summary"):
        hints["merge_result_into"] = "summary"
    if section_map.get("summary") is None and section_map.get("result"):
        hints["merge_summary_into"] = "result"

    steps_entry = section_map.get("steps")
    if steps_entry:
        heading = steps_entry.get("heading") or ""
        if any(kw in heading for kw in _CONTENT_HEADING_KEYWORDS):
            hints.setdefault("uml_default_target", "content")
        elif any(kw in heading for kw in _STEPS_HEADING_KEYWORDS):
            hints.setdefault("uml_default_target", "steps")
        else:
            hints.setdefault("uml_default_target", "content")
    return hints


def _merge_fill_hints(built_hints: dict, meta_hints: dict | None) -> dict:
    """Merge parse-time hints with Agent-provided fill_hints (Agent wins on conflict)."""
    merged = dict(built_hints or {})
    for key, value in (meta_hints or {}).items():
        if value is not None:
            merged[key] = value
    return merged


def _is_content_like_heading(heading: str) -> bool:
    return any(kw in (heading or "") for kw in _CONTENT_HEADING_KEYWORDS)


def _is_steps_like_heading(heading: str) -> bool:
    return any(kw in (heading or "") for kw in _STEPS_HEADING_KEYWORDS)


def _default_diagram_target(section_map: dict, filled_keys: set) -> str | None:
    """Default UML placement: content-like steps > steps-like steps > result > summary."""
    steps_entry = section_map.get("steps")
    if steps_entry and "steps" in filled_keys:
        heading = steps_entry.get("heading") or ""
        if _is_content_like_heading(heading) or not _is_steps_like_heading(heading):
            return "steps"
        return "steps"
    for key in ("result", "summary"):
        if key in filled_keys and section_map.get(key):
            return key
    return None


def _sanitize_diagram_target(
    target_semantic: str,
    section_map: dict,
    filled_keys: set,
) -> str | None:
    """Map placement_hint to a writable section; never objective."""
    hint = (target_semantic or "").strip().lower()
    if not hint or hint in _IMAGE_FORBIDDEN_SEMANTICS:
        return _default_diagram_target(section_map, filled_keys)

    if hint == "content":
        return _default_diagram_target(section_map, filled_keys)
    if hint in ("environment", "design"):
        if section_map.get("steps") and "steps" in filled_keys:
            return "steps"
        return _default_diagram_target(section_map, filled_keys)
    if hint in _CORE_SEMANTICS and hint in filled_keys and section_map.get(hint):
        return hint
    return _default_diagram_target(section_map, filled_keys)


def _resolve_diagram_placements(
    uml_images_b64: list,
    fill_hints: dict,
    section_map: dict,
    filled_keys: set,
) -> list[tuple[str, list]]:
    """Group UML images by target section semantic (after forbidden/downgrade rules)."""
    if not uml_images_b64:
        return []

    diagrams_target = fill_hints.get("diagrams_target")
    if isinstance(diagrams_target, list) and diagrams_target:
        by_target: dict[str, list] = {}
        default_hint = fill_hints.get("uml_default_target") or "content"
        for i, b64 in enumerate(uml_images_b64):
            if not b64:
                continue
            spec = next(
                (t for t in diagrams_target if t.get("image_index") == i),
                None,
            )
            hint = (spec or {}).get("target_semantic") or default_hint
            target = _sanitize_diagram_target(hint, section_map, filled_keys)
            if not target:
                continue
            by_target.setdefault(target, []).append(b64)
        return list(by_target.items())

    default_hint = fill_hints.get("uml_default_target") or "content"
    target = _sanitize_diagram_target(default_hint, section_map, filled_keys)
    if not target:
        return []
    valid = [b for b in uml_images_b64 if b]
    return [(target, valid)] if valid else []


_STEPS_TABLE_LABELS = (
    "实训步骤及内容", "实训步骤", "实训任务", "实训内容",
    "实验步骤及内容", "实验内容",
)
_OBJECTIVE_TABLE_LABELS = ("实验目的", "实验目标", "实训目的")
_NAME_TABLE_LABELS = ()  # experiment_name uses exact label checks in _training_fill_targets

_OBJECTIVE_START_RE = re.compile(
    r"(?:"
    r"(?:[（(]?\s*[一二三四五六七八九十\d]+\s*[)）]\s*)?"
    r"(?:实验目的(?:与原理)?|实训目的)"
    r"|一[、．.]\s*实验目的"
    r"|实验目的[：:]"
    r")",
    re.IGNORECASE,
)
_OBJECTIVE_END_RE = re.compile(
    r"(?:"
    r"(?:[（(]?\s*[一二三四五六七八九十\d]+\s*[)）]\s*)?"
    r"(?:实验(?:内容与步骤|内容|步骤|任务)|实训(?:内容与步骤|内容|步骤))"
    r"|二[、．.]\s*实验(?:内容|步骤)"
    r"|实验(?:内容与步骤|内容|步骤)[：:]"
    r")",
    re.IGNORECASE,
)


def _strip_planner_assignment_prefix(text: str) -> str:
    """Keep only the assignment portion from planner_input_text."""
    t = (text or "").strip()
    if "【实验要求】" not in t:
        return t
    rest = t.split("【实验要求】", 1)[1]
    for marker in ("【待填报告", "【参考资料", "【报告全文】"):
        if marker in rest:
            rest = rest.split(marker, 1)[0]
    return rest.strip()


def _clean_objective_body(body: str, *, max_len: int = 1500) -> str:
    body = re.sub(r"^[：:\s]+", "", (body or "").strip())
    if len(body) > max_len:
        body = body[:max_len].rstrip() + "…"
    return body


def extract_objective_from_assignment(text: str) -> str:
    """Extract the experiment objective section from assignment / planner input."""
    text = _strip_planner_assignment_prefix(text)
    if not text:
        return ""

    start = _OBJECTIVE_START_RE.search(text)
    if not start:
        return ""

    start_pos = start.end()
    end_match = _OBJECTIVE_END_RE.search(text, start_pos)
    end_pos = end_match.start() if end_match else len(text)
    return _clean_objective_body(text[start_pos:end_pos])


def _resolve_objective_text(parsed, metadata=None, ans=None) -> str:
    """Prefer assignment 实验目的; fall back to steps_analysis first paragraph."""
    parsed = parsed or {}
    meta = dict(metadata or {})
    ans_meta = (ans or {}).get("metadata") or {}

    for key in ("objective", "experiment_objective"):
        explicit = (parsed.get(key) or "").strip()
        if explicit:
            return explicit

    for src in (
        meta.get("assignment_text"),
        ans_meta.get("assignment_text"),
        meta.get("planner_input_text"),
        ans_meta.get("planner_input_text"),
    ):
        extracted = extract_objective_from_assignment(src or "")
        if extracted:
            logi("fill", f"objective from assignment excerpt len={len(extracted)}")
            return extracted

    steps = (parsed.get("steps_analysis") or "").strip()
    fallback = steps.split("\n\n")[0].strip() if steps else ""
    return _clean_objective_body(fallback, max_len=800)


def _training_fill_targets(table_map_entries):
    """From table_map entries, determine which cells are fill targets for answers."""
    fill_targets = []
    for entry in table_map_entries or []:
        label = entry.get("label", "")
        table_idx = entry.get("table", 0)
        if any(kw in label for kw in _STEPS_TABLE_LABELS):
            fill_targets.append({"semantic": "steps", **entry})
        elif any(kw in label for kw in _OBJECTIVE_TABLE_LABELS):
            fill_targets.append({"semantic": "objective", **entry})
        elif label == "实验名" or label.strip() == "实验名":
            fill_targets.append({"semantic": "experiment_name", **entry})
        elif label == "实验名称" and table_idx > 0:
            fill_targets.append({"semantic": "experiment_name", **entry})
    return fill_targets


def _training_combined_body(fill_content, *, fill_sections=None):
    """Merge steps/result/summary for a single table body cell."""
    allowed = set(fill_sections) if fill_sections else None
    parts = []
    for key in ("steps", "result", "summary"):
        if allowed is not None and key not in allowed:
            continue
        val = (fill_content.get(key) or "").strip()
        if val:
            parts.append(val)
    return "\n\n".join(parts)


def detect_sections(paragraphs):
    """
    Scan paragraphs for section headers and produce semantic mapping.

    Uses keyword matching on heading text (步骤/结果/总结 etc.) instead of
    fixed numbering (三→steps, 四→result, 五→summary). Also matches
    keyword-only headers without a number prefix (e.g. "实验步骤").

    Returns:
        sections_detected: list of {index, heading, semantic}
        section_map: {steps/result/summary: {type, heading, para_index} or None}
    """
    sections_detected = []
    for idx, para in enumerate(paragraphs):
        text = para.text.strip() if hasattr(para, "text") else str(para).strip()
        if not text or len(text) > 30:
            continue

        semantic = None
        m = _SECTION_NUM_RE.match(text)
        if m:
            num = m.group(1)
            body = text[m.end():].strip()
            if body in _METADATA_LABELS:
                continue
            semantic = _guess_semantic(text)
            # Skip numbered list items (1. / 2.) that are not real section headers
            if num.isdigit() and not semantic:
                continue
        else:
            if text in _METADATA_LABELS:
                continue
            for key, fallback in _FALLBACK_HEADERS.items():
                if text == fallback or text.startswith(fallback + " "):
                    semantic = key
                    break

        if m or semantic:
            sections_detected.append({
                "index": idx,
                "heading": text,
                "semantic": semantic,
            })

    section_map, _ = _build_section_map(sections_detected)
    return sections_detected, section_map


def do_fill(input_path, answers, output_path="", *, fill_sections=None, metadata=None, settings=None):
    if not DOCX_OK:
        return create_txt_report(answers, output_path)

    doc = Document(str(input_path)) if input_path and input_path.exists() else Document()

    meta = dict(metadata or {})
    sections_detected = meta.get("sections_detected") or None

    for ans in answers:
        if not ans:
            continue
        if ans.get("type") == "lab_report":
            fill_lab(doc, ans, fill_sections=fill_sections or ans.get("fill_sections"),
                     metadata=metadata, sections_detected=sections_detected, settings=settings)
        else:
            fill_generic_ans(doc, ans)

    if not output_path:
        stem = input_path.stem if input_path else "实验报告"
        output_path = str(TEMP_DIR / f"{stem}_已完成.docx")

    doc.save(output_path)
    logi("fill", f"已保存: {output_path}")
    return output_path


def _fill_other_section(sec, doc, paras, all_section_indices, ans, *, settings=None):
    """Generate and write content for a non-core section (objective/principles/discussion/appendix/other).

    Calls LLM once per section with context from the original document and AI-generated
    core content, then writes the result via _replace_section.
    """
    if not settings:
        logi("fill", f"skip other section '{sec.get('heading')}': no settings/api_key")
        return

    api_key = (settings.get("api_key") or "").strip()
    if not api_key:
        logi("fill", f"skip other section '{sec.get('heading')}': no api_key")
        return

    heading = sec.get("heading", "")
    idx = sec.get("index")
    semantic = sec.get("semantic") or "other"

    # Extract original text between this section and the next
    next_idx = None
    for si in sorted(all_section_indices):
        if si > idx:
            next_idx = si
            break
    if next_idx is None:
        next_idx = len(paras)

    original_lines = []
    for j in range(idx + 1, min(next_idx, len(paras))):
        t = paras[j].text.strip()
        if t:
            original_lines.append(t)
    original_text = "\n".join(original_lines)

    # Gather AI-generated core content as context
    parsed = ans.get("parsed") or {}
    steps_text = parsed.get("steps_analysis") or ""
    result_text = parsed.get("result_description") or ""
    summary_text = parsed.get("summary") or ""
    ai_content = f"{steps_text}\n\n{result_text}\n\n{summary_text}".strip()

    planner_input = (ans.get("metadata") or {}).get("planner_input_text") or ""

    prompt = (
        f"你是一名大学课程助教。请为以下实验报告的「{heading}」节撰写内容。\n"
        f"\n"
        f"【作业要求】\n"
        f"{planner_input[:2000] or '（未提供）'}\n"
        f"\n"
        f"【该节原文（如有）】\n"
        f"{original_text[:2000] or '（空）'}\n"
        f"\n"
        f"【AI 已生成的解题内容（供参考）】\n"
        f"{ai_content[:3000] or '（未提供）'}\n"
        f"\n"
        f"请输出该节的完整内容（直接可填入 Word 的中文段落）。"
    )

    from llm_client import chat

    try:
        result = chat(
            api_key=api_key,
            provider=settings.get("provider", "deepseek"),
            model=settings.get("model", "deepseek-chat"),
            prompt=prompt,
            custom_url=settings.get("custom_url") or settings.get("customUrl") or "",
            max_tokens=2000,
            phase="fill_other_section",
        )
        content = (result.get("content") or "").strip()
        if content:
            _replace_section(paras, idx, content, all_section_indices=all_section_indices)
            logi("fill", f"filled other section '{heading}' semantic={semantic} len={len(content)}")
        else:
            logi("fill", f"other section '{heading}': LLM returned empty content")
    except Exception as e:
        loge("fill", f"other section '{heading}' LLM error: {e}")


def fill_lab(doc, ans, *, fill_sections=None, metadata=None, sections_detected=None, settings=None):
    parsed = complete_lab_parsed(ans.get("parsed", {}), ans.get("answer", ""))
    logi("fill", f"parsed keys={list(parsed.keys())}")

    include_code = ans.get("include_code", True)

    steps_text = parsed.get("steps_analysis", "")
    if include_code and parsed.get("code"):
        steps_text = steps_text + "\n\n" + parsed["code"]
    steps_text = steps_text.strip()

    objective_text = _resolve_objective_text(parsed, metadata=metadata, ans=ans)

    meta_early = dict(metadata or {})
    experiment_name = (
        meta_early.get("experiment_title")
        or meta_early.get("experiment_name")
        or (ans.get("title") or "").strip()
        or ""
    )

    fill_content = {
        "steps": steps_text,
        "result": (
            parsed.get("result_description", "")
            + "\n\n"
            + parsed.get("expected_output", "")
        ).strip(),
        "summary": parsed.get("summary", ""),
        "objective": objective_text,
        "experiment_name": experiment_name,
    }
    allowed = set(fill_sections) if fill_sections else None
    images_b64 = ans.get("images_b64") or []
    if not images_b64 and ans.get("image_b64"):
        images_b64 = [ans["image_b64"]]
    uml_images_b64 = ans.get("uml_images_b64") or []

    if not any(fill_content.values()):
        fill_content["steps"] = ans.get("answer", "")
        logi("fill", "parsed为空，用answer全文填步骤")

    logi(
        "fill",
        f'fill_content lengths: steps={len(fill_content["steps"])} '
        f'result={len(fill_content["result"])} summary={len(fill_content["summary"])}',
    )

    meta = dict(metadata or {})
    layout = meta.get("report_layout") or ""
    table_only_no_targets = False

    # Auto-detect training_table when metadata was not passed (e.g. /api/fill-report)
    if layout != "training_table":
        from modules.parse_report import _detect_table_layout

        tl = _detect_table_layout(doc)
        if tl.get("report_layout") == "training_table":
            layout = "training_table"
            meta["report_layout"] = layout
            if not meta.get("table_map"):
                meta["table_map"] = tl.get("table_map") or []

    # ── training_table layout: write into table cells ──
    if layout == "training_table":
        table_map = meta.get("table_map") or []
        if _training_fill_targets(table_map):
            table_fill_hints = _merge_fill_hints(
                {},
                meta.get("fill_hints"),
            )
            filled = _fill_training_table(
                doc, fill_content, images_b64, uml_images_b64,
                table_map=table_map,
                fill_sections=fill_sections,
                fill_hints=table_fill_hints,
            )
            has_core_content = any(fill_content.get(k, "").strip() for k in _CORE_SEMANTICS)
            if has_core_content and not filled:
                raise ValueError(
                    "表格模版未能写入内容（未找到可写入的「实验内容/实训步骤及内容」等单元格）。"
                    f" table_map={table_map}"
                )
            return
        table_only_no_targets = True
        logi("fill", "training_table markers without fill targets → paragraph fill")

    # ── paragraph-based fill ──
    paras = list(doc.paragraphs)
    meta = dict(metadata or {})
    working_sections, section_map = _resolve_fill_sections(
        meta, paras, sections_detected=sections_detected
    )
    fill_hints = _merge_fill_hints(
        _build_fill_hints(section_map, sections_detected=working_sections),
        meta.get("fill_hints"),
    )
    logi("fill", f"section_map={ {k: v['heading'] if v else None for k, v in section_map.items()} }")
    logi("fill", f"fill_hints={fill_hints} sections_detected={len(working_sections)}")

    # Merge content when sections are missing (only when not in custom-layout mode)
    for target, source in fill_hints.items():
        if target.startswith("merge_") and source in fill_content:
            source_key = target[len("merge_"):].replace("_into", "")
            dest_key = source
            if fill_content.get(source_key, "").strip():
                sep = "\n\n" if fill_content.get(dest_key, "").strip() else ""
                fill_content[dest_key] = fill_content.get(dest_key, "") + sep + fill_content[source_key]
                logi("fill", f"merged {source_key} → {dest_key} (section missing)")

    # Collect all section indices for boundary detection in _replace_section
    all_indices = {s["index"] for s in working_sections}

    filled_keys = set()

    for sec in working_sections:
        semantic = sec.get("semantic")
        idx = sec.get("index")

        if semantic in _CORE_SEMANTICS:
            # ── core section: use existing fill_content path ──
            if allowed is not None and semantic not in allowed:
                continue
            if semantic in filled_keys:
                continue
            mapped = section_map.get(semantic)
            if not mapped or mapped["para_index"] != idx:
                continue
            content = fill_content.get(semantic, "")
            logi("fill", f"匹配 [{idx}] key={semantic} content_len={len(content)}")
            if content:
                _replace_section(paras, idx, content, section_map=section_map, all_section_indices=all_indices)
                filled_keys.add(semantic)
            if semantic == "result" and images_b64:
                _insert_images_after(doc, paras[idx], images_b64)
        elif semantic == "objective":
            if allowed is not None and "objective" not in allowed:
                continue
            content = fill_content.get("objective", "")
            if content:
                _replace_section(paras, idx, content, all_section_indices=all_indices)
                logi("fill", f"匹配 [{idx}] key=objective content_len={len(content)} (from assignment)")
        else:
            # ── non-core section: generate content via LLM ──
            section_key = semantic or f"sec_{idx}"
            if allowed is not None and section_key not in allowed and semantic not in allowed:
                continue
            _fill_other_section(sec, doc, paras, all_indices, ans, settings=settings)

    # Screenshots when result section is missing → use fill_hints target
    screenshots_placed = "result" in filled_keys
    if not screenshots_placed and images_b64:
        target_key = fill_hints.get("screenshots_target", "summary")
        mapped = section_map.get(target_key)
        if mapped:
            target_para = paras[mapped["para_index"]] if mapped["para_index"] < len(paras) else None
            if target_para:
                _insert_images_after(doc, target_para, images_b64)
                logi("fill", f"screenshots placed in {target_key} (result section missing)")

    if uml_images_b64:
        placements = _resolve_diagram_placements(
            uml_images_b64, fill_hints, section_map, filled_keys,
        )
        for target_semantic, imgs in placements:
            if target_semantic in _IMAGE_FORBIDDEN_SEMANTICS:
                logi("fill", f"跳过禁止插图节: {target_semantic}")
                continue
            _insert_images_at_section_end(
                doc, target_semantic, imgs, section_map=section_map,
            )
            logi("fill", f"已插入 UML 图 {len(imgs)} 张到 {target_semantic}")

    core_filled = filled_keys & _CORE_SEMANTICS
    has_core_content = any(fill_content.get(k, "").strip() for k in _CORE_SEMANTICS)
    if has_core_content and not core_filled:
        headings = [s.get("heading") for s in working_sections]
        if table_only_no_targets and not working_sections:
            logi("fill", "table template has no fill targets and no paragraph sections; skipped")
            return
        raise ValueError(
            "核心节（步骤/结果/总结）未能匹配到文档标题，填表未写入内容。"
            f" 检测到: {headings}"
        )


def _write_table_cell(table, row_idx, col_idx, content):
    """Write text into a table value cell (label in col N → value in col N+1)."""
    if row_idx >= len(table.rows):
        return False
    row = table.rows[row_idx]
    fill_ci = col_idx + 1 if col_idx + 1 < len(row.cells) else col_idx
    if fill_ci >= len(row.cells):
        return False
    cell = row.cells[fill_ci]
    for p in cell.paragraphs:
        p.text = ""
    if cell.paragraphs:
        cell.paragraphs[0].text = content
    else:
        cell.add_paragraph(content)
    return True


def _table_cell_anchor(doc, target):
    """Return the value cell for a table_map target (for in-cell image insertion)."""
    ti = target.get("table", 0)
    ri = target.get("row", 0)
    ci = target.get("col", 0)
    if ti >= len(doc.tables):
        return None
    table = doc.tables[ti]
    if ri >= len(table.rows):
        return None
    row = table.rows[ri]
    fill_ci = ci + 1 if ci + 1 < len(row.cells) else ci
    if fill_ci >= len(row.cells):
        return None
    return row.cells[fill_ci]


def _insert_images_in_cell(cell, images_b64):
    """Insert images as paragraphs inside a table cell (not doc body)."""
    if not cell or not images_b64:
        return 0
    count = 0
    for i, b64 in enumerate(images_b64):
        if not b64:
            continue
        try:
            img_bytes = base64.b64decode(b64)
            img_path = TEMP_DIR / f"cell_img_{i}.png"
            img_path.write_bytes(img_bytes)
            p = cell.add_paragraph()
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(5.5))
            count += 1
        except Exception as e:
            loge("fill", f"单元格插入图片失败: {e}")
    if count:
        logi("fill", f"training_table: inserted {count} image(s) in cell")
    return count


def _fill_training_table(doc, fill_content, images_b64, uml_images_b64, *,
                         table_map=None, fill_sections=None, fill_hints=None):
    """Write answers into table-report cells (实训/实验表格模版).

    Supports ``实验内容`` / ``实验目的`` / ``实验名`` and legacy 实训 markers.
    Screenshots and UML attach after the ``steps`` (body) cell when present.
    """
    targets = _training_fill_targets(table_map)
    if not targets:
        logi("fill", "training_table: no fill-target cells found in table_map")
        return 0

    body_text = _training_combined_body(fill_content, fill_sections=fill_sections)
    semantic_text = {
        "steps": body_text,
        "objective": (fill_content.get("objective") or "").strip(),
        "experiment_name": (fill_content.get("experiment_name") or "").strip(),
    }

    if not any(semantic_text.values()):
        logi("fill", "training_table: no content to fill")
        return 0

    filled_count = 0
    for target in targets:
        semantic = target.get("semantic", "steps")
        content = semantic_text.get(semantic, "")
        if not content:
            continue
        ti = target.get("table")
        ri = target.get("row")
        ci = target.get("col")
        if ti is None or ri is None or ci is None or ti >= len(doc.tables):
            continue
        if _write_table_cell(doc.tables[ti], ri, ci, content):
            filled_count += 1
            logi(
                "fill",
                f"training_table fill: table[{ti}].row({ri}) "
                f"semantic={semantic} label={target.get('label')}",
            )

    if filled_count:
        logi("fill", f"training_table: filled {filled_count} cell(s)")

    steps_target = next((t for t in targets if t.get("semantic") == "steps"), None)
    non_objective = [t for t in targets if t.get("semantic") not in _IMAGE_FORBIDDEN_SEMANTICS]
    attach_target = steps_target or (non_objective[-1] if non_objective else None)

    uml_placements = _resolve_diagram_placements(
        uml_images_b64 or [],
        fill_hints or {},
        {"steps": {"heading": (steps_target or {}).get("label", "实验内容")}},
        {"steps"} if steps_target else set(),
    )
    uml_by_semantic = dict(uml_placements)

    if attach_target:
        cell = _table_cell_anchor(doc, attach_target)
        if cell is not None:
            if images_b64:
                _insert_images_in_cell(cell, images_b64)
            uml_for_steps = uml_by_semantic.get("steps") or (
                uml_images_b64 if not uml_placements else []
            )
            if uml_for_steps:
                _insert_images_in_cell(cell, uml_for_steps)
                logi("fill", f"training_table: inserted {len(uml_for_steps)} UML in steps cell")

    return filled_count


def _replace_section(paras, idx, content, *, section_map=None, all_section_indices=None):
    from docx.oxml import OxmlElement

    section_para = paras[idx]

    # Use all_section_indices first (preferred, includes non-core sections)
    end = idx + 1
    if all_section_indices:
        for si in sorted(all_section_indices):
            if si > idx:
                end = si
                break
    elif section_map:
        next_boundary = None
        for key in ("steps", "result", "summary"):
            mapped = section_map.get(key)
            if mapped and mapped["para_index"] > idx:
                if next_boundary is None or mapped["para_index"] < next_boundary:
                    next_boundary = mapped["para_index"]
        if next_boundary is not None:
            end = next_boundary
    if end == idx + 1:
        # Fallback: number-based boundary detection
        while end < len(paras):
            t = paras[end].text.strip()
            if t and re.match(r"^[一二三四五六七八九十\d]+[、．.]", t) and len(t) < 25:
                break
            end += 1

    for j in range(end - 1, idx, -1):
        p = paras[j]._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    prev = section_para._element
    for line in content.split("\n"):
        np = OxmlElement("w:p")
        nr = OxmlElement("w:r")
        nt = OxmlElement("w:t")
        nt.text = line
        nt.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        nr.append(nt)
        np.append(nr)
        prev.addnext(np)
        prev = np


def _insert_image_after(doc, ref_para, image_b64, index=0):
    img_bytes = base64.b64decode(image_b64)
    img_path = TEMP_DIR / f"insert_img_{index}.png"
    img_path.write_bytes(img_bytes)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))
    # addnext moves p._element from doc body to be next sibling of ref_para
    ref_para._element.addnext(p._element)
    return p


def _insert_images_at_section_end(doc, section_key, images_b64, *, section_map=None):
    paras = list(doc.paragraphs)
    start = None
    if section_map and section_map.get(section_key):
        start = section_map[section_key]["para_index"]
    else:
        for idx, para in enumerate(paras):
            if _is_lab_section_header(para.text.strip(), section_key):
                start = idx
                break
    if start is None:
        return
    end = start + 1
    while end < len(paras):
        t = paras[end].text.strip()
        if t and re.match(r"^[一二三四五六七八九十]+[、．.]", t) and len(t) < 25:
            break
        end += 1
    anchor = paras[end - 1] if end > start + 1 else paras[start]
    _insert_images_after(doc, anchor, images_b64)


def _insert_images_after(doc, ref_para, images_b64):
    try:
        anchor = ref_para
        for i, b64 in enumerate(images_b64):
            if not b64:
                continue
            new_p = _insert_image_after(doc, anchor, b64, index=i)
            anchor = new_p
        logi("fill", f"已插入实验结果截图 {len(images_b64)} 张")
    except Exception as e:
        loge("fill", f"插入图片失败: {e}")


def fill_generic_ans(doc, ans):
    ph = ans.get("placeholder", "")
    text = ans.get("answer", "")
    filled = False
    for para in doc.paragraphs:
        if ph and ph in para.text:
            para.text = para.text.replace(ph, text)
            filled = True
    if not filled:
        doc.add_paragraph(f"\n题目: {ans.get('title', '')}\n{text}")


def create_txt_report(answers, output_path=""):
    lines = ["实验报告\n" + "=" * 40]
    for ans in answers:
        if ans:
            lines.append(f"\n{ans.get('title', '')}\n{ans.get('answer', '')}")
    if not output_path:
        output_path = str(TEMP_DIR / "实验报告_已完成.txt")
    Path(output_path).write_text("\n".join(lines), encoding="utf-8")
    return output_path
