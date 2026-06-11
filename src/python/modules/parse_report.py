"""Word / PDF report parsing."""

import re
from pathlib import Path

from config import DOCX_OK, PDF_OK
from modules.code_cloze import detect_code_cloze

if DOCX_OK:
    from docx import Document

MIN_BODY_CHARS = 200
WARN_LEGACY_DOC = {
    "code": "legacy_doc",
    "message": "旧版 .doc 无法转换。请安装免费的 LibreOffice（https://www.libreoffice.org）或在 Word 中另存为 .docx 后重新上传",
}


def is_legacy_doc(file_name: str) -> bool:
    lower = (file_name or "").lower()
    return lower.endswith(".doc") and not lower.endswith(".docx")


def is_pdf(file_name: str) -> bool:
    return (file_name or "").lower().endswith(".pdf")


def document_format(file_name: str) -> str:
    if is_pdf(file_name):
        return "pdf"
    if is_legacy_doc(file_name):
        return "doc"
    return "docx"


def _pdf_hints_to_warnings(hints: list) -> list:
    warnings = []
    for h in hints or []:
        if isinstance(h, dict) and h.get("message"):
            warnings.append({"code": h.get("code", "pdf_hint"), "message": h["message"]})
    return warnings


def collect_parse_warnings(full_text: str, metadata: dict, file_name: str) -> list:
    """Surface parse quality issues without changing question shape."""
    warnings = []
    if is_legacy_doc(file_name):
        warnings.append(dict(WARN_LEGACY_DOC))

    if is_pdf(file_name):
        if not PDF_OK:
            warnings.append(
                {
                    "code": "pdf_unavailable",
                    "message": "无法读取 PDF（请更新解题能手或联系管理员检查安装包）",
                }
            )
            return warnings
    elif not DOCX_OK:
        warnings.append(
            {
                "code": "docx_unavailable",
                "message": "未安装 python-docx，无法读取 Word 文档",
            }
        )
        return warnings

    body = (full_text or "").strip()
    body_len = len(body)
    has_cover = bool(
        metadata.get("course")
        or metadata.get("experiment_title")
        or metadata.get("student_name")
    )
    image_assets = (metadata.get("image_assets") or [])
    has_images = len(image_assets) > 0
    assignment_images = [img for img in image_assets if img.get("role_guess") == "assignment"]

    if has_cover and body_len < MIN_BODY_CHARS:
        if has_images:
            warnings.append(
                {
                    "code": "short_body_with_images",
                    "message": f"正文较短但检测到 {len(image_assets)} 张嵌入图片"
                    + (f"（{len(assignment_images)} 张疑似题目图）" if assignment_images else "")
                    + "，建议启用 OCR 识别图片中的题目文字",
                }
            )
        else:
            warnings.append(
                {
                    "code": "short_body_with_cover",
                    "message": "已从封面表格读取信息，但正文过短，可能漏读表格或图片中的题目",
                }
            )
    elif body_len < 80:
        if has_images:
            warnings.append(
                {
                    "code": "short_text_with_images",
                    "message": f"正文极短但检测到 {len(image_assets)} 张嵌入图片，题目可能在图片中",
                }
            )
        else:
            warnings.append(
                {
                    "code": "short_text",
                    "message": "提取的正文过短，可能漏读表格/图片题",
                }
            )

    if body and "图" in body and body_len < 500:
        warnings.append(
            {
                "code": "possible_missing_figures",
                "message": "文档含图题但正文较短，图片内文字可能未被读取",
            }
        )

    # IM1: warn when many images detected — suggest OCR
    if len(assignment_images) >= 3:
        warnings.append({
            "code": "multiple_assignment_images",
            "message": f"检测到 {len(assignment_images)} 张疑似题目图片，"
                       "题目要求可能在图片中，建议后续使用 OCR/Vision 识别",
        })

    return warnings


_OCR_ACTION_WARN_CODES = frozenset({
    "short_body_with_images",
    "short_text_with_images",
    "ocr_suggested",
    "possible_missing_figures",
    "multiple_assignment_images",
})


def augment_ocr_action_warnings(
    warnings: list,
    *,
    enable_image_ocr: bool,
    metadata: dict,
) -> None:
    """Add UI action hints for parse warnings that OCR can resolve (IM2-b)."""
    if enable_image_ocr:
        return
    image_assets = metadata.get("image_assets") or []
    if not image_assets:
        return
    ocr_succeeded = bool(metadata.get("assignment_from_images"))
    for w in warnings:
        if not isinstance(w, dict):
            continue
        code = w.get("code")
        if code in _OCR_ACTION_WARN_CODES:
            w["action"] = "enable_ocr_reparse"
        elif code == "pdf_scanned" and not ocr_succeeded:
            w["action"] = "enable_ocr_reparse"


_TRAINING_TABLE_MARKERS = (
    "实训步骤及内容", "实验步骤及内容",
    "实训步骤", "实训任务", "实训内容",
    "实验内容", "实验目的", "实验名",
)

# Longest first so "实验名称" matches before "实验名"
_TRAINING_TABLE_MARKERS_SORTED = tuple(
    sorted(_TRAINING_TABLE_MARKERS, key=len, reverse=True)
)

_TABLE_LAYOUT_KEY_LABELS = (
    "课程名称", "实验序号", "实验名称", "实验名", "实训项目", "实训名称",
    "实训步骤及内容", "实训任务", "实训内容", "实训步骤",
    "实验内容", "实验目的", "实验环境", "指导老师", "指导教师",
    "专业", "学号", "姓名", "班级", "日期", "实验日期",
)


def _cell_texts(row) -> list[str]:
    """Deduplicate merged-cell repeats within a row."""
    seen = []
    for c in row.cells:
        t = c.text.strip()
        if t and (not seen or t != seen[-1]):
            seen.append(t)
    return seen


def _render_table_as_text(table, table_index: int) -> str:
    """Render a docx table as readable text.

    For 2-column tables (key-value style), renders each row as ``【label】value``.
    Otherwise renders rows as pipe-joined cells.
    """
    rows = []
    col_count = max((len(row.cells) for row in table.rows), default=0)

    for ri, row in enumerate(table.rows):
        cells = _cell_texts(row)
        if not cells:
            continue
        if col_count == 2 and len(cells) == 2:
            label, value = cells[0], cells[1]
            if any(kw in label for kw in _TABLE_LAYOUT_KEY_LABELS):
                rows.append(f"【{label}】{value}")
            else:
                rows.append(f"{label}：{value}")
        else:
            rows.append(" | ".join(cells))

    if not rows:
        return ""
    return f"[表格{table_index + 1}]\n" + "\n".join(rows)


# Cell labels that must match exactly (avoid「实验名」⊂「实验名称」误报)
_EXACT_TABLE_MARKERS = frozenset({"实验名"})


def _table_marker_matches(marker: str, cell_text: str) -> bool:
    text = cell_text.strip()
    if marker in _EXACT_TABLE_MARKERS:
        return text == marker
    return marker in text


def _detect_table_layout(doc) -> dict[str, object]:
    """Walk all tables looking for training-report markers.

    Returns a dict with ``report_layout`` and ``table_map`` (cell coordinates
    of each matched marker), or an empty dict for standard paragraph reports.
    """
    table_map_entries: list[dict[str, object]] = []

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                for marker in _TRAINING_TABLE_MARKERS_SORTED:
                    if _table_marker_matches(marker, text):
                        table_map_entries.append({
                            "table": ti,
                            "row": ri,
                            "col": ci,
                            "label": marker,
                            "text_excerpt": text[:200],
                        })
                        break

    if table_map_entries:
        return {
            "report_layout": "training_table",
            "table_map": table_map_entries,
            "metadata_tables": [0, 1] if len(doc.tables) >= 2 else [0],
        }
    return {}


_MONO_FONT_MARKERS = ("courier", "consolas", "monaco", "menlo", "lucida console", "dejavu sans mono")
_CODE_CLOZE_BLANK_RE = re.compile(r"\(\s*\d+\s*\)|（\s*\d+\s*）")
_CODE_SEGMENT_HINT_RE = re.compile(
    r"\b(class|public|private|protected|extends|implements|interface|def|function|import|return)\b|//|/\*",
    re.IGNORECASE,
)


def _iter_docx_body_blocks(doc):
    """Yield Paragraph or Table elements in document body order."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield Paragraph(child, doc)
        elif tag == "tbl":
            yield Table(child, doc)


def _paragraph_is_monospace(para) -> bool:
    style_name = (para.style.name if para.style else "") or ""
    if "code" in style_name.lower() or "pre" in style_name.lower():
        return True
    for run in para.runs:
        name = (run.font.name or "").lower()
        if any(marker in name for marker in _MONO_FONT_MARKERS):
            return True
    return False


def _table_has_monospace(table) -> bool:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if _paragraph_is_monospace(para):
                    return True
    return False


def _table_cell_plain_text(table) -> str:
    """Flatten table cells to plain lines (no DA1 bracket labels)."""
    lines: list[str] = []
    seen_cells: set[str] = set()
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if not cell_text or cell_text in seen_cells:
                continue
            seen_cells.add(cell_text)
            for line in cell_text.splitlines():
                stripped = line.strip()
                if stripped:
                    lines.append(stripped)
    return "\n".join(lines)


def _segment_looks_like_code_cloze(text: str) -> bool:
    if not text:
        return False
    if detect_code_cloze(text).get("is_code_cloze"):
        return True
    return bool(_CODE_SEGMENT_HINT_RE.search(text) and _CODE_CLOZE_BLANK_RE.search(text))


def extract_docx_code_cloze_text(doc) -> str:
    """Collect code-like docx body segments for numbered-blank detection (R5 / Phase D)."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    segments: list[str] = []
    for block in _iter_docx_body_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if _paragraph_is_monospace(block) or _segment_looks_like_code_cloze(text):
                segments.append(text)
        elif isinstance(block, Table):
            plain = _table_cell_plain_text(block)
            if not plain:
                continue
            if _table_has_monospace(block) or _segment_looks_like_code_cloze(plain):
                segments.append(plain)
    return "\n\n".join(segments)


def detect_code_cloze_for_docx(full_text: str, doc) -> dict:
    """Detect code cloze from full docx text and extracted code segments."""
    candidates = [full_text or ""]
    code_text = extract_docx_code_cloze_text(doc)
    if code_text and code_text not in candidates:
        candidates.append(code_text)

    best = detect_code_cloze("")
    for text in candidates:
        result = detect_code_cloze(text)
        if result.get("is_code_cloze") and (
            not best.get("is_code_cloze")
            or result.get("blank_count", 0) > best.get("blank_count", 0)
        ):
            best = result
        elif not best.get("is_code_cloze") and result.get("blank_count", 0) > best.get("blank_count", 0):
            best = result
    return best


def extract_docx(path):
    """提取 Word 文档全文和封面表格元数据。"""
    metadata: dict[str, object] = {}
    if not DOCX_OK:
        return "（无法读取Word文档，请安装 python-docx）", metadata

    doc = Document(str(path))

    if doc.tables:
        for row in doc.tables[0].rows:
            cells = [c.text.strip() for c in row.cells]
            for i, cell in enumerate(cells):
                nxt = cells[i + 1] if i + 1 < len(cells) else ""
                if "课程名称" in cell:
                    metadata["course"] = nxt or cells[-1]
                if "实验序号" in cell or "实验名称" in cell:
                    for c in cells[i + 1 :]:
                        if c and len(c) < 40:
                            metadata["experiment_title"] = c
                            break
                if "专业" in cell and nxt:
                    metadata["major"] = nxt
                if "学号" in cell:
                    metadata["student_id"] = nxt
                if "姓名" in cell:
                    for c in cells[i + 1 :]:
                        if c:
                            metadata["student_name"] = c
                            break

    # Detect table layout and extract all table text
    layout = _detect_table_layout(doc)
    if layout:
        metadata.update(layout)

    table_texts = []
    for ti, table in enumerate(doc.tables):
        rendered = _render_table_as_text(table, ti)
        if rendered:
            table_texts.append(rendered)

    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(lines)
    if table_texts:
        full_text = full_text + "\n\n" + "\n\n".join(table_texts)

    # IM1: extract embedded images (lazy import to avoid circular dep)
    from document.extract_images import extract_docx_images

    image_result = extract_docx_images(path)
    metadata["image_assets"] = image_result.get("image_assets") or []
    metadata["image_bundle_meta"] = image_result.get("image_bundle_meta") or {}

    cloze = detect_code_cloze_for_docx(full_text, doc)
    if cloze.get("is_code_cloze"):
        metadata["code_cloze"] = cloze
        code_text = extract_docx_code_cloze_text(doc)
        if code_text:
            metadata["code_cloze_source"] = "docx_code_segments"

    return full_text, metadata


def parse_document(path, file_name: str | None = None):
    """Extract full_text and metadata from docx or pdf by extension."""
    path = Path(path)
    name = file_name or path.name
    if is_legacy_doc(name):
        from document.convert_doc import convert_doc_to_docx

        converted, err = convert_doc_to_docx(path)
        if converted and converted.exists():
            path = converted
            name = converted.name
        else:
            return "", {}, [{"code": "legacy_doc", "message": err or WARN_LEGACY_DOC["message"]}]
    if is_pdf(name):
        from document.extract_pdf import extract_pdf

        full_text, metadata, hints = extract_pdf(path)
        metadata.setdefault("source_format", "pdf")
        return full_text, metadata, hints
    full_text, metadata = extract_docx(path)
    metadata.setdefault("source_format", "docx")
    return full_text, metadata, []


def extract_document_paragraphs(path: Path, file_name: str) -> list[str]:
    """Paragraph list for combined-layout detection (docx or pdf)."""
    if is_pdf(file_name):
        full_text, _, _ = parse_document(path, file_name)
        return [ln for ln in full_text.split("\n") if ln.strip()]
    return [ln for ln in extract_docx_paragraphs(path)]


def extract_docx_paragraphs(path: Path) -> list[str]:
    if not DOCX_OK:
        full_text, _ = extract_docx(path)
        return [ln for ln in full_text.split("\n") if ln.strip()]
    from docx import Document

    doc = Document(str(path))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _merge_image_read_into_metadata(metadata: dict, image_read: dict) -> None:
    """Persist IM2 OCR fields on parse metadata."""
    if not image_read:
        return
    for key in (
        "image_reading_mode",
        "assignment_from_images",
        "image_sections",
        "image_ocr_merged",
        "image_read_summary",
        "document_assignment_text",
    ):
        if key in image_read:
            metadata[key] = image_read[key]


def build_question_from_document(
    path,
    file_name: str,
    *,
    enable_image_ocr: bool = False,
    ocr_lang: str = "chi_sim+eng",
    ocr_max_pages: int = 20,
    image_reading_mode: str = "ocr_only",
    vision_max_pages: int = 5,
    llm_settings: dict | None = None,
):
    """Parse docx/pdf and return question dict, metadata, full_text, warnings."""
    path = Path(path)
    if is_legacy_doc(file_name):
        from document.convert_doc import convert_doc_to_docx

        converted, err = convert_doc_to_docx(path)
        if not converted or not converted.exists():
            warnings = [{"code": "legacy_doc", "message": err or WARN_LEGACY_DOC["message"]}]
            question = {
                "id": 0,
                "type": "lab_report",
                "title": file_name.replace(".doc", ""),
                "full_text": "",
                "metadata": {},
                "placeholder": "",
                "image_assets": [],
                "image_bundle_meta": {},
            }
            return question, {}, "", warnings
        import shutil

        shutil.copy2(converted, path)
        file_name = converted.name

    full_text, metadata, hints = parse_document(path, file_name)
    from document.image_read import apply_image_reading

    ocr_warnings, image_read = apply_image_reading(
        full_text,
        metadata,
        enable_image_ocr=enable_image_ocr,
        ocr_lang=ocr_lang,
        ocr_max_pages=ocr_max_pages,
        hints=hints,
        image_reading_mode=image_reading_mode,
        vision_max_pages=vision_max_pages,
        llm_settings=llm_settings,
    )
    _merge_image_read_into_metadata(metadata, image_read)
    warnings = collect_parse_warnings(full_text, metadata, file_name)
    warnings.extend(_pdf_hints_to_warnings(hints))
    warnings.extend(ocr_warnings)
    augment_ocr_action_warnings(
        warnings,
        enable_image_ocr=enable_image_ocr,
        metadata=metadata,
    )
    fmt = metadata.get("source_format") or document_format(file_name)
    title_base = file_name.rsplit(".", 1)[0] if "." in file_name else file_name
    image_assets = metadata.get("image_assets") or []
    image_bundle_meta = metadata.get("image_bundle_meta") or {}
    question_metadata = {**metadata, "source_format": fmt}
    question_metadata.pop("image_assets", None)
    question_metadata.pop("image_bundle_meta", None)
    q_type = "lab_report"
    cloze = metadata.get("code_cloze") or detect_code_cloze(full_text)
    if cloze.get("is_code_cloze"):
        q_type = "code_cloze"
        question_metadata["code_cloze"] = cloze

    question = {
        "id": 0,
        "type": q_type,
        "title": metadata.get("experiment_title", title_base),
        "full_text": full_text,
        "metadata": question_metadata,
        "placeholder": "",
        "image_assets": image_assets,
        "image_bundle_meta": image_bundle_meta,
        "assignment_from_images": bool(metadata.get("assignment_from_images")),
        "image_reading_mode": metadata.get("image_reading_mode") or "",
        "image_read_summary": metadata.get("image_read_summary"),
        "image_sections": metadata.get("image_sections") or [],
        "assignment_text": metadata.get("document_assignment_text") or full_text,
    }
    return question, question_metadata, full_text, warnings


def build_question_from_docx(path, file_name: str):
    """Backward-compatible alias for build_question_from_document."""
    return build_question_from_document(path, file_name)


# --- O10 / R8: mixed assignment segment split (theory + code_cloze) ---

_FILL_SECTION_MARKERS = ("实验步骤", "实验结果", "实验总结", "三、", "四、", "五、")

_SEGMENT_HEADING_RES = (
    re.compile(r"^[一二三四五六七八九十]+[、．.]\s*.+"),
    re.compile(r"^第[一二三四五六七八九十\d]+题"),
    re.compile(r"^(简答|填空|代码|编程|选择|判断)题"),
)


def _is_assignment_segment_heading(line: str) -> bool:
    """Exam-style major heading; exclude code-cloze blank-only lines."""
    stripped = (line or "").strip()
    if not stripped or len(stripped) < 2:
        return False
    if _CODE_CLOZE_BLANK_RE.fullmatch(stripped.replace(" ", "")):
        return False
    if re.fullmatch(r"[\(（]\s*\d+\s*[\)）]", stripped):
        return False
    return any(pat.match(stripped) for pat in _SEGMENT_HEADING_RES)


def _split_text_lines_to_segments(lines: list[str]) -> list[dict[str, str]]:
    segments: list[dict[str, str]] = []
    heading = ""
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf, heading
        body = "\n".join(buf).strip()
        if body:
            segments.append({"heading": heading, "text": body})
        buf = []

    for line in lines:
        if _is_assignment_segment_heading(line):
            flush()
            heading = line.strip()
            continue
        buf.append(line)
    flush()
    return segments


def split_text_assignment_segments(text: str) -> list[dict[str, str]]:
    """Split pasted/plain assignment into ordered segments by exam headings."""
    body = (text or "").strip()
    lines = [ln for ln in body.split("\n") if ln.strip()]
    if not lines:
        return []
    segments = _split_text_lines_to_segments(lines)
    if len(segments) > 1:
        return segments
    if detect_code_cloze(body).get("is_code_cloze"):
        return [{"heading": "", "text": body}]
    return _split_by_code_cloze_islands(body)


def _line_in_code_island(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped:
        return False
    if _CODE_CLOZE_BLANK_RE.search(stripped):
        return True
    return bool(_CODE_SEGMENT_HINT_RE.search(stripped))


def _split_by_code_cloze_islands(text: str) -> list[dict[str, str]]:
    """Fallback: prose blocks vs code-cloze code blocks when no section headings."""
    lines = [ln for ln in (text or "").split("\n") if ln.strip()]
    if not lines:
        return []
    segments: list[dict[str, str]] = []
    i = 0
    while i < len(lines):
        prose: list[str] = []
        while i < len(lines) and not _line_in_code_island(lines[i]):
            prose.append(lines[i])
            i += 1
        if prose:
            segments.append({"heading": "", "text": "\n".join(prose)})

        code: list[str] = []
        while i < len(lines):
            if not code and not _line_in_code_island(lines[i]):
                break
            if code and not _line_in_code_island(lines[i]) and not _CODE_CLOZE_BLANK_RE.search(
                lines[i]
            ):
                if _segment_looks_like_code_cloze("\n".join(code)):
                    break
            code.append(lines[i])
            i += 1
        if code:
            body = "\n".join(code)
            if _segment_looks_like_code_cloze(body) or detect_code_cloze(body).get(
                "is_code_cloze"
            ):
                segments.append({"heading": "", "text": body})
            elif segments:
                segments[-1]["text"] = segments[-1]["text"] + "\n" + body
            else:
                segments.append({"heading": "", "text": body})
    return [s for s in segments if (s.get("text") or "").strip()]


def split_docx_assignment_segments(doc) -> list[dict[str, str]]:
    """Split docx body into assignment segments (paragraphs + code tables)."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    segments: list[dict[str, str]] = []
    heading = ""
    buf: list[str] = []

    def flush() -> None:
        nonlocal buf, heading
        body = "\n".join(buf).strip()
        if body:
            segments.append({"heading": heading, "text": body})
        buf = []

    for block in _iter_docx_body_blocks(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if _is_assignment_segment_heading(text):
                flush()
                heading = text
                continue
            buf.append(text)
        elif isinstance(block, Table):
            plain = _table_cell_plain_text(block)
            if not plain:
                continue
            if _table_has_monospace(block) or _segment_looks_like_code_cloze(plain):
                flush()
                segments.append({"heading": heading or "代码填空", "text": plain})
                heading = ""
            else:
                for line in plain.splitlines():
                    stripped = line.strip()
                    if stripped:
                        buf.append(stripped)
    flush()
    if len(segments) <= 1:
        full = "\n\n".join(s["text"] for s in segments)
        return split_text_assignment_segments(full)
    return segments


def classify_assignment_segment(text: str) -> tuple[str, dict[str, object]]:
    """Return question type + segment metadata."""
    body = (text or "").strip()
    if not body:
        return "theory", {}
    cloze = detect_code_cloze(body)
    if cloze.get("is_code_cloze"):
        return "code_cloze", {"code_cloze": cloze}
    if any(m in body for m in _FILL_SECTION_MARKERS):
        return "lab_report", {}
    return "theory", {}


def build_questions_from_segments(
    segments: list[dict[str, str]],
    *,
    title_base: str = "",
) -> list[dict[str, object]]:
    questions: list[dict[str, object]] = []
    for i, seg in enumerate(segments):
        text = (seg.get("text") or "").strip()
        if not text:
            continue
        q_type, seg_meta = classify_assignment_segment(text)
        heading = (seg.get("heading") or "").strip()
        if heading and len(heading) <= 80:
            title = heading
        elif title_base:
            title = f"{title_base} · 第 {len(questions) + 1} 题"
        else:
            title = f"题目 {len(questions) + 1}"
        qmeta: dict[str, object] = {
            "segment_index": len(questions),
            "segment_heading": heading,
        }
        qmeta.update(seg_meta)
        questions.append(
            {
                "id": len(questions),
                "type": q_type,
                "title": title,
                "full_text": text,
                "content": text[:200],
                "metadata": qmeta,
                "placeholder": "",
                "image_assets": [],
                "image_bundle_meta": {},
            }
        )
    return questions


_MIN_THEORY_SEGMENT_CHARS = 80


def should_use_mixed_assignment(questions: list[dict]) -> bool:
    """True when substantial theory text coexists with at least one code_cloze segment."""
    if len(questions) < 2:
        return False
    cloze = [q for q in questions if q.get("type") == "code_cloze"]
    if not cloze:
        return False
    theory_text = "\n\n".join(
        (q.get("full_text") or "").strip()
        for q in questions
        if q.get("type") == "theory"
    ).strip()
    return len(theory_text) >= _MIN_THEORY_SEGMENT_CHARS


def format_mixed_assignment_text(questions: list[dict]) -> str:
    parts: list[str] = []
    for q in questions:
        title = (q.get("title") or f"题目 {int(q.get('id', 0)) + 1}").strip()
        body = (q.get("full_text") or "").strip()
        if body:
            parts.append(f"【{title}】\n{body}")
    return "\n\n".join(parts)


def detect_docx_sections(path):
    """
    Detect section structure from a docx file.

    Returns dict with:
        sections_detected: [{index, heading, semantic}, ...]
        section_map: {steps/result/summary: {type, heading, para_index} or None}
        fill_hints: {screenshots_target, merge_*, ...}
        report_layout: "training_table" | "standard_sections" | "variant_sections" | None
        table_map: [{table, row, col, label, text_excerpt}, ...]
    """
    from modules.fill_report import _build_fill_hints, detect_sections

    if not DOCX_OK:
        return {}

    from docx import Document

    doc = Document(str(path))
    paragraphs = list(doc.paragraphs)
    sections_detected, section_map = detect_sections(paragraphs)
    fill_hints = _build_fill_hints(section_map)

    # Determine report_layout from detected sections + table detection
    layout_info = _detect_table_layout(doc)
    report_layout = layout_info.get("report_layout") or ""
    if not report_layout:
        mapped_count = sum(1 for v in section_map.values() if v)
        if mapped_count == 3:
            has_standard = all(
                section_map[k] and any(
                    str(n) in (section_map[k].get("heading") or "")
                    for n in ["三", "四", "五"]
                )
                for k in ["steps", "result", "summary"]
            )
            report_layout = "standard_sections" if has_standard else "variant_sections"
        elif mapped_count >= 1:
            report_layout = "variant_sections"

    return {
        "sections_detected": sections_detected,
        "section_map": section_map,
        "fill_hints": fill_hints,
        "report_layout": report_layout,
        "table_map": layout_info.get("table_map") or [],
    }
