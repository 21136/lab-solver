"""
PDF fill/export helpers (Phase 3b).

When the uploaded report is PDF, fill_report writes a .docx (paired template or generated shell).
"""

from __future__ import annotations

import uuid
from pathlib import Path
from typing import Any, Optional

from config import DOCX_OK, TEMP_DIR
from modules.fill_report import detect_sections
from modules.parse_report import document_format, is_legacy_doc, is_pdf

_DEFAULT_SECTIONS = (
    "三、实验步骤",
    "四、实验结果",
    "五、实验总结",
)

def resolve_fill_target_info(
    bundles: list[dict[str, Any]],
    primary: dict[str, Any],
) -> dict[str, Any]:
    """
    Resolve where fill_report should write.

    Returns:
        format: always docx for export
        path: existing docx path or None (generate on fill)
        from: docx | user_template | generated
        source_format: pdf | docx (original fill_target upload)
    """
    src_fmt = (primary.get("metadata") or {}).get("source_format") or document_format(
        primary.get("file_name") or ""
    )
    _pdf_export_message = "原版式 PDF 无法直接填回，已按解析出的章节生成 Word 并写入内容"

    def _is_docx_bundle(b: dict[str, Any]) -> bool:
        return document_format(b.get("file_name") or "") == "docx"

    templates = [b for b in bundles if b.get("role") == "fill_template" and _is_docx_bundle(b)]
    if templates:
        target = templates[0]
        return {
            "format": "docx",
            "path": target.get("file_path") or "",
            "from": "user_template",
            "source_format": src_fmt,
            "file_name": target.get("file_name") or "report.docx",
        }

    docx_targets = [
        b for b in bundles if b.get("role") == "fill_target" and _is_docx_bundle(b)
    ]
    if docx_targets:
        target = docx_targets[0]
        return {
            "format": "docx",
            "path": target.get("file_path") or "",
            "from": "docx",
            "source_format": "docx",
            "file_name": target.get("file_name") or "report.docx",
        }

    if src_fmt == "docx" or not is_pdf(primary.get("file_name") or ""):
        return {
            "format": "docx",
            "path": primary.get("file_path") or "",
            "from": "docx",
            "source_format": "docx",
            "file_name": primary.get("file_name") or "report.docx",
        }

    return {
        "format": "docx",
        "path": None,
        "from": "generated",
        "source_format": "pdf",
        "file_name": primary.get("file_name") or "report.pdf",
        "message": _pdf_export_message,
    }


def _section_lines_from_text(full_text: str) -> list[str]:
    """Extract section heading lines using semantic detection (DA2)."""
    lines = [(ln, ln.strip()) for ln in (full_text or "").splitlines()]
    stripped_lines = [s for _, s in lines if s and len(s) < 40]
    if not stripped_lines:
        return []
    sections, _section_map = detect_sections(stripped_lines)
    found = []
    seen = set()
    for sec in sections:
        heading = sec["heading"]
        if heading not in seen:
            found.append(heading)
            seen.add(heading)
    return found


def generate_docx_shell(
    full_text: str,
    metadata: Optional[dict[str, Any]] = None,
    output_path: Optional[Path | str] = None,
) -> Path:
    """Build a minimal .docx from parsed PDF/doc text (Strategy A — export Word)."""
    if not DOCX_OK:
        raise RuntimeError("python-docx 不可用，无法从 PDF 导出 Word")

    from docx import Document

    meta = metadata or {}
    out = Path(output_path) if output_path else TEMP_DIR / f"pdf_shell_{uuid.uuid4().hex[:8]}.docx"

    doc = Document()
    title_bits = [
        meta.get("course"),
        meta.get("experiment_title"),
        meta.get("major"),
    ]
    title = " — ".join(x for x in title_bits if x)
    if title:
        doc.add_paragraph(title)

    sections = _section_lines_from_text(full_text)
    if not sections:
        sections = list(_DEFAULT_SECTIONS)

    for heading in sections:
        doc.add_paragraph(heading)
        doc.add_paragraph("")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def prepare_fill_docx_for_fill(
    input_path: Optional[Path | str],
    file_name: str,
    *,
    source_format: Optional[str] = None,
    paired_docx_path: Optional[Path | str] = None,
    fill_body_text: str = "",
    metadata: Optional[dict[str, Any]] = None,
    shell_output_path: Optional[Path | str] = None,
) -> tuple[Optional[Path], dict[str, Any]]:
    """
    Resolve the .docx path passed to do_fill.

    - Normal .docx upload → same file (unchanged behavior).
    - PDF + paired docx → paired template.
    - PDF only → generated shell from section headers.
    """
    name = file_name or "report.docx"
    fmt = source_format or document_format(name)
    meta = dict(metadata or {})

    if paired_docx_path:
        p = Path(paired_docx_path)
        if p.exists() and p.suffix.lower() == ".docx":
            return p, {
                "format": "docx",
                "from": "user_template",
                "source_format": fmt,
                "export_format": "docx",
            }

    if input_path:
        inp = Path(input_path)
        if inp.exists() and (inp.suffix.lower() == ".docx" or is_legacy_doc(name)):
            if is_legacy_doc(name):
                from .convert_doc import convert_doc_to_docx

                converted, err = convert_doc_to_docx(inp)
                if converted and converted.exists():
                    inp = converted
                else:
                    raise ValueError(err or "旧版 .doc 无法转换，请安装 LibreOffice 或在 Word 中另存为 .docx")
            if paired_docx_path or fmt == "pdf" or is_pdf(name):
                from_kind = "user_template"
            else:
                from_kind = "docx"
            return inp, {
                "format": "docx",
                "from": from_kind,
                "source_format": fmt if from_kind == "user_template" else "docx",
                "export_format": "docx",
            }

    if fmt == "pdf" or is_pdf(name):
        body = fill_body_text or meta.get("fill_body_text") or meta.get("full_text") or ""
        shell_path = generate_docx_shell(
            body,
            metadata=meta,
            output_path=shell_output_path,
        )
        return shell_path, {
            "format": "docx",
            "from": "generated",
            "source_format": "pdf",
            "export_format": "docx",
            "message": "原版式 PDF 无法直接填回，已按解析出的章节生成 Word 并写入内容",
        }

    if input_path and Path(input_path).exists():
        return Path(input_path), {
            "format": "docx",
            "from": "docx",
            "source_format": fmt,
            "export_format": "docx",
        }

    return None, {
        "format": "docx",
        "from": "generated",
        "source_format": fmt,
        "export_format": "docx",
    }
