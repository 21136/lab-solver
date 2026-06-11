"""IM1 + IM2-a: docx embedded images and OCR integration tests."""

import base64
import hashlib
import io
import sys
import tempfile
from pathlib import Path
import pytest
from PIL import Image

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src" / "python"))

from document.extract_images import (
    _body_texts,
    _guess_role,
    _image_pixel_size,
    _nearby_text,
    _sha256_hex,
    extract_docx_images,
)
from document import image_read
from config import PDF_OK
from document.extract_pdf import extract_pdf, render_pdf_pages
from document.image_read import (
    detect_multi_question_in_image,
    merge_assignment_from_images,
    multi_question_image_warnings,
    ocr_batch,
    should_run_ocr,
)
from document.user_upload_images import build_user_upload_assets, process_user_upload_images
from modules.parse_report import build_question_from_document, extract_docx, parse_document

FIXTURES = Path(__file__).resolve().parent / "fixtures" / "image_input"

# ── helpers ────────────────────────────────────────────────────────────


def _make_test_docx(paragraphs_before=0, images=None, paragraphs_after=0):
    """Create a temp docx with optional images at a known position. Returns path."""
    from docx import Document
    from docx.shared import Inches

    if images is None:
        images = []
    doc = Document()
    for text in (paragraphs_before if isinstance(paragraphs_before, list)
                 else [f"Para {i}" for i in range(paragraphs_before)]):
        doc.add_paragraph(str(text))
    for img_path in images:
        doc.add_picture(str(img_path), width=Inches(2))
    for text in (paragraphs_after if isinstance(paragraphs_after, list)
                 else [f"Para after {i}" for i in range(paragraphs_after)]):
        doc.add_paragraph(str(text))

    tmp = tempfile.gettempdir() / Path(f"test_img_{hash(str(images))}.docx")
    # Use a simpler path
    tmp = Path(tempfile.gettempdir()) / f"test_im1_{len(images)}img.docx"
    doc.save(str(tmp))
    return tmp


def _make_test_png(width=200, height=100, color="red"):
    """Create a small test PNG in a temp file. Returns path."""
    img = Image.new("RGB", (width, height), color=color)
    tmp = Path(tempfile.gettempdir()) / f"test_{width}x{height}_{color}.png"
    img.save(str(tmp))
    return tmp


def _sha256_file(path):
    return hashlib.sha256(Path(path).read_bytes()).hexdigest()


# ── unit tests ─────────────────────────────────────────────────────────


class TestSha256:
    def test_deterministic(self):
        assert _sha256_hex(b"hello") == _sha256_hex(b"hello")

    def test_different_data(self):
        assert _sha256_hex(b"hello") != _sha256_hex(b"world")


class TestImagePixelSize:
    def test_known_size(self):
        tmp = _make_test_png(200, 100)
        blob = tmp.read_bytes()
        w, h = _image_pixel_size(blob)
        assert w == 200
        assert h == 100

    def test_invalid_blob(self):
        w, h = _image_pixel_size(b"not an image")
        assert w == 0
        assert h == 0


class TestGuessRole:
    def test_signature_keyword(self):
        assert _guess_role("请在此处签名", 300, 300) == "signature"
        assert _guess_role("盖章处", 300, 300) == "signature"

    def test_signature_size(self):
        assert _guess_role("", 150, 50) == "signature"

    def test_decoration_keyword(self):
        assert _guess_role("学校logo", 300, 300) == "decoration"

    def test_decoration_size(self):
        assert _guess_role("", 40, 40) == "decoration"

    def test_assignment_keyword(self):
        assert _guess_role("实验目的", 300, 300) == "assignment"
        assert _guess_role("电路图如下", 300, 300) == "assignment"

    def test_assignment_by_size(self):
        assert _guess_role("", 500, 400) == "assignment"

    def test_unknown(self):
        assert _guess_role("", 250, 150) == "unknown"


class TestNearbyText:
    def test_surrounding_text(self):
        texts = ["one", "", "two", "three", ""]
        result = _nearby_text(texts, 1)  # center element
        assert "one" in result
        assert "two" in result

    def test_edge_position(self):
        texts = ["first", "second", "third"]
        result = _nearby_text(texts, 0)  # first element
        assert "second" in result
        assert "first" not in result

    def test_empty_texts(self):
        result = _nearby_text([], 0)
        assert result == ""


class TestBodyTexts:
    def test_empty(self):
        assert _body_texts([]) == []


# ── integration tests ──────────────────────────────────────────────────


class TestExtractDocxImages:
    def test_no_images(self):
        path = _make_test_docx(paragraphs_before=3)
        result = extract_docx_images(path)
        assert result["image_assets"] == []
        assert result["image_bundle_meta"]["total"] == 0
        assert result["image_bundle_meta"]["deduped"] == 0

    def test_single_image(self):
        img = _make_test_png(300, 200, "blue")
        path = _make_test_docx(
            paragraphs_before=["实验目的"],
            images=[img],
            paragraphs_after=["请完成实验"],
        )
        result = extract_docx_images(path)
        assert len(result["image_assets"]) == 1
        asset = result["image_assets"][0]
        assert asset["source"] == "docx_inline"
        assert asset["mime"] == "image/png"
        assert asset["order"] == 0
        assert len(asset["bytes_b64"]) > 0
        assert asset["sha256"] == _sha256_file(img)
        assert "实验目的" in asset["nearby_text"]
        assert "请完成实验" in asset["nearby_text"]

    def test_multiple_images_ordered(self):
        img1 = _make_test_png(300, 200, "red")
        img2 = _make_test_png(400, 300, "green")
        img3 = _make_test_png(500, 400, "blue")
        path = _make_test_docx(
            paragraphs_before=["Start"],
            images=[img1, img2, img3],
            paragraphs_after=["End"],
        )
        result = extract_docx_images(path)
        assert len(result["image_assets"]) == 3
        assert result["image_assets"][0]["order"] == 0
        assert result["image_assets"][1]["order"] == 1
        assert result["image_assets"][2]["order"] == 2
        # Verify ordering through sha256
        assert result["image_assets"][0]["sha256"] == _sha256_file(img1)
        assert result["image_assets"][1]["sha256"] == _sha256_file(img2)
        assert result["image_assets"][2]["sha256"] == _sha256_file(img3)

    def test_dedup_same_image(self):
        img = _make_test_png(200, 100, "red")
        path = _make_test_docx(
            paragraphs_before=[],
            images=[img, img],  # same image twice
            paragraphs_after=["End"],
        )
        result = extract_docx_images(path)
        assert len(result["image_assets"]) == 1
        assert result["image_bundle_meta"]["total"] == 2
        assert result["image_bundle_meta"]["deduped"] == 1

    def test_role_assignment_with_context(self):
        img = _make_test_png(600, 400, "white")
        path = _make_test_docx(
            paragraphs_before=["实验目的及要求", "1. 验证欧姆定律"],
            images=[img],
            paragraphs_after=["请完成实验并记录数据"],
        )
        result = extract_docx_images(path)
        assert result["image_assets"][0]["role_guess"] == "assignment"

    def test_large_image_is_assignment(self):
        img = _make_test_png(800, 600)
        path = _make_test_docx(images=[img])
        result = extract_docx_images(path)
        assert result["image_assets"][0]["role_guess"] == "assignment"

    def test_small_image_is_signature(self):
        img = _make_test_png(150, 40)
        path = _make_test_docx(images=[img])
        result = extract_docx_images(path)
        assert result["image_assets"][0]["role_guess"] == "signature"

    def test_bytes_b64_valid_base64(self):
        img = _make_test_png(100, 100)
        path = _make_test_docx(images=[img])
        result = extract_docx_images(path)
        b64 = result["image_assets"][0]["bytes_b64"]
        decoded = base64.b64decode(b64)
        assert len(decoded) > 0
        assert decoded == Path(img).read_bytes()

    def test_page_hint(self):
        img = _make_test_png(100, 100)
        path = _make_test_docx(images=[img])
        result = extract_docx_images(path)
        assert result["image_assets"][0]["page_hint"] >= 1

    def test_ocr_vision_fields_present(self):
        img = _make_test_png(100, 100)
        path = _make_test_docx(images=[img])
        result = extract_docx_images(path)
        assert result["image_assets"][0]["ocr_text"] == ""
        assert result["image_assets"][0]["vision_summary"] == ""


# ── parse_report integration tests ─────────────────────────────────────


class TestParseReportIntegration:
    def test_extract_docx_has_image_assets(self):
        img = _make_test_png(300, 200)
        path = _make_test_docx(
            paragraphs_before=["题目要求"],
            images=[img],
        )
        full_text, metadata = extract_docx(path)
        assert "image_assets" in metadata
        assert len(metadata["image_assets"]) == 1
        assert "image_bundle_meta" in metadata

    def test_extract_docx_no_images(self):
        path = _make_test_docx(paragraphs_before=["纯文本报告"], images=[])
        full_text, metadata = extract_docx(path)
        assert metadata.get("image_assets") == []
        assert metadata.get("image_bundle_meta", {}).get("total") == 0

    def test_build_question_has_image_fields(self):
        img = _make_test_png(300, 200)
        path = _make_test_docx(
            paragraphs_before=["实验报告正文"],
            images=[img],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path, "test.docx"
        )
        assert "image_assets" in question
        assert len(question["image_assets"]) == 1
        assert "image_bundle_meta" in question

    def test_build_question_no_images(self):
        path = _make_test_docx(paragraphs_before=["纯文本"])
        question, metadata, full_text, warnings = build_question_from_document(
            path, "test.docx"
        )
        assert question.get("image_assets") == []
        assert question["image_bundle_meta"]["total"] == 0

    def test_warnings_with_images(self):
        img = _make_test_png(500, 400)
        path = _make_test_docx(
            paragraphs_before=["题目"],
            images=[img],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path, "test.docx"
        )
        assert len(warnings) >= 1
        # Should mention images since text is short
        codes = [w.get("code") for w in warnings]
        assert "short_text_with_images" in codes or "short_text" in codes

    def test_multiple_assignment_images_warning(self):
        img = _make_test_png(600, 400)
        path = _make_test_docx(
            paragraphs_before=["题目"],
            images=[img, _make_test_png(600, 400, "blue"),
                    _make_test_png(600, 400, "green")],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path, "test.docx"
        )
        codes = [w.get("code") for w in warnings]
        assert "multiple_assignment_images" in codes

    def test_standard_docx_regression(self):
        """V1 standard docx with no images must still parse correctly."""
        path = _make_test_docx(
            paragraphs_before=[
                "实验目的：验证欧姆定律",
                "实验原理：R=U/I",
                "三、实验步骤",
                "1. 连接电路",
                "2. 测量数据",
                "四、实验结果",
                "数据如下...",
                "五、实验总结",
                "本次实验成功验证了欧姆定律",
            ]
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path, "test.docx"
        )
        assert len(full_text) > 50
        assert question.get("image_assets") == []
        assert question["image_bundle_meta"]["total"] == 0
        # Standard sections should still be detected
        assert "实验目的" in full_text
        assert "实验步骤" in full_text


# ── multi-doc integration ──────────────────────────────────────────────


class TestMultiDocIntegration:
    def test_parse_single_file_has_image_assets(self):
        """parse_single_file should propagate image_assets to metadata."""
        img = _make_test_png(300, 200)
        docx_path = _make_test_docx(
            paragraphs_before=["实验要求"],
            images=[img],
        )
        from agent.parse_documents import parse_single_file

        file_bytes = docx_path.read_bytes()
        bundle = parse_single_file(file_bytes, "test.docx", role="fill_target")
        assert "image_assets" in bundle["metadata"]
        assert len(bundle["metadata"]["image_assets"]) == 1

    def test_parse_documents_list_collects_images(self):
        """parse_documents_list should aggregate images from all docs."""
        img = _make_test_png(300, 200)
        docx_path = _make_test_docx(
            paragraphs_before=["实验步骤", "实验内容较多"],
            images=[img],
        )
        from agent.parse_documents import parse_documents_list

        docs = [{
            "file_data": base64.b64encode(docx_path.read_bytes()).decode(),
            "file_name": "report.docx",
            "role": "fill_target",
        }]
        result = parse_documents_list(docs)
        # The fill_target's bundle should be in _bundles
        bundles = result.get("_bundles") or []
        assert len(bundles) >= 1
        meta = bundles[0].get("metadata") or {}
        assert len(meta.get("image_assets") or []) == 1


# ── IM2-a OCR tests ───────────────────────────────────────────────────


def _mock_ocr_text_for_asset(asset):
    """Return deterministic OCR based on fixture sha or generic text."""
    return {
        "ocr_text": "实验目的：验证欧姆定律。要求：测量电阻电压。",
        "ocr_confidence": 0.88,
        "ocr_status": "ok",
        "ocr_engine": "tesseract",
        "ocr_lang": "chi_sim+eng",
        "ocr_error": "",
    }


@pytest.fixture
def mock_tesseract(monkeypatch):
    monkeypatch.setattr(image_read, "OCR_OK", True)

    def _fake(asset, **kwargs):
        return _mock_ocr_text_for_asset(asset)

    monkeypatch.setattr(image_read, "ocr_image_asset", _fake)


class TestImageReadMerge:
    def test_merge_appends_to_body(self):
        assets = [{
            "id": "img_001",
            "order": 0,
            "ocr_text": "题目在图中",
            "ocr_status": "ok",
            "nearby_text": "实验要求",
        }]
        out = merge_assignment_from_images("正文很短", assets)
        assert "正文很短" in out["assignment_text"]
        assert "题目在图中" in out["assignment_text"]
        assert "--- 图 1（OCR）---" in out["assignment_text"]
        assert out["assignment_from_images"] is True
        assert len(out["image_sections"]) == 1

    def test_merge_ocr_only_when_body_empty(self):
        assets = [{
            "id": "img_001",
            "order": 0,
            "ocr_text": "纯图题目",
            "ocr_status": "ok",
        }]
        out = merge_assignment_from_images("", assets)
        assert out["assignment_text"] == "--- 图 1（OCR）---\n\n纯图题目"
        assert out["assignment_from_images"] is True

    def test_merge_skips_empty_status(self):
        assets = [{
            "id": "img_001",
            "order": 0,
            "ocr_text": "",
            "ocr_status": "empty",
        }]
        out = merge_assignment_from_images("正文", assets)
        assert out["assignment_text"] == "正文"
        assert out["assignment_from_images"] is False


class TestOcrTrigger:
    def test_auto_trigger_short_body_assignment_image(self):
        assets = [{"role_guess": "assignment", "order": 0}]
        assert should_run_ocr(50, assets, enable_image_ocr=False) is True

    def test_no_trigger_sufficient_body(self):
        assets = [{"role_guess": "assignment", "order": 0}]
        assert should_run_ocr(500, assets, enable_image_ocr=False) is False

    def test_enable_flag_triggers_unknown(self):
        assets = [{"role_guess": "unknown", "order": 0}]
        assert should_run_ocr(500, assets, enable_image_ocr=True) is True

    def test_skip_signature(self):
        assets = [{"role_guess": "signature", "order": 0}]
        assert should_run_ocr(10, assets, enable_image_ocr=True) is False

    def test_pdf_scanned_auto_triggers(self):
        assets = [{"role_guess": "assignment", "source": "pdf_page_render", "order": 0}]
        hints = [{"code": "pdf_scanned", "message": "scan"}]
        assert should_run_ocr(0, assets, enable_image_ocr=False, hints=hints) is True


class TestIm2bSettings:
    def test_settings_schema_ocr_defaults(self):
        from settings_schema import SETTINGS_DEFAULTS, SETTINGS_SCHEMA_VERSION

        assert SETTINGS_SCHEMA_VERSION == 10
        assert SETTINGS_DEFAULTS["autoFastTierForLightQuestions"] is True
        assert SETTINGS_DEFAULTS["enableParallelModuleSteps"] is True
        assert SETTINGS_DEFAULTS["solvePipelineVersion"] == "v4"
        assert SETTINGS_DEFAULTS["autoRemediate"] is True
        assert SETTINGS_DEFAULTS["autoRemediateMaxRounds"] == 1
        assert SETTINGS_DEFAULTS["enableImageOcr"] is False
        assert SETTINGS_DEFAULTS["imageOcrLang"] == "chi_sim+eng"
        assert SETTINGS_DEFAULTS["imageReadingMode"] == "ocr_only"
        assert SETTINGS_DEFAULTS["imageOcrMaxPages"] == 20
        assert SETTINGS_DEFAULTS["imageVisionMaxPages"] == 5


class TestIm2bOcrWarnings:
    def test_ocr_suggested_has_action_when_ocr_off(self):
        long_body = ["这是一段足够长的正文内容，用于避免自动 OCR 触发。"] * 3
        img = _make_test_png(250, 150, "gray")
        path = _make_test_docx(
            paragraphs_before=long_body,
            images=[img],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "warn_ocr.docx",
            enable_image_ocr=False,
        )
        assert len(question.get("image_assets") or []) == 1
        actionable = [w for w in warnings if w.get("action") == "enable_ocr_reparse"]
        assert len(actionable) >= 1

    def test_no_ocr_action_when_ocr_on(self, mock_tesseract):
        img = FIXTURES / "ocr_simple_zh.png"
        if not img.exists():
            pytest.skip("fixture missing")
        path = _make_test_docx(paragraphs_before=["题"], images=[img])
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "ocr_on.docx",
            enable_image_ocr=True,
        )
        assert not any(w.get("action") == "enable_ocr_reparse" for w in warnings)
        assert metadata.get("assignment_from_images") or question.get("assignment_from_images")


class TestOcrIntegration:
    def test_o1_short_body_ocr_enabled(self, mock_tesseract):
        img = FIXTURES / "ocr_simple_zh.png"
        if not img.exists():
            pytest.skip("fixture missing")
        path = _make_test_docx(
            paragraphs_before=["封面"],
            images=[img],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "short_ocr.docx",
            enable_image_ocr=True,
        )
        assert "欧姆定律" in (metadata.get("image_ocr_merged") or "")
        asset = question["image_assets"][0]
        assert asset.get("ocr_status") == "ok"
        assert asset.get("ocr_text")

    def test_o2_long_body_no_ocr_by_default(self, mock_tesseract):
        long_body = ["实验目的：验证欧姆定律"] + [
            f"段落 {i}：详细实验说明与数据记录要求。" for i in range(20)
        ]
        img = _make_test_png(600, 400)
        path = _make_test_docx(paragraphs_before=long_body, images=[img])
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "long_body.docx",
            enable_image_ocr=False,
        )
        asset = question["image_assets"][0]
        assert asset.get("ocr_text", "") == ""
        assert not metadata.get("assignment_from_images")

    def test_o4_no_tesseract_graceful(self, monkeypatch):
        monkeypatch.setattr(image_read, "OCR_OK", False)
        img = _make_test_png(600, 400)
        path = _make_test_docx(paragraphs_before=["短"], images=[img])
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "no_ocr.docx",
            enable_image_ocr=True,
        )
        codes = [w.get("code") for w in warnings]
        assert "ocr_unavailable" in codes
        assert question["image_assets"][0].get("ocr_status") == "skipped"

    def test_standard_docx_no_ocr_regression(self, mock_tesseract):
        path = _make_test_docx(
            paragraphs_before=[
                "实验目的：验证欧姆定律",
                "实验原理：R=U/I",
                "三、实验步骤",
                "1. 连接电路",
                "2. 测量数据",
                "四、实验结果",
                "数据如下...",
                "五、实验总结",
                "本次实验成功验证了欧姆定律",
            ],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path, "standard.docx"
        )
        assert question.get("image_assets") == []
        assert not metadata.get("assignment_from_images")

    def test_planner_input_includes_ocr(self, mock_tesseract):
        img = FIXTURES / "ocr_simple_zh.png"
        if not img.exists():
            pytest.skip("fixture missing")
        docx_path = _make_test_docx(paragraphs_before=["题"], images=[img])
        from agent.parse_documents import parse_single_file

        bundle = parse_single_file(
            docx_path.read_bytes(),
            "ocr_fill.docx",
            role="fill_target",
            enable_image_ocr=True,
        )
        assert "欧姆定律" in bundle.get("assignment_text", "")
        assert "欧姆定律" in bundle.get("planner_input_text", "")
        assert bundle.get("assignment_from_images") is True

    def test_ocr_dedup_by_sha256(self, mock_tesseract):
        img = _make_test_png(600, 400)
        path = _make_test_docx(
            paragraphs_before=["短"],
            images=[img, img],
            paragraphs_after=["尾"],
        )
        result = extract_docx_images(path)
        assets, summary = ocr_batch(
            result["image_assets"],
            enable_image_ocr=True,
            body_len=10,
        )
        assert summary["ocr_attempted"] == 1
        assert assets[0]["ocr_text"] == assets[0]["ocr_text"]

    def test_o5_no_llm_import_in_image_read_ocr_path(self):
        """OCR-only path must not eagerly import llm_client at module load."""
        import importlib

        mod = importlib.import_module("document.image_read")
        source = Path(mod.__file__).read_text(encoding="utf-8")
        assert "from llm_client import" not in source.split("def vision_image_asset")[0]


# ── IM3 scanned PDF tests ──────────────────────────────────────────────


def _ensure_scanned_pdf_fixture() -> Path:
    pdf = FIXTURES / "scanned_5page.pdf"
    if not pdf.exists():
        import importlib.util

        spec = importlib.util.spec_from_file_location(
            "gen_fixtures", FIXTURES / "gen_fixtures.py"
        )
        mod = importlib.util.module_from_spec(spec)
        assert spec.loader is not None
        spec.loader.exec_module(mod)
        mod.make_scanned_pdf(pdf)
    return pdf


def _mock_ocr_by_page(asset):
    page = asset.get("page_hint") or asset.get("order", 0) + 1
    texts = {
        1: "第1页：实验目的。验证欧姆定律。",
        2: "第2页：实验原理。R 等于 U 除以 I。",
        3: "第3页：实验步骤。连接电路并测量。",
        4: "第4页：实验数据。记录电压电流值。",
        5: "第5页：实验总结。完成本次实验。",
    }
    return {
        "ocr_text": texts.get(page, f"page {page}"),
        "ocr_confidence": 0.9,
        "ocr_status": "ok",
        "ocr_engine": "tesseract",
        "ocr_lang": "chi_sim+eng",
        "ocr_error": "",
    }


@pytest.fixture
def mock_tesseract_pages(monkeypatch):
    monkeypatch.setattr(image_read, "OCR_OK", True)

    def _fake(asset, **kwargs):
        return _mock_ocr_by_page(asset)

    monkeypatch.setattr(image_read, "ocr_image_asset", _fake)


class TestExtractPdfRender:
    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_scanned_pdf_renders_pages(self):
        pdf = _ensure_scanned_pdf_fixture()
        full_text, metadata, hints = extract_pdf(pdf)
        assert len(full_text.strip()) < 80
        assets = metadata.get("image_assets") or []
        assert len(assets) == 5
        assert all(a.get("source") == "pdf_page_render" for a in assets)
        assert [a.get("page_hint") for a in assets] == [1, 2, 3, 4, 5]
        codes = [h.get("code") for h in hints]
        assert "pdf_scanned" in codes

    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_render_pdf_pages_order(self):
        import fitz

        pdf = _ensure_scanned_pdf_fixture()
        doc = fitz.open(str(pdf))
        try:
            assets = render_pdf_pages(doc)
        finally:
            doc.close()
        assert len(assets) == 5
        assert assets[0]["order"] == 0
        assert assets[4]["order"] == 4
        assert all(len(a.get("bytes_b64") or "") > 0 for a in assets)


class TestIm3ScannedPdfOcr:
    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_o3_i3_five_page_merge(self, mock_tesseract_pages):
        pdf = _ensure_scanned_pdf_fixture()
        question, metadata, full_text, warnings = build_question_from_document(
            pdf,
            "scanned_5page.pdf",
            enable_image_ocr=False,
        )
        assets = question.get("image_assets") or []
        assert len(assets) == 5
        assert metadata.get("assignment_from_images") is True
        sections = metadata.get("image_sections") or []
        assert len(sections) == 5
        assignment = metadata.get("document_assignment_text") or question.get("assignment_text") or ""
        for n in range(1, 6):
            assert f"第{n}页" in assignment
        summary = metadata.get("image_read_summary") or {}
        assert summary.get("ocr_attempted") == 5
        assert summary.get("ocr_ok") == 5

    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_pdf_scanned_auto_ocr_without_enable_flag(self, mock_tesseract_pages):
        pdf = _ensure_scanned_pdf_fixture()
        question, metadata, full_text, warnings = build_question_from_document(
            pdf,
            "scanned_5page.pdf",
            enable_image_ocr=False,
        )
        assert question["image_assets"][0].get("ocr_status") == "ok"
        assert metadata.get("assignment_from_images") is True

    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_pdf_scanned_action_when_ocr_unavailable(self, monkeypatch):
        monkeypatch.setattr(image_read, "OCR_OK", False)
        pdf = _ensure_scanned_pdf_fixture()
        question, metadata, full_text, warnings = build_question_from_document(
            pdf,
            "scanned_5page.pdf",
            enable_image_ocr=False,
        )
        codes = [w.get("code") for w in warnings]
        assert "pdf_scanned" in codes
        assert "ocr_unavailable" in codes
        scanned_warn = next(w for w in warnings if w.get("code") == "pdf_scanned")
        assert scanned_warn.get("action") == "enable_ocr_reparse"
        assert len(question.get("image_assets") or []) == 5

    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_planner_input_includes_scanned_ocr(self, mock_tesseract_pages):
        pdf = _ensure_scanned_pdf_fixture()
        from agent.parse_documents import parse_single_file

        bundle = parse_single_file(
            pdf.read_bytes(),
            "scanned_5page.pdf",
            role="fill_target",
            enable_image_ocr=False,
        )
        assert "第1页" in bundle.get("planner_input_text", "")
        assert "第5页" in bundle.get("planner_input_text", "")
        assert bundle.get("assignment_from_images") is True

    @pytest.mark.skipif(not PDF_OK, reason="pymupdf not installed")
    def test_no_pdf_scanned_action_after_successful_ocr(self, mock_tesseract_pages):
        pdf = _ensure_scanned_pdf_fixture()
        question, metadata, full_text, warnings = build_question_from_document(
            pdf,
            "scanned_5page.pdf",
            enable_image_ocr=False,
        )
        scanned_warn = next(w for w in warnings if w.get("code") == "pdf_scanned")
        assert scanned_warn.get("action") != "enable_ocr_reparse"


# ── IM4 user upload assignment images ─────────────────────────────────


def _ensure_i4_png_fixtures() -> list[Path]:
    paths = [FIXTURES / f"assignment_page{i}.png" for i in range(1, 5)]
    if all(p.exists() for p in paths):
        return paths
    import importlib.util

    spec = importlib.util.spec_from_file_location(
        "gen_fixtures", FIXTURES / "gen_fixtures.py"
    )
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod.make_i4_assignment_pngs(FIXTURES)


def _png_items(paths: list[Path], *, include_flags: list[bool] | None = None) -> list[dict]:
    items = []
    for i, path in enumerate(paths):
        item = {
            "id": f"ui_{i}",
            "file_name": path.name,
            "file_data": base64.b64encode(path.read_bytes()).decode(),
            "order": i,
        }
        if include_flags is not None:
            item["include_in_ocr"] = include_flags[i]
        items.append(item)
    return items


def _mock_ocr_i4_page(asset):
    label = (asset.get("nearby_text") or asset.get("file_name") or "")
    texts = {
        "1": "图1：实验目的。掌握欧姆定律验证方法。",
        "2": "图2：实验器材。电源、电阻箱、电压表。",
        "3": "图3：实验步骤。按电路图连接并测量。",
        "4": "图4：数据记录。填写下表并计算误差。",
    }
    page_key = "1"
    for key in ("4", "3", "2", "1"):
        if f"page{key}" in label:
            page_key = key
            break
    return {
        "ocr_text": texts[page_key],
        "ocr_confidence": 0.9,
        "ocr_status": "ok",
        "ocr_engine": "tesseract",
        "ocr_lang": "chi_sim+eng",
        "ocr_error": "",
    }


@pytest.fixture
def mock_tesseract_i4(monkeypatch):
    monkeypatch.setattr(image_read, "OCR_OK", True)

    def _fake(asset, **kwargs):
        return _mock_ocr_i4_page(asset)

    monkeypatch.setattr(image_read, "ocr_image_asset", _fake)


class TestUserUploadAssets:
    def test_build_user_upload_ordered_deduped(self):
        png = _make_test_png(400, 300, "red")
        blob = png.read_bytes()
        b64 = base64.b64encode(blob).decode()
        items = [
            {"file_name": "a.png", "file_data": b64, "order": 1},
            {"file_name": "b.png", "file_data": b64, "order": 0},
            {"file_name": "c.png", "file_data": b64, "order": 2},
        ]
        result = build_user_upload_assets(items)
        assets = result["image_assets"]
        assert len(assets) == 1
        assert assets[0]["source"] == "user_upload"
        assert assets[0]["role_guess"] == "assignment"
        assert assets[0]["include_in_ocr"] is True

    def test_include_in_ocr_flag(self):
        png = _make_test_png(400, 300, "blue")
        b64 = base64.b64encode(png.read_bytes()).decode()
        result = build_user_upload_assets([
            {"file_name": "skip.png", "file_data": b64, "include_in_ocr": False},
        ])
        assert result["image_assets"][0]["include_in_ocr"] is False


class TestIm4UserUploadParse:
    def test_i4_four_images_merge_and_planner(self, mock_tesseract_i4):
        paths = _ensure_i4_png_fixtures()
        from agent.parse_documents import parse_assignment_images_only

        parsed = parse_assignment_images_only(
            _png_items(paths),
            enable_image_ocr=True,
        )
        assignment = parsed.get("assignment_text") or ""
        assert "图1" in assignment
        assert "图4" in assignment
        assert parsed.get("assignment_from_images") is True
        planner = parsed.get("planner_input_text") or ""
        assert "图1" in planner
        assert "图3" in planner
        assets = (parsed.get("metadata") or {}).get("image_assets") or []
        assert len(assets) == 4
        assert all(a.get("source") == "user_upload" for a in assets)

    def test_i4_reorder_changes_merge_order(self, mock_tesseract_i4):
        paths = _ensure_i4_png_fixtures()
        reordered = [paths[3], paths[1], paths[0], paths[2]]
        result = process_user_upload_images(
            _png_items(reordered),
            enable_image_ocr=True,
        )
        assignment = result.get("assignment_text") or ""
        pos1 = assignment.find("图4")
        pos2 = assignment.find("图2")
        pos3 = assignment.find("图1")
        assert pos1 >= 0 and pos2 > pos1 and pos3 > pos2

    def test_i4_include_ocr_false_skips_image(self, mock_tesseract_i4):
        paths = _ensure_i4_png_fixtures()[:2]
        result = process_user_upload_images(
            _png_items(paths, include_flags=[True, False]),
            enable_image_ocr=True,
        )
        assignment = result.get("assignment_text") or ""
        assert "图1" in assignment
        assert "图2" not in assignment
        assert result["image_read_summary"]["ocr_attempted"] == 1

    def test_i4_merge_with_document_assignment(self, mock_tesseract_i4):
        paths = _ensure_i4_png_fixtures()[:2]
        from agent.parse_documents import _apply_user_upload_assignment_images

        base = {
            "assignment_text": "粘贴的题目要求",
            "layout": "fill_only",
            "metadata": {"image_assets": []},
            "warnings": [],
            "documents": [],
            "_bundles": [],
        }
        merged = _apply_user_upload_assignment_images(
            base,
            _png_items(paths),
            enable_image_ocr=True,
        )
        text = merged["assignment_text"]
        assert "粘贴的题目要求" in text
        assert "图1" in text
        assert "图2" in text
        assert "图1" in merged["planner_input_text"]

    def test_i4_auto_ocr_without_enable_flag(self, mock_tesseract_i4):
        paths = _ensure_i4_png_fixtures()[:1]
        result = process_user_upload_images(
            _png_items(paths),
            enable_image_ocr=False,
        )
        assert result.get("assignment_from_images") is True
        assert "图1" in (result.get("assignment_text") or "")


# ── IM5 Vision / hybrid tests ─────────────────────────────────────────


VISION_LLM_SETTINGS = {
    "api_key": "test-key",
    "provider": "openai",
    "model": "gpt-4o",
}


def _ensure_vision_blank_fixture() -> Path:
    path = FIXTURES / "vision_blank_page.png"
    if not path.exists():
        import importlib.util

        spec = importlib.util.spec_from_file_location(
            "gen_fixtures", FIXTURES / "gen_fixtures.py"
        )
        mod = importlib.util.module_from_spec(spec)
        assert spec.loader is not None
        spec.loader.exec_module(mod)
        mod.make_vision_blank_png(path)
    return path


def _ensure_vision_page_fixtures(count: int = 6) -> list[Path]:
    paths = [FIXTURES / f"vision_page{i}.png" for i in range(1, count + 1)]
    if all(p.exists() for p in paths):
        return paths
    import importlib.util

    spec = importlib.util.spec_from_file_location(
        "gen_fixtures", FIXTURES / "gen_fixtures.py"
    )
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(mod)
    return mod.make_i5_vision_pages(FIXTURES, count=count)


def _mock_empty_ocr(asset, **kwargs):
    return {
        "ocr_text": "",
        "ocr_confidence": 0.0,
        "ocr_status": "empty",
        "ocr_engine": "tesseract",
        "ocr_lang": "chi_sim+eng",
        "ocr_error": "",
    }


@pytest.fixture
def mock_empty_ocr(monkeypatch):
    monkeypatch.setattr(image_read, "OCR_OK", True)
    monkeypatch.setattr(image_read, "ocr_image_asset", _mock_empty_ocr)


@pytest.fixture
def mock_chat_vision(monkeypatch):
    def _fake(settings, **kwargs):
        return {
            "content": "实验目的：通过 Vision 识别到的题目文字。",
            "reasoning_content": "",
            "phase": "vision_read",
            "finish_reason": "stop",
            "usage": {},
        }

    import llm_client

    monkeypatch.setattr(llm_client, "supports_vision", lambda s: True)
    monkeypatch.setattr(llm_client, "chat_vision", _fake)


@pytest.fixture
def mock_chat_vision_by_page(monkeypatch):
    def _fake(settings, **kwargs):
        b64 = kwargs.get("image_b64") or ""
        idx = len(b64) % 6 + 1
        return {
            "content": f"Vision页{idx}：实验要求段落 {idx}。",
            "reasoning_content": "",
            "phase": "vision_read",
            "finish_reason": "stop",
            "usage": {},
        }

    import llm_client

    monkeypatch.setattr(llm_client, "supports_vision", lambda s: True)
    monkeypatch.setattr(llm_client, "chat_vision", _fake)


class TestLlmClientVision:
    def test_supports_vision_openai(self):
        from llm_client import supports_vision

        assert supports_vision({"provider": "openai", "model": "gpt-4o"}) is True
        assert supports_vision({"provider": "deepseek", "model": "deepseek-chat"}) is False

    def test_build_vision_user_content(self):
        from llm_client import build_vision_user_content

        parts = build_vision_user_content("read this", image_b64="abc123", mime="image/png")
        assert parts[0]["type"] == "text"
        assert parts[1]["type"] == "image_url"
        assert "abc123" in parts[1]["image_url"]["url"]


class TestIm5HybridVision:
    def test_i5_hybrid_empty_ocr_uses_vision(self, mock_empty_ocr, mock_chat_vision):
        img = _ensure_vision_blank_fixture()
        path = _make_test_docx(paragraphs_before=["短"], images=[img])
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "hybrid.docx",
            enable_image_ocr=True,
            image_reading_mode="hybrid",
            vision_max_pages=5,
            llm_settings=VISION_LLM_SETTINGS,
        )
        assignment = metadata.get("document_assignment_text") or ""
        assert "Vision 识别" in assignment
        assert metadata.get("image_reading_mode") == "hybrid"
        section = (metadata.get("image_sections") or [])[0]
        assert section.get("source") == "vision"
        assert question["image_assets"][0].get("vision_status") == "ok"

    def test_hybrid_keeps_ocr_when_ok(self, mock_tesseract, mock_chat_vision):
        img = FIXTURES / "ocr_simple_zh.png"
        if not img.exists():
            pytest.skip("fixture missing")
        path = _make_test_docx(paragraphs_before=["题"], images=[img])
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "hybrid_ok.docx",
            enable_image_ocr=True,
            image_reading_mode="hybrid",
            llm_settings=VISION_LLM_SETTINGS,
        )
        assignment = metadata.get("document_assignment_text") or ""
        assert "欧姆定律" in assignment
        assert "--- 图 1（OCR）---" in assignment
        summary = metadata.get("image_read_summary") or {}
        assert summary.get("vision_attempted", 0) == 0

    def test_ocr_only_never_calls_vision(self, mock_empty_ocr, monkeypatch):
        called = {"n": 0}

        def _fake(*args, **kwargs):
            called["n"] += 1
            return {"content": "should not run"}

        import llm_client

        monkeypatch.setattr(llm_client, "chat_vision", _fake)
        img = _ensure_vision_blank_fixture()
        path = _make_test_docx(paragraphs_before=["短"], images=[img])
        build_question_from_document(
            path,
            "ocr_only.docx",
            enable_image_ocr=True,
            image_reading_mode="ocr_only",
            llm_settings=VISION_LLM_SETTINGS,
        )
        assert called["n"] == 0

    def test_vision_unavailable_warns(self, mock_empty_ocr, monkeypatch):
        import llm_client

        monkeypatch.setattr(llm_client, "supports_vision", lambda s: False)
        img = _ensure_vision_blank_fixture()
        path = _make_test_docx(paragraphs_before=["短"], images=[img])
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "no_vision_model.docx",
            enable_image_ocr=True,
            image_reading_mode="hybrid",
            llm_settings={"provider": "deepseek", "model": "deepseek-chat", "api_key": "k"},
        )
        codes = [w.get("code") for w in warnings]
        assert "vision_unavailable" in codes
        assert not metadata.get("assignment_from_images")

    def test_standard_docx_regression_with_hybrid_settings(
        self, mock_tesseract, mock_chat_vision
    ):
        path = _make_test_docx(
            paragraphs_before=[
                "实验目的：验证欧姆定律",
                "实验原理：R=U/I",
                "三、实验步骤",
                "1. 连接电路",
                "2. 测量数据",
                "四、实验结果",
                "数据如下...",
                "五、实验总结",
                "本次实验成功验证了欧姆定律",
            ],
        )
        question, metadata, full_text, warnings = build_question_from_document(
            path,
            "standard_hybrid.docx",
            image_reading_mode="hybrid",
            llm_settings=VISION_LLM_SETTINGS,
        )
        assert question.get("image_assets") == []
        assert not metadata.get("assignment_from_images")
        summary = metadata.get("image_read_summary")
        assert summary is None or summary.get("vision_attempted", 0) == 0


class TestIm5VisionLimit:
    def test_i6_vision_limit_exceeded_warns(
        self, mock_empty_ocr, mock_chat_vision_by_page
    ):
        paths = _ensure_vision_page_fixtures(6)
        from document.user_upload_images import process_user_upload_images

        result = process_user_upload_images(
            _png_items(paths),
            enable_image_ocr=True,
            image_reading_mode="vision",
            vision_max_pages=3,
            llm_settings=VISION_LLM_SETTINGS,
        )
        summary = result.get("image_read_summary") or {}
        assert summary.get("vision_attempted") == 3
        assert summary.get("vision_limit_exceeded") == 3
        codes = [w.get("code") for w in result.get("warnings") or []]
        assert "vision_limit_exceeded" in codes
        skipped = [
            a for a in result.get("image_assets") or []
            if a.get("vision_error") == "vision_max_pages_exceeded"
        ]
        assert len(skipped) == 3

    def test_vision_dedup_by_sha256(self, mock_empty_ocr, monkeypatch):
        calls = {"n": 0}

        def _fake(settings, **kwargs):
            calls["n"] += 1
            return {
                "content": "Vision dedup text",
                "reasoning_content": "",
                "phase": "vision_read",
                "finish_reason": "stop",
                "usage": {},
            }

        import llm_client

        monkeypatch.setattr(llm_client, "supports_vision", lambda s: True)
        monkeypatch.setattr(llm_client, "chat_vision", _fake)

        png = _make_test_png(400, 300, "red")
        blob = png.read_bytes()
        b64 = base64.b64encode(blob).decode()
        items = [
            {"file_name": "a.png", "file_data": b64, "order": 0},
            {"file_name": "b.png", "file_data": b64, "order": 1},
        ]
        from document.user_upload_images import process_user_upload_images

        result = process_user_upload_images(
            items,
            enable_image_ocr=True,
            image_reading_mode="vision",
            llm_settings=VISION_LLM_SETTINGS,
        )
        assert calls["n"] == 1
        assert len(result.get("image_assets") or []) == 1


# ── UI polish: multi-question warn + plan assignment override ─────────


class TestMultiQuestionInImage:
    def test_detect_numbered_sections(self):
        text = "一、实验目的\n做某某实验\n二、实验步骤\n按手册操作\n三、思考题\n简述原理"
        assert detect_multi_question_in_image(text) is True

    def test_short_text_not_flagged(self):
        assert detect_multi_question_in_image("一、短题") is False

    def test_warning_emitted_no_split(self):
        sections = [{
            "image_id": "img_001",
            "text": "一、第一题\n内容足够长的一、第一题内容\n二、第二题\n更多内容让整体超过四十个字符",
            "source": "ocr",
        }]
        warns = multi_question_image_warnings(sections)
        assert len(warns) == 1
        assert warns[0]["code"] == "multi_question_in_image"
        assert "不会自动拆分" in warns[0]["message"]


class TestAssignmentTextOverride:
    def test_apply_override_rebuilds_planner_input(self):
        from agent.parse_documents import apply_assignment_text_override

        bundle = {
            "assignment_text": "旧题干",
            "report_text": "填表正文",
            "fill_body_text": "填表正文",
            "layout": "combined",
            "planner_input_text": "旧 planner",
        }
        apply_assignment_text_override(bundle, "用户校对后的题干")
        assert bundle["assignment_text"] == "用户校对后的题干"
        assert "用户校对后的题干" in bundle["planner_input_text"]
        assert "填表正文" in bundle["planner_input_text"]
