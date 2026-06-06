"""Phase B — UML diagram placement in fill_report."""

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src" / "python"))

import pytest
from docx import Document

from modules.fill_report import (
    _IMAGE_FORBIDDEN_SEMANTICS,
    _resolve_diagram_placements,
    _sanitize_diagram_target,
    fill_lab,
)

FIXTURES = ROOT / "tests" / "fixtures"
TINY_PNG = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _section_map(steps_heading="三、实验内容", result_heading="四、实验结果", summary_heading="五、实验总结"):
    return {
        "steps": {"type": "paragraph", "heading": steps_heading, "para_index": 2},
        "result": {"type": "paragraph", "heading": result_heading, "para_index": 5},
        "summary": {"type": "paragraph", "heading": summary_heading, "para_index": 8},
    }


class TestDiagramTargetHelpers:
    def test_objective_is_forbidden(self):
        assert "objective" in _IMAGE_FORBIDDEN_SEMANTICS

    def test_sanitize_downgrades_objective(self):
        sm = _section_map()
        target = _sanitize_diagram_target("objective", sm, {"steps", "result", "summary"})
        assert target == "steps"

    def test_default_prefers_content_semantic(self):
        sm = _section_map(steps_heading="三、实验内容")
        target = _sanitize_diagram_target("content", sm, {"steps"})
        assert target == "steps"

    def test_diagrams_target_per_index(self):
        sm = _section_map()
        fill_hints = {
            "diagrams_target": [
                {"image_index": 0, "target_semantic": "result"},
                {"image_index": 1, "target_semantic": "steps"},
            ],
            "uml_default_target": "content",
        }
        placements = _resolve_diagram_placements(
            ["img0", "img1"],
            fill_hints,
            sm,
            {"steps", "result", "summary"},
        )
        by_target = dict(placements)
        assert by_target["result"] == ["img0"]
        assert by_target["steps"] == ["img1"]


class TestUmlNotInObjective:
    def test_paragraph_report_skips_objective_for_uml(self):
        doc = Document()
        doc.add_paragraph("一、实验目的")
        doc.add_paragraph("（占位，将被替换）")
        doc.add_paragraph("三、实验内容")
        doc.add_paragraph("（占位）")
        doc.add_paragraph("四、实验结果")
        doc.add_paragraph("（占位）")
        doc.add_paragraph("五、实验总结")
        doc.add_paragraph("（占位）")

        ans = {
            "type": "lab_report",
            "parsed": {
                "steps_analysis": "需求分析：绘制用例图。",
                "result_description": "用例图已绘制。",
                "summary": "掌握用例建模。",
            },
            "uml_images_b64": [TINY_PNG],
            "include_code": False,
        }
        metadata = {
            "sections_detected": [
                {"index": 0, "heading": "一、实验目的", "semantic": "objective"},
                {"index": 2, "heading": "三、实验内容", "semantic": "steps"},
                {"index": 4, "heading": "四、实验结果", "semantic": "result"},
                {"index": 6, "heading": "五、实验总结", "semantic": "summary"},
            ],
            "section_map": _section_map(),
            "fill_hints": {
                "diagrams_target": [{"image_index": 0, "target_semantic": "objective"}],
            },
        }
        fill_lab(doc, ans, metadata=metadata)

        objective_idx = 0
        content_idx = 2
        objective_paras = 0
        content_paras = 0
        for i, p in enumerate(doc.paragraphs):
            has_img = any(
                run._element.xpath(".//a:blip")
                for run in p.runs
            ) if p.runs else False
            if not has_img:
                continue
            if i <= objective_idx + 2:
                objective_paras += 1
            if i >= content_idx:
                content_paras += 1
        assert objective_paras == 0, "UML must not appear in 实验目的"
        assert content_paras >= 1, "UML should fall back to 实验内容/steps"

    def test_training_table_uml_not_in_objective_cell(self):
        fixture = FIXTURES / "lab_report_table.docx"
        if not fixture.exists():
            pytest.skip("run python tests/generate_fixtures.py first")

        from docx import Document as D
        doc = D(str(fixture))
        ans = {
            "type": "lab_report",
            "parsed": {
                "steps_analysis": "ER 图设计…",
                "result_description": "",
                "summary": "完成数据库设计",
            },
            "uml_images_b64": [TINY_PNG],
            "include_code": False,
        }
        metadata = {
            "report_layout": "training_table",
            "table_map": [
                {"table": 0, "row": 5, "col": 0, "label": "实验目的"},
                {"table": 0, "row": 6, "col": 0, "label": "实验内容"},
            ],
            "fill_hints": {"uml_default_target": "content"},
        }
        fill_lab(doc, ans, metadata=metadata)

        objective_cell = doc.tables[0].rows[5].cells[1]
        body_cell = doc.tables[0].rows[6].cells[1]
        obj_img_paras = sum(
            1 for p in objective_cell.paragraphs
            if any(run._element.xpath(".//a:blip") for run in p.runs)
        )
        body_img_paras = sum(
            1 for p in body_cell.paragraphs
            if any(run._element.xpath(".//a:blip") for run in p.runs)
        )
        assert obj_img_paras == 0
        assert body_img_paras >= 1
