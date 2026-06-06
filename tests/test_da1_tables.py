"""DA1 — table text extraction and training_table layout detection."""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src" / "python"))

import pytest
from modules.parse_report import (
    _cell_texts,
    _detect_table_layout,
    _render_table_as_text,
    extract_docx,
)

FIXTURES = ROOT / "tests" / "fixtures"


def _load_doc(path):
    from docx import Document
    return Document(str(path))


class TestCellTexts:
    def test_dedup_merged_cells(self):
        """Mock row cells where merged cells produce duplicate text."""
        class FakeCell:
            def __init__(self, text):
                self.text = text
        from unittest.mock import MagicMock
        row = MagicMock()
        row.cells = [FakeCell("A"), FakeCell("A"), FakeCell("B"), FakeCell("C"), FakeCell("C")]
        assert _cell_texts(row) == ["A", "B", "C"]

    def test_empty_cells_skipped(self):
        from unittest.mock import MagicMock
        class FakeCell:
            def __init__(self, text):
                self.text = text
        row = MagicMock()
        row.cells = [FakeCell(""), FakeCell("X"), FakeCell(""), FakeCell("")]
        assert _cell_texts(row) == ["X"]


class TestRenderTableAsText:
    def test_two_col_key_value(self, tmp_path):
        """2-column table with known labels uses bracket format."""
        from docx import Document
        d = Document()
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "实训项目"
        t.rows[0].cells[1].text = "JSP 文件上传"
        t.rows[1].cells[0].text = "学号"
        t.rows[1].cells[1].text = "2021001"
        out = _render_table_as_text(t, 0)
        assert "【实训项目】" in out
        assert "JSP 文件上传" in out
        assert "【学号】" in out

    def test_multi_col_pipe_join(self):
        from docx import Document
        d = Document()
        t = d.add_table(rows=2, cols=3)
        t.rows[0].cells[0].text = "A"
        t.rows[0].cells[1].text = "B"
        t.rows[0].cells[2].text = "C"
        t.rows[1].cells[0].text = "1"
        t.rows[1].cells[1].text = "2"
        t.rows[1].cells[2].text = "3"
        out = _render_table_as_text(t, 0)
        assert "A | B | C" in out
        assert "1 | 2 | 3" in out

    def test_empty_table(self):
        from docx import Document
        d = Document()
        t = d.add_table(rows=0, cols=0)
        assert _render_table_as_text(t, 0) == ""


class TestDetectTableLayout:
    def test_training_table_detected(self):
        doc = _load_doc(FIXTURES / "training_table.docx")
        result = _detect_table_layout(doc)
        assert result["report_layout"] == "training_table"
        assert isinstance(result["table_map"], list)
        assert len(result["table_map"]) >= 1
        markers = {e["label"] for e in result["table_map"]}
        assert "实训任务" in markers or "实训步骤及内容" in markers

    def test_lab_report_table_detected(self):
        fixture = FIXTURES / "lab_report_table.docx"
        if not fixture.exists():
            pytest.skip("run python tests/generate_fixtures.py first")
        doc = _load_doc(fixture)
        result = _detect_table_layout(doc)
        assert result["report_layout"] == "training_table"
        markers = {e["label"] for e in result["table_map"]}
        assert "实验内容" in markers
        assert "实验目的" in markers

    def test_standard_docx_not_training(self):
        doc = _load_doc(FIXTURES / "programming_lab.docx")
        result = _detect_table_layout(doc)
        assert result == {}


class TestExtractDocx:
    def test_training_table_text_visible(self):
        text, meta = extract_docx(FIXTURES / "training_table.docx")
        assert "实训项目" in text
        assert "JSP" in text
        assert "FileUpload" in text or "文件上传" in text
        assert "Tomcat" in text
        assert meta.get("report_layout") == "training_table"

    def test_standard_docx_regression(self):
        """Standard paragraph-based report must still parse correctly."""
        text, meta = extract_docx(FIXTURES / "programming_lab.docx")
        assert "实验目的" in text
        assert "顺序结构程序设计" in text or "Java" in text
        assert "三、" in text or "实验步骤" in text
        assert "四、" in text or "实验结果" in text
        assert meta.get("report_layout") is None  # standard layout: no override

    def test_metadata_unchanged_shape(self):
        """metadata still has expected fields for standard reports."""
        _, meta = extract_docx(FIXTURES / "programming_lab.docx")
        assert "course" in meta or "experiment_title" in meta
        assert meta.get("report_layout") is None

    def test_cover_table_metadata(self):
        """Cover table metadata extraction still works."""
        _, meta = extract_docx(FIXTURES / "training_table.docx")
        assert meta.get("course") == "Java Web 开发"
        assert "第十周实训" in meta.get("experiment_title", "")
        assert meta.get("major") == "计算机科学与技术"
