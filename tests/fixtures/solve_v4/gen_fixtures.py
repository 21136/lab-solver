"""Generate solve_v4 golden docx fixtures (run: python tests/fixtures/solve_v4/gen_fixtures.py)."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document

DIR = Path(__file__).resolve().parent


def _cover(doc: Document, course: str, title: str) -> None:
    table = doc.add_table(rows=3, cols=2)
    for i, (k, v) in enumerate(
        [("课程名称", course), ("实验名称", title), ("专业", "计算机科学与技术")]
    ):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v


def _save_docx(name: str, builder) -> None:
    path = DIR / name
    builder().save(path)
    print(f"wrote {path}")


def _code_response(code_files: list, main_file: str, language: str) -> str:
    return json.dumps(
        {"code_files": code_files, "main_file": main_file, "language": language},
        ensure_ascii=False,
    )


def _report_response(**fields) -> str:
    base = {
        "steps_analysis": "实验步骤说明。",
        "result_description": "运行结果符合预期。",
        "summary": "实验总结。",
    }
    base.update(fields)
    return json.dumps(base, ensure_ascii=False)


FIXTURES = [
    {
        "id": "01_fifo_lru",
        "docx": "01_fifo_lru.docx",
        "preferred_lang": "java",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "Main.java",
                        "code": (
                            "public class Main {\n"
                            "  public static void main(String[] args) {\n"
                            '    System.out.println("FIFO hits: 3");\n'
                            '    System.out.println("LRU hits: 4");\n'
                            "  }\n"
                            "}\n"
                        ),
                    }
                ],
                "Main.java",
                "java",
            ),
            _report_response(
                result_description="FIFO 命中率 3，LRU 命中率 4。",
                expected_output="FIFO hits: 3\nLRU hits: 4\n",
            ),
        ],
        "assertions": {"stdout_contains": "FIFO hits"},
    },
    {
        "id": "02_factory_singleton",
        "docx": "02_factory_singleton.docx",
        "preferred_lang": "java",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "Main.java",
                        "code": (
                            "interface Product { String name(); }\n"
                            "class ConcreteProduct implements Product {\n"
                            '  public String name() { return "product"; }\n'
                            "}\n"
                            "class Factory {\n"
                            "  static Product create() { return new ConcreteProduct(); }\n"
                            "}\n"
                            "public class Main {\n"
                            "  public static void main(String[] args) {\n"
                            "    Product p = Factory.create();\n"
                            "    System.out.println(p.name());\n"
                            "  }\n"
                            "}\n"
                        ),
                    }
                ],
                "Main.java",
                "java",
            ),
            _report_response(result_description="工厂创建产品并输出 product。"),
        ],
        "assertions": {"stdout_contains": "product"},
    },
    {
        "id": "03_thread_join",
        "docx": "03_thread_join.docx",
        "preferred_lang": "java",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "Main.java",
                        "code": (
                            "public class Main {\n"
                            "  public static void main(String[] args) throws Exception {\n"
                            '    Thread t = new Thread(() -> System.out.println("worker"));\n'
                            "    t.start();\n"
                            "    t.join();\n"
                            '    System.out.println("done");\n'
                            "  }\n"
                            "}\n"
                        ),
                    }
                ],
                "Main.java",
                "java",
            ),
            _report_response(result_description="主线程等待子线程结束后输出 done。"),
        ],
        "assertions": {"stdout_contains": "done"},
    },
    {
        "id": "04_sort_c",
        "docx": "04_sort_c.docx",
        "preferred_lang": "c",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "main.c",
                        "code": (
                            "#include <stdio.h>\n"
                            "int main() {\n"
                            "  int a[] = {3, 1, 2};\n"
                            "  for (int i = 0; i < 3; i++)\n"
                            "    for (int j = i + 1; j < 3; j++)\n"
                            "      if (a[i] > a[j]) { int t = a[i]; a[i] = a[j]; a[j] = t; }\n"
                            '  printf("sorted\\n");\n'
                            "  return 0;\n"
                            "}\n"
                        ),
                    }
                ],
                "main.c",
                "c",
            ),
            _report_response(result_description="冒泡排序后输出 sorted。"),
        ],
        "assertions": {"stdout_contains": "sorted"},
    },
    {
        "id": "05_file_io_python",
        "docx": "05_file_io_python.docx",
        "preferred_lang": "python",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "main.py",
                        "code": (
                            'with open("golden_io.txt", "w", encoding="utf-8") as f:\n'
                            '    f.write("lab-data")\n'
                            'with open("golden_io.txt", encoding="utf-8") as f:\n'
                            "    print(f.read())\n"
                        ),
                    }
                ],
                "main.py",
                "python",
            ),
            _report_response(result_description="文件读写输出 lab-data。"),
        ],
        "assertions": {"stdout_contains": "lab-data"},
    },
    {
        "id": "06_theory_only",
        "docx": "06_theory_only.docx",
        "preferred_lang": "python",
        "expected_code_status": "skipped",
        "mock_llm": [
            _report_response(
                steps_analysis="分析 TCP 三次握手各阶段报文特征。",
                result_description="SYN、SYN-ACK、ACK 依次交互建立连接。",
                summary="理解了握手时序与标志位含义。",
            ),
        ],
        "assertions": {},
    },
    {
        "id": "07_web_simulation",
        "docx": "07_web_simulation.docx",
        "preferred_lang": "java",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "Main.java",
                        "code": (
                            "public class Main {\n"
                            "  public static void main(String[] args) {\n"
                            '    System.out.println("模拟 Web 请求处理 OK");\n'
                            "  }\n"
                            "}\n"
                        ),
                    }
                ],
                "Main.java",
                "java",
            ),
            _report_response(result_description="内存模拟 Web 请求，无 Servlet 容器。"),
        ],
        "assertions": {"stdout_contains": "Web"},
    },
    {
        "id": "08_multifile_java",
        "docx": "08_multifile_java.docx",
        "preferred_lang": "java",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "Helper.java",
                        "code": (
                            "class Helper {\n"
                            '  static String greet() { return "hello"; }\n'
                            "}\n"
                        ),
                    },
                    {
                        "name": "Main.java",
                        "code": (
                            "public class Main {\n"
                            "  public static void main(String[] args) {\n"
                            "    System.out.println(Helper.greet());\n"
                            "  }\n"
                            "}\n"
                        ),
                    },
                ],
                "Main.java",
                "java",
            ),
            _report_response(result_description="多文件包结构编译运行输出 hello。"),
        ],
        "assertions": {"stdout_contains": "hello"},
    },
    {
        "id": "09_linked_list_cpp",
        "docx": "09_linked_list_cpp.docx",
        "preferred_lang": "cpp",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "main.cpp",
                        "code": (
                            "#include <iostream>\n"
                            "struct Node { int v; Node* next; };\n"
                            "int main() {\n"
                            "  Node n{1, nullptr};\n"
                            '  std::cout << "list ok" << std::endl;\n'
                            "  return 0;\n"
                            "}\n"
                        ),
                    }
                ],
                "main.cpp",
                "cpp",
            ),
            _report_response(result_description="链表节点创建成功，输出 list ok。"),
        ],
        "assertions": {"stdout_contains": "list ok"},
    },
    {
        "id": "10_no_emoji",
        "docx": "10_no_emoji.docx",
        "preferred_lang": "java",
        "expected_code_status": "verified",
        "mock_llm": [
            _code_response(
                [
                    {
                        "name": "Main.java",
                        "code": (
                            "public class Main {\n"
                            "  public static void main(String[] args) {\n"
                            '    System.out.println("run ok");\n'
                            "  }\n"
                            "}\n"
                        ),
                    }
                ],
                "Main.java",
                "java",
            ),
            _report_response(
                result_description="程序正常结束，输出 run ok。",
                summary="未使用 emoji 符号。",
            ),
        ],
        "assertions": {
            "stdout_contains": "run ok",
            "no_emoji_in_result": True,
        },
    },
]


def _doc_builders() -> dict[str, callable]:
    return {
        "01_fifo_lru.docx": lambda: _fifo_doc(),
        "02_factory_singleton.docx": lambda: _factory_doc(),
        "03_thread_join.docx": lambda: _thread_doc(),
        "04_sort_c.docx": lambda: _sort_doc(),
        "05_file_io_python.docx": lambda: _file_io_doc(),
        "06_theory_only.docx": lambda: _theory_doc(),
        "07_web_simulation.docx": lambda: _web_doc(),
        "08_multifile_java.docx": lambda: _multifile_doc(),
        "09_linked_list_cpp.docx": lambda: _cpp_doc(),
        "10_no_emoji.docx": lambda: _emoji_doc(),
    }


def _fifo_doc():
    doc = Document()
    _cover(doc, "操作系统", "实验一 页面置换 FIFO/LRU")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("实现 FIFO 与 LRU 页面置换算法，统计并打印命中次数。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("使用 Java 编写程序，模拟页面访问序列并对比两种算法命中率。")
    doc.add_paragraph("三、实验步骤")
    doc.add_paragraph("1. 实现 FIFO 队列与 LRU 访问时间更新；2. 编译运行并记录输出。")
    return doc


def _factory_doc():
    doc = Document()
    _cover(doc, "设计模式", "实验二 简单工厂与单例")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握简单工厂模式创建产品对象。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("使用 Java 实现 Factory 与 Product 接口，main 方法打印产品名称。")
    return doc


def _thread_doc():
    doc = Document()
    _cover(doc, "Java 程序设计", "实验三 多线程 join")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握线程创建与 join 等待。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("编写 Java 程序：启动子线程打印 worker，主线程 join 后打印 done。")
    return doc


def _sort_doc():
    doc = Document()
    _cover(doc, "C 语言程序设计", "实验四 排序算法")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("实现冒泡排序算法。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("使用 C 语言对整型数组排序，完成后打印 sorted。")
    return doc


def _file_io_doc():
    doc = Document()
    _cover(doc, "Python 程序设计", "实验五 文件读写")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握 Python 文件读写。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("编写 Python 程序：写入 golden_io.txt 再读出并打印内容。")
    return doc


def _theory_doc():
    doc = Document()
    _cover(doc, "计算机网络", "实验六 协议分析")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("理解 TCP 三次握手过程，用文字描述各阶段报文特征。")
    doc.add_paragraph("二、实验原理")
    doc.add_paragraph("TCP 建立连接需要 SYN、SYN-ACK、ACK 三次交互。")
    doc.add_paragraph("三、实验步骤")
    doc.add_paragraph("使用 Wireshark 抓包并分析握手报文，撰写分析报告即可。")
    return doc


def _web_doc():
    doc = Document()
    _cover(doc, "Java Web 开发", "实验七 Web 请求模拟")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("理解 Web 请求处理流程（本实验为内存模拟，非真实 Servlet 部署）。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("使用 Java 编写控制台程序模拟 Web 请求处理并打印结果。")
    return doc


def _multifile_doc():
    doc = Document()
    _cover(doc, "Java 程序设计", "实验八 多文件包结构")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握多文件 Java 工程组织。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("创建 Helper.java 与 Main.java，Main 调用 Helper.greet() 并打印。")
    return doc


def _cpp_doc():
    doc = Document()
    _cover(doc, "数据结构", "实验九 链表")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握链表基本结构。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("使用 C++ 定义 Node 结构并打印 list ok。")
    return doc


def _emoji_doc():
    doc = Document()
    _cover(doc, "Java 程序设计", "实验十 控制台输出")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("编写可编译运行的 Java 程序，控制台输出不含 emoji。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("实现 main 方法打印 run ok，报告正文亦勿使用 emoji。")
    return doc


def main():
    builders = _doc_builders()
    for spec in FIXTURES:
        _save_docx(spec["docx"], builders[spec["docx"]])
    manifest = {"version": 1, "fixtures": FIXTURES}
    manifest_path = DIR / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"wrote {manifest_path}")


if __name__ == "__main__":
    main()
