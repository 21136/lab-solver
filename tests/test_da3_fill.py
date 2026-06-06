"""DA3 — fill_report adaptation: training_table fill, missing section merge, fill_hints."""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src" / "python"))

import pytest
from modules.fill_report import (
    _build_fill_hints,
    _build_section_map,
    _resolve_objective_text,
    _training_fill_targets,
    _replace_section,
    detect_sections,
    do_fill,
    extract_objective_from_assignment,
    fill_lab,
)

FIXTURES = ROOT / "tests" / "fixtures"


# ── helpers ──

def _load_doc(path):
    from docx import Document
    return Document(str(path))


def _make_fake_ans(steps="步骤内容", result="结果内容", summary="总结内容", code=None):
    parsed = {
        "steps_analysis": steps,
        "result_description": result,
        "expected_output": "",
        "summary": summary,
        "code": code or "",
        "diagrams": [],
    }
    return {"type": "lab_report", "parsed": parsed, "answer": steps, "include_code": False}


# ── _build_fill_hints ──


class TestBuildFillHints:
    def test_all_sections_present(self):
        sm = {
            "steps": {"type": "paragraph", "heading": "三、实验步骤", "para_index": 5},
            "result": {"type": "paragraph", "heading": "四、实验结果", "para_index": 10},
            "summary": {"type": "paragraph", "heading": "五、实验总结", "para_index": 15},
        }
        hints = _build_fill_hints(sm)
        assert hints.get("uml_default_target") == "steps"
        assert "merge_" not in "".join(hints.keys())
        assert "screenshots_target" not in hints

    def test_result_missing(self):
        sm = {
            "steps": {"type": "paragraph", "heading": "三、实验步骤", "para_index": 5},
            "result": None,
            "summary": {"type": "paragraph", "heading": "四、实验总结", "para_index": 10},
        }
        hints = _build_fill_hints(sm)
        assert hints["screenshots_target"] == "summary"
        assert hints["merge_result_into"] == "summary"

    def test_steps_missing(self):
        sm = {
            "steps": None,
            "result": {"type": "paragraph", "heading": "三、实验结果", "para_index": 5},
            "summary": {"type": "paragraph", "heading": "四、实验总结", "para_index": 10},
        }
        hints = _build_fill_hints(sm)
        assert hints["merge_steps_into"] == "result"

    def test_summary_missing(self):
        sm = {
            "steps": {"type": "paragraph", "heading": "三、实验步骤", "para_index": 5},
            "result": {"type": "paragraph", "heading": "四、实验结果", "para_index": 10},
            "summary": None,
        }
        hints = _build_fill_hints(sm)
        assert hints["merge_summary_into"] == "result"

    def test_only_steps(self):
        sm = {
            "steps": {"type": "paragraph", "heading": "三、实验内容", "para_index": 2},
            "result": None,
            "summary": None,
        }
        hints = _build_fill_hints(sm)
        assert hints["screenshots_target"] == "steps"


# ── objective extraction ──


DESIGN_PATTERN_ASSIGNMENT = """实验项目二： 创建型设计模式实验；
（1）实验目的与原理
①结合实例,熟练绘制创建型设计模式结构图。
②结合实例,熟练使用Java面向对象编程语言实现创建型设计模式。
③通过编程实践，理解每一种创建型设计模式的概念和内涵、结构、优缺点以及应用场景。
（2）实验内容与步骤
①使用简单工厂模式设计一个可以创建不同几何形状的绘图工具类。
"""


class TestObjectiveExtraction:
    def test_extract_from_assignment_section(self):
        obj = extract_objective_from_assignment(DESIGN_PATTERN_ASSIGNMENT)
        assert "熟练绘制创建型设计模式结构图" in obj
        assert "简单工厂模式" not in obj

    def test_resolve_prefers_assignment_over_steps(self):
        parsed = {"steps_analysis": "简单工厂模式实现…\n\n更多步骤"}
        meta = {"assignment_text": DESIGN_PATTERN_ASSIGNMENT}
        obj = _resolve_objective_text(parsed, metadata=meta)
        assert "熟练使用Java面向对象编程语言" in obj
        assert "简单工厂模式实现" not in obj


# ── _training_fill_targets ──


class TestTrainingFillTargets:
    def test_finds_steps_cell(self):
        entries = [
            {"table": 1, "row": 0, "col": 0, "label": "实训项目"},
            {"table": 1, "row": 4, "col": 0, "label": "实训步骤及内容"},
            {"table": 1, "row": 5, "col": 0, "label": "实训任务"},
        ]
        targets = _training_fill_targets(entries)
        assert len(targets) == 2  # "实训步骤及内容" + "实训任务"
        assert targets[0]["label"] == "实训步骤及内容"
        assert targets[0]["semantic"] == "steps"
        assert targets[1]["label"] == "实训任务"
        assert targets[1]["semantic"] == "steps"

    def test_finds_experiment_content_cell(self):
        entries = [
            {"table": 0, "row": 0, "col": 0, "label": "学号"},
            {"table": 0, "row": 5, "col": 0, "label": "实验目的"},
            {"table": 0, "row": 6, "col": 0, "label": "实验内容"},
        ]
        targets = _training_fill_targets(entries)
        assert len(targets) == 2
        assert targets[0]["semantic"] == "objective"
        assert targets[0]["label"] == "实验目的"
        assert targets[1]["semantic"] == "steps"
        assert targets[1]["label"] == "实验内容"

    def test_experiment_name_skips_cover_table(self):
        entries = [
            {"table": 0, "row": 1, "col": 0, "label": "实验名称"},
            {"table": 1, "row": 0, "col": 0, "label": "实验名称"},
        ]
        targets = _training_fill_targets(entries)
        assert len(targets) == 1
        assert targets[0]["table"] == 1
        assert targets[0]["semantic"] == "experiment_name"

    def test_empty_entries(self):
        assert _training_fill_targets([]) == []
        assert _training_fill_targets(None) == []

    def test_no_matching_label(self):
        entries = [
            {"table": 0, "row": 0, "col": 0, "label": "课程名称"},
            {"table": 0, "row": 1, "col": 0, "label": "学号"},
        ]
        assert _training_fill_targets(entries) == []


# ── _replace_section boundary detection ──


class TestReplaceSectionBoundary:
    def test_uses_section_map_boundary(self):
        """_replace_section should use next section from section_map, not number regex."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("三、实验步骤")
        doc.add_paragraph("old step content")
        doc.add_paragraph("more old content")
        doc.add_paragraph("四、实验总结")

        paras = list(doc.paragraphs)
        section_map = {
            "steps": {"type": "paragraph", "heading": "三、实验步骤", "para_index": 0},
            "result": None,
            "summary": {"type": "paragraph", "heading": "四、实验总结", "para_index": 3},
        }
        _replace_section(paras, 0, "new content", section_map=section_map)

        texts = [p.text for p in doc.paragraphs]
        # Should have kept the header + new content + next section header
        assert "三、实验步骤" in texts
        assert "new content" in texts
        assert "四、实验总结" in texts
        # old content should be removed
        assert "old step content" not in texts

    def test_fallback_number_boundary(self):
        """Without section_map, falls back to number-based boundary."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("三、实验步骤")
        doc.add_paragraph("old content")
        doc.add_paragraph("四、实验结果")

        paras = list(doc.paragraphs)
        _replace_section(paras, 0, "new content")

        texts = [p.text for p in doc.paragraphs]
        assert "new content" in texts
        assert "四、实验结果" in texts
        assert "old content" not in texts


# ── detect_sections (DA2 regression) ──


class TestDetectSectionsRegression:
    def test_standard_three_sections(self):
        """T3: standard 三/四/五 maps correctly."""
        doc = _load_doc(FIXTURES / "programming_lab.docx")
        paras = list(doc.paragraphs)
        sections, sm = detect_sections(paras)
        assert sm["steps"] is not None, "steps should be detected"
        assert sm["result"] is not None, "result should be detected"
        assert sm["summary"] is not None, "summary should be detected"

    def test_variant_four_sections(self):
        """T2: 四=实验总结, not mapped to result."""
        doc = _load_doc(FIXTURES / "variant_four_sections.docx")
        paras = list(doc.paragraphs)
        sections, sm = detect_sections(paras)
        assert sm["summary"] is not None, "summary should be detected (四=总结)"
        assert sm["summary"]["heading"] == "四、实验总结"
        assert sm["result"] is None, "result should be null (no independent result section)"

    def test_variant_three_sections(self):
        """T4: 二/三/四 map correctly."""
        doc = _load_doc(FIXTURES / "variant_three_sections.docx")
        paras = list(doc.paragraphs)
        sections, sm = detect_sections(paras)
        assert sm["steps"] is not None
        assert sm["result"] is not None
        assert sm["summary"] is not None


# ── fill_lab paragraph-based ──


class TestFillLabParagraph:
    def test_standard_fill(self, tmp_path):
        """Standard three-section report fills all sections correctly."""
        doc = _load_doc(FIXTURES / "programming_lab.docx")
        ans = _make_fake_ans(
            steps="分析：使用顺序结构。\n代码略。",
            result="运行结果：输出正确。",
            summary="掌握了顺序结构编程。",
        )
        fill_lab(doc, ans)
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "分析：使用顺序结构" in texts
        assert "运行结果：输出正确" in texts
        assert "掌握了顺序结构编程" in texts

    def test_fill_respects_fill_sections(self, tmp_path):
        """Only fill the specified sections."""
        doc = _load_doc(FIXTURES / "programming_lab.docx")
        ans = _make_fake_ans(
            steps="步骤内容",
            result="结果内容",
            summary="总结内容",
        )
        fill_lab(doc, ans, fill_sections=["steps"])
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "步骤内容" in texts
        assert "结果内容" not in texts
        assert "总结内容" not in texts

    def test_variant_four_sections_merge(self, tmp_path):
        """T2: No separate result → result content merges into summary."""
        doc = _load_doc(FIXTURES / "variant_four_sections.docx")
        ans = _make_fake_ans(
            steps="干涉实验步骤",
            result="干涉条纹清晰，波长测量准确。",
            summary="理解了光的干涉原理。",
        )
        fill_lab(doc, ans)
        texts = "\n".join(p.text for p in doc.paragraphs)
        # Steps should be in the document
        assert "干涉实验步骤" in texts
        # Summary section should contain both result + summary (merged)
        assert "干涉条纹清晰" in texts
        assert "理解了光的干涉原理" in texts


# ── fill_lab training_table ──


class TestFillLabTrainingTable:
    def test_training_table_fill(self, tmp_path):
        """Content goes into table cells for training_table layout."""
        doc = _load_doc(FIXTURES / "training_table.docx")
        ans = _make_fake_ans(
            steps="JSP文件上传代码：\n<%@ page ... %>\n配置web.xml...",
            result="运行截图见附件",
            summary="掌握了JSP文件上传技术",
        )
        metadata = {
            "report_layout": "training_table",
            "table_map": [
                {"table": 1, "row": 0, "col": 0, "label": "实训项目"},
                {"table": 1, "row": 1, "col": 0, "label": "实训任务"},
                {"table": 1, "row": 2, "col": 0, "label": "实训步骤及内容"},
                {"table": 1, "row": 3, "col": 0, "label": "实训总结"},
            ],
        }
        fill_lab(doc, ans, metadata=metadata)

        # Content should be in the value column (col 1) of the fill-target row
        cell_text = doc.tables[1].rows[2].cells[1].text
        assert "JSP" in cell_text or "文件上传" in cell_text
        assert "运行截图" in cell_text
        assert "掌握了" in cell_text

    def test_training_table_fill_no_targets(self, tmp_path):
        """When no fill-target cells found, should not crash."""
        doc = _load_doc(FIXTURES / "training_table.docx")
        ans = _make_fake_ans(steps="test")
        metadata = {
            "report_layout": "training_table",
            "table_map": [
                {"table": 1, "row": 0, "col": 0, "label": "实训项目"},
            ],
        }
        # Should not raise
        fill_lab(doc, ans, metadata=metadata)

    def test_training_table_respects_fill_sections(self, tmp_path):
        """Only fill allowed sections in training table."""
        doc = _load_doc(FIXTURES / "training_table.docx")
        ans = _make_fake_ans(
            steps="步骤ABC",
            result="结果DEF",
            summary="总结GHI",
        )
        metadata = {
            "report_layout": "training_table",
            "table_map": [
                {"table": 1, "row": 2, "col": 0, "label": "实训步骤及内容"},
            ],
        }
        fill_lab(doc, ans, fill_sections=["steps"], metadata=metadata)
        cell_text = doc.tables[1].rows[2].cells[1].text
        assert "步骤ABC" in cell_text
        assert "结果DEF" not in cell_text
        assert "总结GHI" not in cell_text

    def test_lab_report_table_fill(self, tmp_path):
        """超星式表格：实验名 / 实验目的 / 实验内容 分格写入。"""
        fixture = FIXTURES / "lab_report_table.docx"
        if not fixture.exists():
            pytest.skip("run python tests/generate_fixtures.py first")
        doc = _load_doc(fixture)
        ans = _make_fake_ans(
            steps="简单工厂模式实现…\npublic class Main { … }",
            result="运行输出见附件",
            summary="掌握了六种创建型模式",
        )
        ans["images_b64"] = [
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
        ]
        ans["uml_images_b64"] = [
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
        ]
        metadata = {
            "report_layout": "training_table",
            "experiment_title": "实验项目二：创建型设计模式实验",
            "assignment_text": DESIGN_PATTERN_ASSIGNMENT,
            "table_map": [
                {"table": 0, "row": 4, "col": 0, "label": "实验名"},
                {"table": 0, "row": 5, "col": 0, "label": "实验目的"},
                {"table": 0, "row": 6, "col": 0, "label": "实验内容"},
            ],
        }
        fill_lab(doc, ans, metadata=metadata)

        assert "创建型设计模式" in doc.tables[0].rows[4].cells[1].text
        objective_cell = doc.tables[0].rows[5].cells[1].text
        assert "熟练绘制创建型设计模式结构图" in objective_cell
        assert "简单工厂模式实现" not in objective_cell
        body = doc.tables[0].rows[6].cells[1].text
        assert "简单工厂模式" in body
        assert "运行输出" in body
        assert "掌握了" in body
        body_cell = doc.tables[0].rows[6].cells[1]
        # text + at least 2 image paragraphs (screenshot + uml)
        assert len(body_cell.paragraphs) >= 3


# ── do_fill integration ──


class TestDoFillIntegration:
    def test_standard_do_fill(self, tmp_path):
        """do_fill with standard docx produces output file."""
        out = tmp_path / "output.docx"
        ans = _make_fake_ans(
            steps="步骤",
            result="结果",
            summary="总结",
        )
        result = do_fill(FIXTURES / "programming_lab.docx", [ans], str(out))
        assert Path(result).exists()
        assert Path(result).suffix == ".docx"

        # Verify content
        from docx import Document
        doc = Document(str(result))
        texts = "\n".join(p.text for p in doc.paragraphs)
        assert "步骤" in texts
        assert "结果" in texts
        assert "总结" in texts

    def test_do_fill_with_training_metadata(self, tmp_path):
        """do_fill passes metadata through to fill_lab for training_table."""
        out = tmp_path / "output.docx"
        ans = _make_fake_ans(
            steps="JSP实现代码",
            result="部署成功",
            summary="完成实训",
        )
        metadata = {
            "report_layout": "training_table",
            "table_map": [
                {"table": 1, "row": 2, "col": 0, "label": "实训步骤及内容"},
            ],
        }
        result = do_fill(FIXTURES / "training_table.docx", [ans], str(out), metadata=metadata)
        assert Path(result).exists()

        from docx import Document
        doc = Document(str(result))
        cell_text = doc.tables[1].rows[2].cells[1].text
        assert "JSP实现代码" in cell_text
