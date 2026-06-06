"""
Phase 3b — PDF export / paired docx fill (no LLM).

Usage:
  python tests/test_phase3b_pdf_export.py
"""

import sys
import time
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src" / "python"))

from config import DOCX_OK, PDF_OK, TEMP_DIR  # noqa: E402
from agent.document_store import put_bundle, resolve_agent_context  # noqa: E402
from agent.parse_documents import parse_documents_list  # noqa: E402
from document.pdf_export import (  # noqa: E402
    generate_docx_shell,
    prepare_fill_docx_for_fill,
    resolve_fill_target_info,
)
from modules.fill_report import do_fill  # noqa: E402


def _make_pdf(path: Path) -> None:
    if not PDF_OK:
        return
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "Course: Networks\n三、实验步骤\n(fill here)\n四、实验结果\n五、实验总结\n",
        fontsize=11,
    )
    doc.save(str(path))
    doc.close()


def _make_docx(path: Path, body: str = "三、实验步骤\n\n四、实验结果\n\n五、实验总结\n") -> None:
    if not DOCX_OK:
        return
    from docx import Document

    doc = Document()
    for line in body.splitlines():
        doc.add_paragraph(line)
    doc.save(str(path))


def test_generate_docx_shell():
    if not DOCX_OK:
        print("SKIP: docx unavailable")
        return
    out = TEMP_DIR / "test_phase3b_shell.docx"
    path = generate_docx_shell("三、实验步骤\n四、实验结果\n五、实验总结", {"course": "OS"}, out)
    assert path.exists()
    from docx import Document

    texts = [p.text for p in Document(str(path)).paragraphs]
    assert any("三" in t for t in texts)
    assert any("OS" in t for t in texts)


def test_prepare_fill_pdf_generates_docx():
    if not DOCX_OK:
        print("SKIP: docx unavailable")
        return
    inp, info = prepare_fill_docx_for_fill(
        None,
        "report.pdf",
        source_format="pdf",
        fill_body_text="三、实验步骤\n四、实验结果\n五、实验总结",
    )
    assert inp is not None
    assert inp.suffix == ".docx"
    assert info["from"] == "generated"
    assert info["source_format"] == "pdf"


def test_prepare_fill_docx_unchanged():
    if not DOCX_OK:
        print("SKIP: docx unavailable")
        return
    path = TEMP_DIR / "test_phase3b_plain.docx"
    _make_docx(path)
    inp, info = prepare_fill_docx_for_fill(path, "plain.docx", source_format="docx")
    assert inp == path
    assert info["from"] == "docx"


def test_paired_docx_template():
    if not DOCX_OK or not PDF_OK:
        print("SKIP: docx/pdf unavailable")
        return
    pdf = TEMP_DIR / "test_phase3b_pair.pdf"
    docx = TEMP_DIR / "test_phase3b_pair_tpl.docx"
    _make_pdf(pdf)
    _make_docx(docx)

    inp, info = prepare_fill_docx_for_fill(
        pdf,
        "pair.pdf",
        source_format="pdf",
        paired_docx_path=docx,
        fill_body_text="三、实验步骤",
    )
    assert inp == docx
    assert info["from"] == "user_template"


def test_resolve_fill_target_pdf_with_template():
    pdf_bundle = {
        "document_id": "p1",
        "role": "fill_target",
        "file_name": "lab.pdf",
        "file_path": "/tmp/lab.pdf",
        "metadata": {"source_format": "pdf"},
    }
    tpl_bundle = {
        "document_id": "t1",
        "role": "fill_template",
        "file_name": "blank.docx",
        "file_path": "/tmp/blank.docx",
        "metadata": {"source_format": "docx"},
    }
    info = resolve_fill_target_info([pdf_bundle, tpl_bundle], pdf_bundle)
    assert info["from"] == "user_template"
    assert info["path"] == "/tmp/blank.docx"


def test_multi_doc_pdf_fill_template():
    if not DOCX_OK or not PDF_OK:
        print("SKIP: docx/pdf unavailable")
        return
    import base64

    pdf = TEMP_DIR / "test_phase3b_pdf_target.pdf"
    docx = TEMP_DIR / "test_phase3b_tpl.docx"
    _make_pdf(pdf)
    _make_docx(docx)

    parsed = parse_documents_list(
        [
            {
                "id": "p1",
                "role": "fill_target",
                "file_name": "report.pdf",
                "file_data": base64.b64encode(pdf.read_bytes()).decode(),
            },
            {
                "id": "t1",
                "role": "fill_template",
                "file_name": "blank.docx",
                "file_data": base64.b64encode(docx.read_bytes()).decode(),
            },
        ]
    )
    assert parsed["fill_target_info"]["from"] == "user_template"
    assert parsed["fill_target_info"]["path"] == parsed["_bundles"][1]["file_path"]
    assert parsed["fill_target"]["fill_docx_from"] == "user_template"
    assert parsed["fill_target"]["source_format"] == "pdf"


def test_multi_doc_assignment_pdf_fill_docx():
    if not DOCX_OK or not PDF_OK:
        print("SKIP: docx/pdf unavailable")
        return
    import base64

    pdf = TEMP_DIR / "test_phase3b_assign.pdf"
    docx = TEMP_DIR / "test_phase3b_fill.docx"
    _make_pdf(pdf)
    _make_docx(docx)

    parsed = parse_documents_list(
        [
            {
                "id": "a1",
                "role": "assignment",
                "file_name": "assign.pdf",
                "file_data": base64.b64encode(pdf.read_bytes()).decode(),
            },
            {
                "id": "f1",
                "role": "fill_target",
                "file_name": "blank.docx",
                "file_data": base64.b64encode(docx.read_bytes()).decode(),
            },
        ]
    )
    assert parsed["fill_target"]["source_format"] == "docx"
    assert parsed["fill_target"]["export_format"] == "docx"
    assert parsed.get("assignment_text")


def test_fill_report_pdf_only_payload():
    if not DOCX_OK:
        print("SKIP: docx unavailable")
        return
    answers = [
        {
            "type": "lab_report",
            "parsed": {
                "steps_analysis": "步骤说明",
                "result_description": "结果",
                "summary": "总结",
            },
        }
    ]
    inp, info = prepare_fill_docx_for_fill(
        None,
        "only.pdf",
        source_format="pdf",
        fill_body_text="三、实验步骤\n四、实验结果\n五、实验总结",
    )
    out = do_fill(inp, answers, str(TEMP_DIR / "test_phase3b_filled.docx"))
    assert Path(out).exists()
    assert info["from"] == "generated"


def test_agent_context_pdf_template():
    b_pdf = {
        "document_id": "p1",
        "role": "fill_target",
        "file_name": "r.pdf",
        "file_path": str(TEMP_DIR / "r.pdf"),
        "report_text": "三、实验步骤",
        "fill_body_text": "三、实验步骤",
        "full_text": "三、实验步骤",
        "metadata": {"source_format": "pdf"},
        "question": {"type": "lab_report"},
        "warnings": [],
        "created_at": time.time(),
    }
    b_docx = {
        "document_id": "t1",
        "role": "fill_template",
        "file_name": "t.docx",
        "file_path": str(TEMP_DIR / "t.docx"),
        "metadata": {"source_format": "docx"},
        "created_at": time.time(),
    }
    put_bundle(b_pdf)
    put_bundle(b_docx)
    ctx = resolve_agent_context(["p1", "t1"])
    assert ctx["fill_target_info"]["from"] == "user_template"
    assert ctx["fill_target_info"]["path"] == b_docx["file_path"]


def main():
    test_generate_docx_shell()
    test_prepare_fill_pdf_generates_docx()
    test_prepare_fill_docx_unchanged()
    test_paired_docx_template()
    test_resolve_fill_target_pdf_with_template()
    test_multi_doc_pdf_fill_template()
    test_multi_doc_assignment_pdf_fill_docx()
    test_fill_report_pdf_only_payload()
    test_agent_context_pdf_template()
    print("test_phase3b_pdf_export: OK")


if __name__ == "__main__":
    main()
