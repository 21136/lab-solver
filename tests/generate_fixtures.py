"""Generate golden docx fixtures (run once: python tests/generate_fixtures.py)."""

from pathlib import Path

from docx import Document

FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"


def _cover_table(doc, course, exp_title):
    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("课程名称", course),
        ("实验名称", exp_title),
        ("专业", "计算机科学与技术"),
        ("姓名", "测试学生"),
    ]
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v


def programming_lab():
    doc = Document()
    _cover_table(doc, "Java 程序设计", "实验一 顺序结构程序设计")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握 Java 基本语法，编写并运行顺序结构程序。")
    doc.add_paragraph("二、实验内容")
    doc.add_paragraph("使用 Java 编写程序，输入两个整数并输出它们的和。")
    doc.add_paragraph("三、实验内容及步骤")
    doc.add_paragraph("1. 创建工程并编写 main 方法；2. 编译运行并截图。")
    doc.add_paragraph("四、实验结果")
    doc.add_paragraph("（填写运行截图与输出）")
    doc.add_paragraph("五、实验总结")
    doc.add_paragraph("（填写心得）")
    return doc


def theory_lab():
    doc = Document()
    _cover_table(doc, "计算机网络", "实验二 协议分析")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("理解 TCP 三次握手过程，并能用文字描述各阶段报文特征。")
    doc.add_paragraph("二、实验原理")
    doc.add_paragraph("TCP 建立连接需要 SYN、SYN-ACK、ACK 三次交互。")
    doc.add_paragraph("三、实验步骤")
    doc.add_paragraph("使用 Wireshark 抓包并分析握手报文。")
    doc.add_paragraph("四、实验结果与分析")
    doc.add_paragraph("根据抓包结果论述握手时序与标志位含义。")
    doc.add_paragraph("五、思考题")
    doc.add_paragraph("为何需要三次握手而不是两次？")
    return doc


def combined_lab():
    doc = Document()
    _cover_table(doc, "软件工程", "实验三 需求分析与原型")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("完成需求分析文档，并用 Python 实现简易原型界面。")
    doc.add_paragraph("二、实验原理")
    doc.add_paragraph("结合用例图与活动图描述系统流程。")
    doc.add_paragraph("三、实验内容及步骤")
    doc.add_paragraph("第一部分：撰写需求说明；第二部分：编写 Python 原型并运行截图。")
    doc.add_paragraph("四、实验结果")
    doc.add_paragraph("提交文档与程序运行截图。")
    doc.add_paragraph("五、实验总结")
    doc.add_paragraph("总结需求分析与实现过程中的问题。")
    return doc


SAMPLE_CLOZE_LINES = [
    "public class MainControllerCenter {",
    "    ( 1 ) MainControllerCenter instance;",
    "    private ( 2 ) MainControllerCenter() {}",
    "    public static MainControllerCenter getInstance() {",
    "        if (instance == null) {",
    "            instance = ( 3 ) MainControllerCenter();",
    "        }",
    "        return instance;",
    "    }",
    "}",
]


def mixed_theory_cloze_paste_text() -> str:
    """Plain-text golden for split_text_assignment_segments (aligned with mixed_theory_cloze docx)."""
    cloze = "\n".join(SAMPLE_CLOZE_LINES)
    return (
        "一、简答题\n"
        "1. 简要说明外观模式（Facade）的意图及适用场景，并举例说明子系统如何被封装。\n"
        "2. 比较工厂方法模式与抽象工厂模式的异同，说明各自适用于何种产品族扩展场景。\n"
        "3. 简述单例模式懒汉式与饿汉式实现的线程安全差异及适用条件。\n"
        "二、代码填空\n"
        "阅读下列 Java 代码，完成 1–3 空（Singleton）。\n"
        f"{cloze}"
    )


def mixed_theory_cloze():
    """O10/R8: theory short answers + Singleton code cloze in one exam docx."""
    doc = Document()
    doc.add_paragraph("软件工程期末复习 — 混排样卷")
    doc.add_paragraph("一、简答题")
    doc.add_paragraph("1. 简要说明外观模式（Facade）的意图及适用场景。")
    doc.add_paragraph("2. 比较工厂方法模式与抽象工厂模式的异同。")
    doc.add_paragraph("二、代码填空")
    doc.add_paragraph("阅读下列 Java 代码，完成 1–3 空（Singleton）。")
    table = doc.add_table(rows=0, cols=1)
    for line in SAMPLE_CLOZE_LINES:
        table.add_row().cells[0].text = line
    return doc


def code_cloze_singleton():
    """Singleton 代码填空 — docx import golden (R5 / Phase D)."""
    doc = Document()
    doc.add_paragraph("设计模式复习 — Singleton 代码填空")
    table = doc.add_table(rows=0, cols=1)
    for line in SAMPLE_CLOZE_LINES:
        table.add_row().cells[0].text = line
    return doc


def training_table():
    """Simulate a 实训周 table-based report where the task is inside a table cell."""
    doc = Document()
    _cover_table(doc, "Java Web 开发", "第十周实训")
    doc.add_paragraph("实训周报告 —— 请按表格内容完成")
    doc.add_paragraph("")

    table = doc.add_table(rows=4, cols=2)
    rows_data = [
        ("实训项目", "JSP 文件上传与下载"),
        ("实训任务", "新建 JSP10 项目，配置 Tomcat，实现 FileUpload 功能并截图"),
        ("实训步骤及内容", "1. 创建 Dynamic Web Project\n2. 编写 upload.jsp 与 DownloadServlet\n3. 部署到 Tomcat 9\n4. 浏览器测试并截图"),
        ("实训总结", "（请在此填写实训心得与遇到的问题）"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    return doc


def variant_four_sections():
    """T2: 一至四节，四=实验总结，无第五节。V1 会把四误映射到 result。"""
    doc = Document()
    _cover_table(doc, "大学物理", "实验四 光的干涉")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("观察光的干涉现象，测量光波波长。")
    doc.add_paragraph("二、实验原理")
    doc.add_paragraph("两列相干光相遇时产生干涉条纹，条纹间距与波长有关。")
    doc.add_paragraph("三、实验内容及步骤")
    doc.add_paragraph("1. 调整迈克尔逊干涉仪；2. 测量干涉条纹间距；3. 记录数据。")
    doc.add_paragraph("四、实验总结")
    doc.add_paragraph("（填写实验心得与思考题答案）")
    return doc


def variant_three_sections():
    """T4: 三节（二/三/四），步骤+结果+总结。不同节号同样需要正确映射。"""
    doc = Document()
    _cover_table(doc, "数字电路", "实验二 组合逻辑电路")
    doc.add_paragraph("一、实验目的")
    doc.add_paragraph("掌握组合逻辑电路的设计方法。")
    doc.add_paragraph("二、实验步骤")
    doc.add_paragraph("1. 根据真值表设计电路；2. 在实验箱上连线验证；3. 记录输入输出波形。")
    doc.add_paragraph("三、实验结果")
    doc.add_paragraph("（填写真值表与波形图）")
    doc.add_paragraph("四、实验总结")
    doc.add_paragraph("（填写实验心得）")
    return doc


def lab_report_table():
    """Table-only lab report (超星常见): 实验名 / 实验目的 / 实验内容，无三/四/五段落。"""
    doc = Document()
    doc.add_paragraph("“设计模式”实验报告")
    table = doc.add_table(rows=8, cols=4)
    rows_data = [
        ("学号", "", "姓名", ""),
        ("班级", "", "专业", ""),
        ("实验日期", "", "指导老师", ""),
        ("实验环境", "", "", ""),
        ("实验名", "", "", ""),
        ("实验目的", "", "", ""),
        ("实验内容", "", "", ""),
        ("", "", "", ""),
    ]
    for i, cells in enumerate(rows_data):
        for j, text in enumerate(cells):
            table.rows[i].cells[j].text = text
    return doc


def main():
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    specs = [
        ("programming_lab.docx", programming_lab),
        ("theory_lab.docx", theory_lab),
        ("combined_lab.docx", combined_lab),
        ("training_table.docx", training_table),
        ("lab_report_table.docx", lab_report_table),
        ("variant_four_sections.docx", variant_four_sections),
        ("variant_three_sections.docx", variant_three_sections),
        ("code_cloze_singleton.docx", code_cloze_singleton),
        ("mixed_theory_cloze.docx", mixed_theory_cloze),
    ]
    for name, builder in specs:
        out = FIXTURES_DIR / name
        builder().save(out)
        print(f"wrote {out}")


if __name__ == "__main__":
    main()
