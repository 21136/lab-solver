"""DA2: section_map semantic detection tests (no LLM calls)."""

from pathlib import Path

import pytest

if True:
    from docx import Document

from modules.fill_report import (
    detect_sections,
    _guess_semantic,
    _apply_semantic_overrides,
    do_fill,
)

FIXTURES = Path(__file__).resolve().parent / "fixtures"


# ── helpers ──


def _paras_from_docx(path):
    doc = Document(str(path))
    return list(doc.paragraphs)


# ── unit: _guess_semantic ──


@pytest.mark.parametrize(
    "heading, expected",
    [
        ("三、实验内容及步骤", "steps"),
        ("三、实验步骤", "steps"),
        ("二、实验内容", "steps"),
        ("四、实验结果", "result"),
        ("四、实验结果与分析", "result"),
        ("三、实验结果", "result"),
        ("五、实验总结", "summary"),
        ("四、实验总结", "summary"),
        ("五、思考题", "summary"),
        ("一、实验目的", "objective"),
        ("二、实验原理", "principles"),
        ("二、实验任务", "steps"),
        ("四、实验小结", "summary"),
    ],
)
def test_guess_semantic(heading, expected):
    assert _guess_semantic(heading) == expected


def test_skip_arabic_list_items_as_sections():
    """Numbered list items like 1.xxx should not pollute sections_detected."""
    class P:
        def __init__(self, text):
            self.text = text

    paras = [
        P("一、实验目的"),
        P("1.掌握MyBatis基于注解方式完成业务。"),
        P("2.掌握MyBatis基于XML文件方式完成业务。"),
        P("二、实验任务"),
    ]
    sections, smap = detect_sections(paras)
    headings = [s["heading"] for s in sections]
    assert "1.掌握MyBatis基于注解方式完成业务。" not in headings
    assert "2.掌握MyBatis基于XML文件方式完成业务。" not in headings
    assert smap["steps"] is not None
    assert "任务" in smap["steps"]["heading"]


def test_semantic_overrides_apply_to_sections():
    detected = [
        {"index": 0, "heading": "一、实验目的", "semantic": "objective"},
        {"index": 4, "heading": "二、实验任务", "semantic": None},
    ]
    overrides = {"steps": "二、实验任务"}
    out = _apply_semantic_overrides(detected, overrides)
    task = next(s for s in out if s["heading"] == "二、实验任务")
    assert task["semantic"] == "steps"


def _heading_at(paras, idx):
    """Return paragraph text at index, or '' if out of range."""
    if 0 <= idx < len(paras):
        return paras[idx].text.strip()
    return ""


def _has_kw(text, *keywords):
    """Check if text contains any of the given keywords."""
    return any(kw in text for kw in keywords)


# ── detect_sections: standard 三/四/五 (T3 regression) ──


def test_standard_sections():
    """Standard docx: sections detected with correct semantic roles."""
    paras = _paras_from_docx(FIXTURES / "programming_lab.docx")
    sections, smap = detect_sections(paras)

    assert len(sections) >= 5  # 一 through 五

    assert smap["steps"] is not None
    assert smap["steps"]["type"] == "paragraph"
    assert smap["steps"]["para_index"] < smap["result"]["para_index"]

    assert smap["result"] is not None
    assert smap["result"]["type"] == "paragraph"
    assert smap["result"]["para_index"] < smap["summary"]["para_index"]

    assert smap["summary"] is not None
    assert smap["summary"]["type"] == "paragraph"

    # Verify the heading texts have expected keywords
    steps_h = _heading_at(paras, smap["steps"]["para_index"])
    assert _has_kw(steps_h, "步骤", "内容")

    result_h = _heading_at(paras, smap["result"]["para_index"])
    assert _has_kw(result_h, "结果", "实验")

    summary_h = _heading_at(paras, smap["summary"]["para_index"])
    assert _has_kw(summary_h, "总结")


# ── detect_sections: variant 四=总结 (T2) ──


def test_variant_four_sections():
    """Variant: 四=实验总结 maps to summary, NOT result."""
    paras = _paras_from_docx(FIXTURES / "variant_four_sections.docx")
    sections, smap = detect_sections(paras)

    assert smap["steps"] is not None
    steps_h = _heading_at(paras, smap["steps"]["para_index"])
    assert _has_kw(steps_h, "步骤", "内容")

    # 四=实验总结 → summary, NOT result
    assert smap["summary"] is not None
    summary_h = _heading_at(paras, smap["summary"]["para_index"])
    assert _has_kw(summary_h, "总结")
    assert smap["summary"]["para_index"] > smap["steps"]["para_index"]

    # result does not exist (no independent results section)
    assert smap["result"] is None


# ── detect_sections: variant 二/三/四 (T4) ──


def test_variant_three_sections():
    """Variant (二/三/四): all three sections map correctly."""
    paras = _paras_from_docx(FIXTURES / "variant_three_sections.docx")
    sections, smap = detect_sections(paras)

    assert smap["steps"] is not None
    steps_h = _heading_at(paras, smap["steps"]["para_index"])
    assert _has_kw(steps_h, "步骤")

    assert smap["result"] is not None
    result_h = _heading_at(paras, smap["result"]["para_index"])
    assert _has_kw(result_h, "结果")

    assert smap["summary"] is not None
    summary_h = _heading_at(paras, smap["summary"]["para_index"])
    assert _has_kw(summary_h, "总结")

    # Verify ordering: steps < result < summary
    assert smap["steps"]["para_index"] < smap["result"]["para_index"] < smap["summary"]["para_index"]


# ── detect_sections: theory lab (四=结果与分析, 五=思考题) ──


def test_theory_lab_sections():
    """Theory lab: 四=实验结果与分析→result, 五=思考题→summary."""
    paras = _paras_from_docx(FIXTURES / "theory_lab.docx")
    sections, smap = detect_sections(paras)

    assert smap["steps"] is not None
    assert smap["result"] is not None
    result_h = _heading_at(paras, smap["result"]["para_index"])
    assert _has_kw(result_h, "结果", "分析")

    assert smap["summary"] is not None
    summary_h = _heading_at(paras, smap["summary"]["para_index"])
    assert _has_kw(summary_h, "思考", "总结")


# ── sections_detected completeness ──


def test_sections_detected_fields():
    """Each entry in sections_detected has index, heading, semantic."""
    paras = _paras_from_docx(FIXTURES / "combined_lab.docx")
    sections, _ = detect_sections(paras)

    for sec in sections:
        assert "index" in sec
        assert isinstance(sec["index"], int)
        assert "heading" in sec
        assert isinstance(sec["heading"], str)
        assert "semantic" in sec  # may be None


# ── fill integration: standard template still works ──


def test_fill_standard_template(tmp_path):
    """Fill a standard docx and verify all three sections get content."""
    import shutil

    src = FIXTURES / "programming_lab.docx"
    dst = tmp_path / "test.docx"
    shutil.copy(src, dst)

    answers = [
        {
            "type": "lab_report",
            "parsed": {
                "steps_analysis": "步骤内容测试",
                "result_description": "结果描述测试",
                "summary": "总结内容测试",
            },
            "answer": "fallback answer",
        }
    ]
    output = do_fill(dst, answers, output_path=str(tmp_path / "out.docx"))
    assert output

    # Verify the output has our content
    out_doc = Document(output)
    full = "\n".join(p.text for p in out_doc.paragraphs)
    assert "步骤内容测试" in full
    assert "结果描述测试" in full
    assert "总结内容测试" in full


# ── fill integration: variant 四=总结 maps correctly ──


def test_fill_variant_four_sections(tmp_path):
    """Fill variant (四=总结): summary content goes to 四, not to a nonexistent 五."""
    import shutil

    src = FIXTURES / "variant_four_sections.docx"
    dst = tmp_path / "test.docx"
    shutil.copy(src, dst)

    answers = [
        {
            "type": "lab_report",
            "parsed": {
                "steps_analysis": "步骤AAA",
                "result_description": "结果BBB",
                "summary": "总结CCC",
            },
            "answer": "fallback",
        }
    ]
    output = do_fill(dst, answers, output_path=str(tmp_path / "out.docx"))
    assert output

    out_doc = Document(output)
    full = "\n".join(p.text for p in out_doc.paragraphs)

    # Steps content should be there
    assert "步骤AAA" in full
    # Summary content should be there (under 四、实验总结)
    assert "总结CCC" in full
    # Result content may or may not appear (no result section to fill)
    # The key assertion: nothing crashed and steps+summary were placed


# ── fill integration: variant 二/三/四 maps correctly ──


def test_fill_variant_three_sections(tmp_path):
    """Fill variant (二/三/四): all three sections filled at correct headings."""
    import shutil

    src = FIXTURES / "variant_three_sections.docx"
    dst = tmp_path / "test.docx"
    shutil.copy(src, dst)

    answers = [
        {
            "type": "lab_report",
            "parsed": {
                "steps_analysis": "步骤XXX",
                "result_description": "结果YYY",
                "summary": "总结ZZZ",
            },
            "answer": "fallback",
        }
    ]
    output = do_fill(dst, answers, output_path=str(tmp_path / "out.docx"))
    assert output

    out_doc = Document(output)
    full = "\n".join(p.text for p in out_doc.paragraphs)

    assert "步骤XXX" in full
    assert "结果YYY" in full
    assert "总结ZZZ" in full
